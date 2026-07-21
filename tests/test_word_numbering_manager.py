from __future__ import annotations

import unittest

from docx import Document as WordDocument
from docx.oxml.ns import qn

from src.exporter.word_numbering_manager import (
    WordNumberingManager,
)
from src.models.list_item import (
    ListMarkerKind,
    ListMarkerSource,
)
from src.models.list_sequence import (
    ListContainerType,
    ListSequence,
    ListSequenceItem,
)


def make_number_sequence(
    sequence_id: int = 1,
    start_at: int = 1,
) -> ListSequence:
    sequence = ListSequence(
        sequence_id=sequence_id,
        page_number=1,
        list_type="number",
        container_type=(
            ListContainerType.PAGE_BODY
        ),
        container_id=1,
        container_left=50.0,
        container_right=550.0,
        start_at=start_at,
        maximum_level=2,
    )

    sequence.items.extend([
        ListSequenceItem(
            page_number=1,
            paragraph_region_number=1,
            item_index=0,
            level=0,
            marker="1.",
            marker_kind=(
                ListMarkerKind.DECIMAL
            ),
            marker_source=(
                ListMarkerSource.TEXT
            ),
            indent=80.0,
            numeric_value=1,
        ),

        ListSequenceItem(
            page_number=1,
            paragraph_region_number=2,
            item_index=1,
            level=1,
            marker="a)",
            marker_kind=(
                ListMarkerKind.LOWER_ALPHA
            ),
            marker_source=(
                ListMarkerSource.TEXT
            ),
            indent=115.0,
            numeric_value=1,
        ),

        ListSequenceItem(
            page_number=1,
            paragraph_region_number=3,
            item_index=2,
            level=2,
            marker="(i)",
            marker_kind=(
                ListMarkerKind.LOWER_ROMAN
            ),
            marker_source=(
                ListMarkerSource.TEXT
            ),
            indent=150.0,
            numeric_value=1,
        ),
    ])

    return sequence


def make_bullet_sequence() -> ListSequence:
    sequence = ListSequence(
        sequence_id=1,
        page_number=1,
        list_type="bullet",
        container_type=(
            ListContainerType.PAGE_BODY
        ),
        container_id=1,
        container_left=50.0,
        container_right=550.0,
        maximum_level=1,
    )

    sequence.items.extend([
        ListSequenceItem(
            page_number=1,
            paragraph_region_number=1,
            item_index=0,
            level=0,
            marker="•",
            marker_kind=(
                ListMarkerKind.BULLET
            ),
            marker_source=(
                ListMarkerSource.TEXT
            ),
            indent=80.0,
        ),

        ListSequenceItem(
            page_number=1,
            paragraph_region_number=2,
            item_index=1,
            level=1,
            marker="◦",
            marker_kind=(
                ListMarkerKind.BULLET
            ),
            marker_source=(
                ListMarkerSource.TEXT
            ),
            indent=115.0,
        ),
    ])

    return sequence


class WordNumberingManagerTests(
    unittest.TestCase
):

    def _resolve_abstract_number(
        self,
        manager: WordNumberingManager,
        number_id: int,
    ):
        number_element = next(
            element
            for element in (
                manager.numbering_element
                .findall(
                    qn("w:num")
                )
            )
            if int(
                element.get(
                    qn("w:numId")
                )
            )
            == number_id
        )

        abstract_reference = (
            number_element.find(
                qn("w:abstractNumId")
            )
        )

        abstract_number_id = int(
            abstract_reference.get(
                qn("w:val")
            )
        )

        return next(
            element
            for element in (
                manager.numbering_element
                .findall(
                    qn("w:abstractNum")
                )
            )
            if int(
                element.get(
                    qn(
                        "w:abstractNumId"
                    )
                )
            )
            == abstract_number_id
        )

    def test_multilevel_sequence_creates_three_levels(
        self,
    ) -> None:
        word_document = WordDocument()

        manager = (
            WordNumberingManager(
                word_document
            )
        )

        number_id = manager.create_sequence(
            make_number_sequence()
        )

        abstract_number = (
            self._resolve_abstract_number(
                manager,
                number_id,
            )
        )

        levels = abstract_number.findall(
            qn("w:lvl")
        )

        self.assertEqual(
            len(levels),
            3,
        )

        multi_level_type = (
            abstract_number.find(
                qn("w:multiLevelType")
            )
        )

        self.assertEqual(
            multi_level_type.get(
                qn("w:val")
            ),
            "multilevel",
        )

    def test_number_formats_follow_marker_kinds(
        self,
    ) -> None:
        word_document = WordDocument()

        manager = (
            WordNumberingManager(
                word_document
            )
        )

        number_id = manager.create_sequence(
            make_number_sequence()
        )

        abstract_number = (
            self._resolve_abstract_number(
                manager,
                number_id,
            )
        )

        levels = abstract_number.findall(
            qn("w:lvl")
        )

        formats = [
            level.find(
                qn("w:numFmt")
            ).get(
                qn("w:val")
            )
            for level in levels
        ]

        self.assertEqual(
            formats,
            [
                "decimal",
                "lowerLetter",
                "lowerRoman",
            ],
        )

    def test_marker_punctuation_is_preserved(
        self,
    ) -> None:
        word_document = WordDocument()

        manager = WordNumberingManager(
            word_document
        )

        number_id = manager.create_sequence(
            make_number_sequence()
        )

        abstract_number = (
            self._resolve_abstract_number(
                manager,
                number_id,
            )
        )

        level_texts = [
            level.find(
                qn("w:lvlText")
            ).get(
                qn("w:val")
            )
            for level in (
                abstract_number.findall(
                    qn("w:lvl")
                )
            )
        ]

        self.assertEqual(
            level_texts,
            [
                "%1.",
                "%2)",
                "(%3)",
            ],
        )

    def test_bullet_characters_are_preserved(
        self,
    ) -> None:
        word_document = WordDocument()

        manager = WordNumberingManager(
            word_document
        )

        number_id = manager.create_sequence(
            make_bullet_sequence()
        )

        abstract_number = (
            self._resolve_abstract_number(
                manager,
                number_id,
            )
        )

        level_texts = [
            level.find(
                qn("w:lvlText")
            ).get(
                qn("w:val")
            )
            for level in (
                abstract_number.findall(
                    qn("w:lvl")
                )
            )
        ]

        self.assertEqual(
            level_texts,
            [
                "•",
                "◦",
            ],
        )

    def test_separate_sequences_receive_unique_num_ids(
        self,
    ) -> None:
        word_document = WordDocument()

        manager = WordNumberingManager(
            word_document
        )

        first_number_id = (
            manager.create_sequence(
                make_number_sequence(
                    sequence_id=1
                )
            )
        )

        second_number_id = (
            manager.create_sequence(
                make_number_sequence(
                    sequence_id=2
                )
            )
        )

        self.assertNotEqual(
            first_number_id,
            second_number_id,
        )

    def test_compatible_sequences_reuse_abstract_definition(
        self,
    ) -> None:
        word_document = WordDocument()

        manager = WordNumberingManager(
            word_document
        )

        initial_abstract_count = len(
            manager.numbering_element.findall(
                qn("w:abstractNum")
            )
        )

        manager.create_sequence(
            make_number_sequence(
                sequence_id=1
            )
        )

        after_first_count = len(
            manager.numbering_element.findall(
                qn("w:abstractNum")
            )
        )

        manager.create_sequence(
            make_number_sequence(
                sequence_id=2
            )
        )

        after_second_count = len(
            manager.numbering_element.findall(
                qn("w:abstractNum")
            )
        )

        self.assertGreater(
            after_first_count,
            initial_abstract_count,
        )

        self.assertEqual(
            after_second_count,
            after_first_count,
        )

    def test_top_level_start_override_is_created(
        self,
    ) -> None:
        word_document = WordDocument()

        manager = WordNumberingManager(
            word_document
        )

        number_id = manager.create_sequence(
            make_number_sequence(
                start_at=4
            )
        )

        number_element = next(
            element
            for element in (
                manager.numbering_element
                .findall(
                    qn("w:num")
                )
            )
            if int(
                element.get(
                    qn("w:numId")
                )
            )
            == number_id
        )

        override = (
            number_element.find(
                qn("w:lvlOverride")
            )
        )

        self.assertIsNotNone(
            override
        )

        self.assertEqual(
            override.get(
                qn("w:ilvl")
            ),
            "0",
        )

        start_override = (
            override.find(
                qn("w:startOverride")
            )
        )

        self.assertEqual(
            start_override.get(
                qn("w:val")
            ),
            "4",
        )

    def test_abstract_numbers_remain_before_number_instances(
        self,
    ) -> None:
        word_document = WordDocument()

        manager = WordNumberingManager(
            word_document
        )

        manager.create_sequence(
            make_number_sequence()
        )

        children = list(
            manager.numbering_element
        )

        abstract_indexes = [
            index
            for index, child
            in enumerate(children)
            if child.tag
            == qn("w:abstractNum")
        ]

        number_indexes = [
            index
            for index, child
            in enumerate(children)
            if child.tag
            == qn("w:num")
        ]

        self.assertTrue(
            abstract_indexes
        )

        self.assertTrue(
            number_indexes
        )

        self.assertLess(
            max(
                abstract_indexes
            ),
            min(
                number_indexes
            ),
        )

    def test_unicode_bullet_does_not_use_symbol_font(
        self,
    ) -> None:
        word_document = WordDocument()
    
        manager = WordNumberingManager(
            word_document
        )
    
        number_id = manager.create_sequence(
            sequence=make_bullet_sequence(),
            marker_font_name="Symbol",
            marker_font_size=11.0,
        )
    
        abstract_number = (
            self._resolve_abstract_number(
                manager,
                number_id,
            )
        )
    
        level = abstract_number.find(
            qn("w:lvl")
        )
    
        run_properties = level.find(
            qn("w:rPr")
        )
    
        fonts = run_properties.find(
            qn("w:rFonts")
        )
    
        self.assertEqual(
            fonts.get(
                qn("w:ascii")
            ),
            "Arial",
        )
    
        self.assertEqual(
            fonts.get(
                qn("w:hAnsi")
            ),
            "Arial",
        )


if __name__ == "__main__":
    unittest.main()