from __future__ import annotations

import unittest

from docx import (
    Document,
)
from docx.oxml.ns import (
    qn,
)

from src.exporter.editable_floating_image_renderer import (
    EditableFloatingImageRenderer,
    EditableFloatingImageWrapMode,
)
from src.models.editable_image import (
    EditableImage,
    EditableImageDisposition,
    EditableImageExtractionMode,
    EditableImagePayloadStatus,
    EditableImagePlacement,
    EditableImageRole,
)
from src.models.geometry.rectangle import (
    Rectangle,
)


def create_test_png() -> bytes:
    try:
        import pymupdf

    except ImportError:
        import fitz as pymupdf

    pixmap = pymupdf.Pixmap(
        pymupdf.csRGB,
        pymupdf.IRect(
            0,
            0,
            4,
            2,
        ),
        False,
    )

    pixmap.clear_with(
        255
    )

    return pixmap.tobytes(
        "png"
    )


PNG_BYTES = create_test_png()


def make_image(
    **overrides,
) -> EditableImage:
    values = {
        "page_number": 1,
        "image_id": "image:1:1",
        "bbox": Rectangle(
            left=72.0,
            top=144.0,
            right=272.0,
            bottom=244.0,
        ),
        "extraction_mode": (
            EditableImageExtractionMode
            .DIRECT_BYTES
        ),
        "disposition": (
            EditableImageDisposition.NATIVE
        ),
        "placement": (
            EditableImagePlacement.FLOATING
        ),
        "role": EditableImageRole.CONTENT,
        "payload": PNG_BYTES,
        "payload_status": (
            EditableImagePayloadStatus.READY
        ),
        "payload_mime_type": "image/png",
        "extension": "png",
        "pixel_width": 4,
        "pixel_height": 2,
        "confidence": 0.95,
        "payload_confidence": 0.95,
    }

    values.update(
        overrides
    )

    return EditableImage(
        **values
    )


def make_page_bbox() -> Rectangle:
    return Rectangle(
        left=0.0,
        top=0.0,
        right=600.0,
        bottom=800.0,
    )


class EditableFloatingImageRendererTests(
    unittest.TestCase
):

    def test_floating_image_creates_anchor(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableFloatingImageRenderer
            .render(
                container=document,
                image=make_image(),
                page_bbox=make_page_bbox(),
            )
        )

        self.assertEqual(
            result.anchor.tag,
            qn(
                "wp:anchor"
            ),
        )

        self.assertIn(
            "wp:anchor",
            result.paragraph._p.xml,
        )

        self.assertNotIn(
            "wp:inline",
            result.paragraph._p.xml,
        )

    def test_page_relative_offsets_are_written(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableFloatingImageRenderer
            .render(
                container=document,
                image=make_image(),
                page_bbox=make_page_bbox(),
            )
        )

        horizontal_position = (
            result.anchor.find(
                qn(
                    "wp:positionH"
                )
            )
        )

        vertical_position = (
            result.anchor.find(
                qn(
                    "wp:positionV"
                )
            )
        )

        horizontal_offset = (
            horizontal_position.find(
                qn(
                    "wp:posOffset"
                )
            )
        )

        vertical_offset = (
            vertical_position.find(
                qn(
                    "wp:posOffset"
                )
            )
        )

        self.assertEqual(
            horizontal_offset.text,
            str(
                int(
                    72.0
                    * 12_700
                )
            ),
        )

        self.assertEqual(
            vertical_offset.text,
            str(
                int(
                    144.0
                    * 12_700
                )
            ),
        )

    def test_nonzero_page_origin_is_normalized(
        self,
    ) -> None:
        document = Document()

        image = make_image(
            bbox=Rectangle(
                left=60.0,
                top=120.0,
                right=260.0,
                bottom=220.0,
            ),
        )

        page_bbox = Rectangle(
            left=10.0,
            top=20.0,
            right=610.0,
            bottom=820.0,
        )

        result = (
            EditableFloatingImageRenderer
            .render(
                container=document,
                image=image,
                page_bbox=page_bbox,
            )
        )

        self.assertEqual(
            result.horizontal_offset,
            50.0,
        )

        self.assertEqual(
            result.vertical_offset,
            100.0,
        )

    def test_image_scales_to_page_boundary(
        self,
    ) -> None:
        document = Document()

        image = make_image(
            bbox=Rectangle(
                left=300.0,
                top=100.0,
                right=700.0,
                bottom=300.0,
            ),
        )

        result = (
            EditableFloatingImageRenderer
            .render(
                container=document,
                image=image,
                page_bbox=make_page_bbox(),
            )
        )

        self.assertAlmostEqual(
            result.rendered_width,
            300.0,
            places=2,
        )

        self.assertAlmostEqual(
            result.rendered_height,
            150.0,
            places=2,
        )

    def test_small_image_is_not_upscaled(
        self,
    ) -> None:
        document = Document()

        image = make_image(
            bbox=Rectangle(
                left=10.0,
                top=20.0,
                right=50.0,
                bottom=40.0,
            ),
        )

        result = (
            EditableFloatingImageRenderer
            .render(
                container=document,
                image=image,
                page_bbox=make_page_bbox(),
            )
        )

        self.assertEqual(
            result.rendered_width,
            40.0,
        )

        self.assertEqual(
            result.rendered_height,
            20.0,
        )

    def test_content_image_uses_square_wrap(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableFloatingImageRenderer
            .render(
                container=document,
                image=make_image(
                    role=EditableImageRole.CONTENT,
                ),
                page_bbox=make_page_bbox(),
            )
        )

        self.assertEqual(
            result.wrap_mode,
            EditableFloatingImageWrapMode
            .SQUARE,
        )

        self.assertIn(
            "wp:wrapSquare",
            result.paragraph._p.xml,
        )

    def test_logo_uses_no_wrap(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableFloatingImageRenderer
            .render(
                container=document,
                image=make_image(
                    role=EditableImageRole.LOGO,
                ),
                page_bbox=make_page_bbox(),
            )
        )

        self.assertEqual(
            result.wrap_mode,
            EditableFloatingImageWrapMode.NONE,
        )

        self.assertIn(
            "wp:wrapNone",
            result.paragraph._p.xml,
        )

    def test_explicit_wrap_mode_is_preserved(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableFloatingImageRenderer
            .render(
                container=document,
                image=make_image(
                    role=EditableImageRole.CONTENT,
                ),
                page_bbox=make_page_bbox(),
                wrap_mode=(
                    EditableFloatingImageWrapMode
                    .NONE
                ),
            )
        )

        self.assertEqual(
            result.wrap_mode,
            EditableFloatingImageWrapMode.NONE,
        )

    def test_inline_image_is_rejected(
        self,
    ) -> None:
        document = Document()

        with self.assertRaises(
            ValueError
        ):
            EditableFloatingImageRenderer.render(
                container=document,
                image=make_image(
                    placement=(
                        EditableImagePlacement.INLINE
                    ),
                ),
                page_bbox=make_page_bbox(),
            )

    def test_skipped_image_is_rejected(
        self,
    ) -> None:
        document = Document()

        with self.assertRaises(
            ValueError
        ):
            EditableFloatingImageRenderer.render(
                container=document,
                image=make_image(
                    disposition=(
                        EditableImageDisposition.SKIP
                    ),
                ),
                page_bbox=make_page_bbox(),
            )

    def test_unresolved_payload_is_rejected(
        self,
    ) -> None:
        document = Document()

        with self.assertRaises(
            ValueError
        ):
            EditableFloatingImageRenderer.render(
                container=document,
                image=make_image(
                    payload=None,
                    payload_status=(
                        EditableImagePayloadStatus
                        .FAILED
                    ),
                ),
                page_bbox=make_page_bbox(),
            )

    def test_rotated_image_is_deferred(
        self,
    ) -> None:
        document = Document()

        with self.assertRaises(
            ValueError
        ):
            EditableFloatingImageRenderer.render(
                container=document,
                image=make_image(
                    rotation=90.0,
                ),
                page_bbox=make_page_bbox(),
            )

    def test_invalid_payload_rolls_back_paragraph(
        self,
    ) -> None:
        document = Document()

        paragraph_count_before = len(
            document.paragraphs
        )

        with self.assertRaises(
            ValueError
        ):
            EditableFloatingImageRenderer.render(
                container=document,
                image=make_image(
                    payload=b"invalid-image",
                ),
                page_bbox=make_page_bbox(),
            )

        self.assertEqual(
            len(
                document.paragraphs
            ),
            paragraph_count_before,
        )


if __name__ == "__main__":
    unittest.main()