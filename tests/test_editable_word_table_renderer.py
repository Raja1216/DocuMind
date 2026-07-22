from __future__ import annotations

import unittest

from docx import Document as WordDocument
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
)
from docx.oxml.ns import qn

from src.exporter.editable_word_table_renderer import (
    EditableWordTableRenderer,
)
from src.models.color.rgb_color import (
    RGBColor,
)
from src.models.editable_table import (
    EditableBorderLineStyle,
    EditableCellHorizontalAlignment,
    EditableCellVerticalAlignment,
    EditableTable,
    EditableTableBorder,
    EditableTableCell,
    EditableTableCellBorders,
    EditableTableCellPadding,
    EditableTableCellParagraph,
    EditableTableColumn,
    EditableTableDisposition,
    EditableTableRow,
)
from src.models.geometry.rectangle import (
    Rectangle,
)
from src.models.text_run import (
    TextRun,
)


def make_border(
    *,
    style: EditableBorderLineStyle = (
        EditableBorderLineStyle.SINGLE
    ),
    color: str = "000000",
    width: float = 0.5,
) -> EditableTableBorder:
    return EditableTableBorder(
        style=style,
        color=color,
        width=width,
        confidence=0.95,
    )


def make_borders() -> EditableTableCellBorders:
    return EditableTableCellBorders(
        top=make_border(),
        right=make_border(),
        bottom=make_border(),
        left=make_border(),
    )


def make_table(
    *,
    rows: int = 2,
    columns: int = 2,
) -> EditableTable:
    table = EditableTable(
        page_number=1,
        table_id="table:1:1",
        bbox=Rectangle(
            left=0.0,
            top=0.0,
            right=float(
                columns * 100
            ),
            bottom=float(
                rows * 40
            ),
        ),
        row_count=rows,
        column_count=columns,
        confidence=0.95,
    )

    for row_index in range(
        rows
    ):
        table.add_row(
            EditableTableRow(
                row_index=row_index,
                top=float(
                    row_index * 40
                ),
                bottom=float(
                    (row_index + 1)
                    * 40
                ),
                is_header=(
                    row_index == 0
                ),
                confidence=0.95,
            )
        )

    for column_index in range(
        columns
    ):
        table.add_column(
            EditableTableColumn(
                column_index=column_index,
                left=float(
                    column_index * 100
                ),
                right=float(
                    (column_index + 1)
                    * 100
                ),
                confidence=0.95,
            )
        )

    return table


def add_cell(
    table: EditableTable,
    row: int,
    column: int,
    *,
    text: str = "",
    row_span: int = 1,
    column_span: int = 1,
    content_paragraphs=None,
    horizontal_alignment=(
        EditableCellHorizontalAlignment.LEFT
    ),
    vertical_alignment=(
        EditableCellVerticalAlignment.TOP
    ),
    fill_color: str | None = None,
    borders=None,
    padding=None,
) -> EditableTableCell:
    cell = EditableTableCell(
        row_index=row,
        column_index=column,
        bbox=Rectangle(
            left=float(
                column * 100
            ),
            top=float(
                row * 40
            ),
            right=float(
                (
                    column
                    + column_span
                )
                * 100
            ),
            bottom=float(
                (
                    row
                    + row_span
                )
                * 40
            ),
        ),
        text=text,
        row_span=row_span,
        column_span=column_span,
        content_paragraphs=list(
            content_paragraphs
            or []
        ),
        horizontal_alignment=(
            horizontal_alignment
        ),
        vertical_alignment=(
            vertical_alignment
        ),
        fill_color=fill_color,
        borders=(
            borders
            or make_borders()
        ),
        padding=(
            padding
            or EditableTableCellPadding()
        ),
        confidence=0.95,
    )

    table.add_cell(
        cell
    )

    return cell


def complete_simple_table(
    table: EditableTable,
) -> None:
    for row_index in range(
        table.row_count
    ):
        for column_index in range(
            table.column_count
        ):
            if (
                row_index,
                column_index,
            ) in table.occupied_positions:
                continue

            add_cell(
                table,
                row_index,
                column_index,
                text=(
                    f"{row_index},"
                    f"{column_index}"
                ),
            )


class EditableWordTableRendererTests(
    unittest.TestCase
):

    def test_renders_native_table_with_fixed_layout(
        self,
    ) -> None:
        model = make_table()
        complete_simple_table(
            model
        )

        document = WordDocument()

        word_table = (
            EditableWordTableRenderer.render(
                container=document,
                table=model,
                available_width=300.0,
            )
        )

        self.assertEqual(
            len(
                document.tables
            ),
            1,
        )

        self.assertEqual(
            len(
                word_table.rows
            ),
            2,
        )

        self.assertEqual(
            len(
                word_table.columns
            ),
            2,
        )

        layout = (
            word_table._tbl.tblPr
            .first_child_found_in(
                "w:tblLayout"
            )
        )

        self.assertIsNotNone(
            layout
        )

        self.assertEqual(
            layout.get(
                qn(
                    "w:type"
                )
            ),
            "fixed",
        )

        self.assertFalse(
            word_table.autofit
        )

    def test_column_widths_are_scaled_to_available_width(
        self,
    ) -> None:
        model = make_table(
            rows=1,
            columns=2,
        )

        model.columns[0].right = (
            150.0
        )

        model.columns[1].left = (
            150.0
        )

        model.columns[1].right = (
            300.0
        )

        complete_simple_table(
            model
        )

        document = WordDocument()

        word_table = (
            EditableWordTableRenderer.render(
                container=document,
                table=model,
                available_width=180.0,
            )
        )

        table_width = (
            word_table._tbl.tblPr
            .first_child_found_in(
                "w:tblW"
            )
        )

        self.assertEqual(
            int(
                table_width.get(
                    qn(
                        "w:w"
                    )
                )
            ),
            3600,
        )

    def test_formatted_runs_are_preserved(
        self,
    ) -> None:
        model = make_table(
            rows=1,
            columns=1,
        )

        paragraph = (
            EditableTableCellParagraph(
                text="Bold normal",
                runs=[
                    TextRun(
                        text="Bold",
                        font_name=(
                            "Helvetica-Bold"
                        ),
                        font_size=11.2,
                        color=RGBColor(
                            255,
                            0,
                            0,
                        ),
                        bold=True,
                    ),
                    TextRun(
                        text=" normal",
                        font_name="Helvetica",
                        font_size=10.0,
                        color=RGBColor(
                            0,
                            0,
                            0,
                        ),
                    ),
                ],
                confidence=0.95,
            )
        )

        add_cell(
            model,
            0,
            0,
            content_paragraphs=[
                paragraph
            ],
        )

        document = WordDocument()

        word_table = (
            EditableWordTableRenderer.render(
                container=document,
                table=model,
            )
        )

        word_paragraph = (
            word_table.cell(
                0,
                0,
            ).paragraphs[0]
        )

        self.assertEqual(
            len(
                word_paragraph.runs
            ),
            2,
        )

        self.assertTrue(
            word_paragraph
            .runs[0]
            .bold
        )

        self.assertEqual(
            word_paragraph
            .runs[0]
            .font.name,
            "Arial",
        )

        self.assertEqual(
            str(
                word_paragraph
                .runs[0]
                .font.color.rgb
            ),
            "FF0000",
        )

    def test_cell_alignment_is_applied(
        self,
    ) -> None:
        model = make_table(
            rows=1,
            columns=1,
        )

        add_cell(
            model,
            0,
            0,
            text="Centered",
            horizontal_alignment=(
                EditableCellHorizontalAlignment
                .CENTER
            ),
            vertical_alignment=(
                EditableCellVerticalAlignment
                .BOTTOM
            ),
        )

        document = WordDocument()

        word_table = (
            EditableWordTableRenderer.render(
                container=document,
                table=model,
            )
        )

        word_cell = (
            word_table.cell(
                0,
                0,
            )
        )

        self.assertEqual(
            word_cell.vertical_alignment,
            WD_CELL_VERTICAL_ALIGNMENT
            .BOTTOM,
        )

        self.assertEqual(
            word_cell.paragraphs[0]
            .alignment,
            WD_ALIGN_PARAGRAPH
            .CENTER,
        )

    def test_fill_borders_and_padding_are_written(
        self,
    ) -> None:
        model = make_table(
            rows=1,
            columns=1,
        )

        borders = (
            EditableTableCellBorders(
                top=make_border(
                    style=(
                        EditableBorderLineStyle
                        .DOUBLE
                    ),
                    color="123456",
                    width=1.0,
                ),
                right=make_border(
                    style=(
                        EditableBorderLineStyle
                        .DASHED
                    ),
                ),
                bottom=make_border(
                    style=(
                        EditableBorderLineStyle
                        .DOTTED
                    ),
                ),
                left=make_border(
                    style=(
                        EditableBorderLineStyle
                        .NONE
                    ),
                    width=0.0,
                ),
            )
        )

        add_cell(
            model,
            0,
            0,
            text="Styled",
            fill_color="ABCDEF",
            borders=borders,
            padding=(
                EditableTableCellPadding(
                    top=1.0,
                    right=2.0,
                    bottom=3.0,
                    left=4.0,
                )
            ),
        )

        document = WordDocument()

        word_table = (
            EditableWordTableRenderer.render(
                container=document,
                table=model,
            )
        )

        cell_properties = (
            word_table.cell(
                0,
                0,
            )._tc.tcPr
        )

        shading = (
            cell_properties
            .first_child_found_in(
                "w:shd"
            )
        )

        self.assertEqual(
            shading.get(
                qn(
                    "w:fill"
                )
            ),
            "ABCDEF",
        )

        margins = (
            cell_properties
            .first_child_found_in(
                "w:tcMar"
            )
        )

        self.assertEqual(
            margins.find(
                qn(
                    "w:left"
                )
            ).get(
                qn(
                    "w:w"
                )
            ),
            "80",
        )

        border_xml = (
            cell_properties
            .first_child_found_in(
                "w:tcBorders"
            )
        )

        self.assertEqual(
            border_xml.find(
                qn(
                    "w:top"
                )
            ).get(
                qn(
                    "w:val"
                )
            ),
            "double",
        )

        self.assertEqual(
            border_xml.find(
                qn(
                    "w:left"
                )
            ).get(
                qn(
                    "w:val"
                )
            ),
            "nil",
        )

    def test_rectangular_merged_cell_is_rendered(
        self,
    ) -> None:
        model = make_table()

        add_cell(
            model,
            0,
            0,
            text="Merged",
            row_span=2,
            column_span=2,
        )

        document = WordDocument()

        word_table = (
            EditableWordTableRenderer.render(
                container=document,
                table=model,
            )
        )

        top_left = word_table.cell(
            0,
            0,
        )

        bottom_right = (
            word_table.cell(
                1,
                1,
            )
        )

        self.assertIs(
            top_left._tc,
            bottom_right._tc,
        )

        self.assertEqual(
            top_left.text,
            "Merged",
        )

    def test_header_row_and_minimum_height_are_written(
        self,
    ) -> None:
        model = make_table(
            rows=1,
            columns=1,
        )

        add_cell(
            model,
            0,
            0,
            text="Header",
        )

        document = WordDocument()

        word_table = (
            EditableWordTableRenderer.render(
                container=document,
                table=model,
            )
        )

        row_properties = (
            word_table.rows[0]
            ._tr
            .get_or_add_trPr()
        )

        self.assertIsNotNone(
            row_properties.find(
                qn(
                    "w:tblHeader"
                )
            )
        )

        self.assertIsNotNone(
            row_properties.find(
                qn(
                    "w:cantSplit"
                )
            )
        )

        self.assertIsNotNone(
            word_table.rows[0]
            .height
        )

    def test_plain_text_fallback_is_rendered(
        self,
    ) -> None:
        model = make_table(
            rows=1,
            columns=1,
        )

        add_cell(
            model,
            0,
            0,
            text="Fallback text",
        )

        document = WordDocument()

        word_table = (
            EditableWordTableRenderer.render(
                container=document,
                table=model,
            )
        )

        self.assertEqual(
            word_table.cell(
                0,
                0,
            ).text,
            "Fallback text",
        )

    def test_visual_fallback_table_is_rejected(
        self,
    ) -> None:
        model = make_table(
            rows=1,
            columns=1,
        )

        add_cell(
            model,
            0,
            0,
            text="Not editable",
        )

        model.disposition = (
            EditableTableDisposition
            .VISUAL_FALLBACK
        )

        document = WordDocument()

        with self.assertRaises(
            ValueError
        ):
            EditableWordTableRenderer.render(
                container=document,
                table=model,
            )


if __name__ == "__main__":
    unittest.main()
