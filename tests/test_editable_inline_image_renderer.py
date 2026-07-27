from __future__ import annotations

import base64
import unittest

from docx import (
    Document,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
)

from src.exporter.editable_inline_image_renderer import (
    EditableInlineImageRenderer,
)
from src.models.editable_image import (
    EditableImage,
    EditableImageDisposition,
    EditableImageExtractionMode,
    EditableImageHorizontalAlignment,
    EditableImagePayloadStatus,
    EditableImagePlacement,
)
from src.models.geometry.rectangle import (
    Rectangle,
)


PNG_BYTES = base64.b64decode(
    (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAAB"
        "CAQAAAC1HAwCAAAAC0lEQVR42mP8"
        "/x8AAusB9Y9Z5ZkAAAAASUVORK5CYII="
    )
)


def make_image(
    **overrides,
) -> EditableImage:
    values = {
        "page_number": 1,
        "image_id": "image:1:1",
        "bbox": Rectangle(
            left=0.0,
            top=0.0,
            right=200.0,
            bottom=100.0,
        ),
        "extraction_mode": (
            EditableImageExtractionMode
            .DIRECT_BYTES
        ),
        "disposition": (
            EditableImageDisposition.NATIVE
        ),
        "placement": (
            EditableImagePlacement.INLINE
        ),
        "horizontal_alignment": (
            EditableImageHorizontalAlignment
            .LEFT
        ),
        "payload": PNG_BYTES,
        "payload_status": (
            EditableImagePayloadStatus.READY
        ),
        "payload_mime_type": "image/png",
        "extension": "png",
        "pixel_width": 1,
        "pixel_height": 1,
        "confidence": 0.95,
        "payload_confidence": 0.95,
    }

    values.update(
        overrides
    )

    return EditableImage(
        **values
    )


class EditableInlineImageRendererTests(
    unittest.TestCase
):

    def test_native_inline_image_is_rendered(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableInlineImageRenderer
            .render(
                container=document,
                image=make_image(),
                available_width=400.0,
            )
        )

        self.assertEqual(
            len(
                document.inline_shapes
            ),
            1,
        )

        self.assertEqual(
            result.rendered_width,
            200.0,
        )

        self.assertEqual(
            result.rendered_height,
            100.0,
        )

    def test_region_rendered_payload_is_allowed(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableInlineImageRenderer
            .render(
                container=document,
                image=make_image(
                    extraction_mode=(
                        EditableImageExtractionMode
                        .PAGE_REGION
                    ),
                    disposition=(
                        EditableImageDisposition
                        .REGION_FALLBACK
                    ),
                    payload_status=(
                        EditableImagePayloadStatus
                        .REGION_RENDERED
                    ),
                ),
                available_width=400.0,
            )
        )

        self.assertIsNotNone(
            result.inline_shape
        )

    def test_image_scales_down_to_available_width(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableInlineImageRenderer
            .render(
                container=document,
                image=make_image(),
                available_width=100.0,
            )
        )

        self.assertAlmostEqual(
            result.rendered_width,
            100.0,
            places=2,
        )

        self.assertAlmostEqual(
            result.rendered_height,
            50.0,
            places=2,
        )

    def test_image_scales_down_to_available_height(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableInlineImageRenderer
            .render(
                container=document,
                image=make_image(),
                available_width=400.0,
                available_height=25.0,
            )
        )

        self.assertAlmostEqual(
            result.rendered_width,
            50.0,
            places=2,
        )

        self.assertAlmostEqual(
            result.rendered_height,
            25.0,
            places=2,
        )

    def test_small_image_is_not_upscaled(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableInlineImageRenderer
            .render(
                container=document,
                image=make_image(
                    bbox=Rectangle(
                        left=0.0,
                        top=0.0,
                        right=40.0,
                        bottom=20.0,
                    ),
                ),
                available_width=400.0,
            )
        )

        self.assertEqual(
            result.rendered_width,
            40.0,
        )

        self.assertEqual(
            result.rendered_height,
            20.0,
        )

    def test_center_alignment_is_applied(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableInlineImageRenderer
            .render(
                container=document,
                image=make_image(
                    horizontal_alignment=(
                        EditableImageHorizontalAlignment
                        .CENTER
                    ),
                ),
                available_width=400.0,
            )
        )

        self.assertEqual(
            result.paragraph.alignment,
            WD_ALIGN_PARAGRAPH.CENTER,
        )

    def test_right_alignment_is_applied(
        self,
    ) -> None:
        document = Document()

        result = (
            EditableInlineImageRenderer
            .render(
                container=document,
                image=make_image(
                    horizontal_alignment=(
                        EditableImageHorizontalAlignment
                        .RIGHT
                    ),
                ),
                available_width=400.0,
            )
        )

        self.assertEqual(
            result.paragraph.alignment,
            WD_ALIGN_PARAGRAPH.RIGHT,
        )

    def test_non_inline_image_is_rejected(
        self,
    ) -> None:
        document = Document()

        with self.assertRaises(
            ValueError
        ):
            EditableInlineImageRenderer.render(
                container=document,
                image=make_image(
                    placement=(
                        EditableImagePlacement
                        .FLOATING
                    ),
                ),
                available_width=400.0,
            )

    def test_skipped_image_is_rejected(
        self,
    ) -> None:
        document = Document()

        with self.assertRaises(
            ValueError
        ):
            EditableInlineImageRenderer.render(
                container=document,
                image=make_image(
                    disposition=(
                        EditableImageDisposition
                        .SKIP
                    ),
                ),
                available_width=400.0,
            )

    def test_unresolved_payload_is_rejected(
        self,
    ) -> None:
        document = Document()

        with self.assertRaises(
            ValueError
        ):
            EditableInlineImageRenderer.render(
                container=document,
                image=make_image(
                    payload_status=(
                        EditableImagePayloadStatus
                        .UNRESOLVED
                    ),
                ),
                available_width=400.0,
            )

    def test_rotated_image_is_rejected(
        self,
    ) -> None:
        document = Document()

        with self.assertRaises(
            ValueError
        ):
            EditableInlineImageRenderer.render(
                container=document,
                image=make_image(
                    rotation=90.0,
                ),
                available_width=400.0,
            )

    def test_invalid_payload_rolls_back_paragraph(
        self,
    ) -> None:
        document = Document()

        paragraph_count_before = len(
            document.paragraphs
        )

        with self.assertRaises(
            ValueError
        ):
            EditableInlineImageRenderer.render(
                container=document,
                image=make_image(
                    payload=b"invalid-image",
                ),
                available_width=400.0,
            )

        self.assertEqual(
            len(
                document.paragraphs
            ),
            paragraph_count_before,
        )

        self.assertEqual(
            len(
                document.inline_shapes
            ),
            0,
        )

    def test_table_cell_can_be_used_as_container(
        self,
    ) -> None:
        document = Document()

        table = document.add_table(
            rows=1,
            cols=1,
        )

        word_cell = table.cell(
            0,
            0,
        )

        result = (
            EditableInlineImageRenderer
            .render(
                container=word_cell,
                image=make_image(),
                available_width=100.0,
            )
        )

        self.assertIsNotNone(
            result.inline_shape
        )

        self.assertGreaterEqual(
            len(
                word_cell.paragraphs
            ),
            2,
        )


if __name__ == "__main__":
    unittest.main()