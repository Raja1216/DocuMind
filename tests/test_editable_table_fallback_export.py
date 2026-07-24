from __future__ import annotations

import unittest

from types import SimpleNamespace
from unittest.mock import Mock, patch

import fitz

from docx import Document as WordDocument
from docx.oxml.ns import qn

from src.exporter.docx_exporter import (
    DocxExporter,
)
from src.exporter.editable_page_render_resolver import (
    EditablePageRenderPlan,
    EditableRenderAction,
    EditableRenderInstruction,
)
from src.exporter.editable_table_export_coordinator import (
    EditableTableExportCoordinator,
)
from src.exporter.editable_table_region_fallback_renderer import (
    EditableTableRegionFallbackRenderer,
)
from src.exporter.editable_word_table_renderer import (
    EditableWordTableRenderer,
)
from src.models.editable_table import (
    EditableTable,
    EditableTableCell,
    EditableTableColumn,
    EditableTableDisposition,
    EditableTableRow,
)
from src.models.editable_table_export import (
    EditableTableExportMode,
)
from src.models.editable_table_validation import (
    EditableTableRenderDecision,
    EditableTableValidationReport,
)
from src.models.geometry.rectangle import (
    Rectangle,
)


def make_pdf_page():
    document = fitz.open()

    page = document.new_page(
        width=400.0,
        height=300.0,
    )

    page.draw_rect(
        fitz.Rect(
            50.0,
            60.0,
            350.0,
            160.0,
        ),
        color=(0.0, 0.0, 0.0),
        width=1.0,
    )

    page.draw_line(
        fitz.Point(200.0, 60.0),
        fitz.Point(200.0, 160.0),
        color=(0.0, 0.0, 0.0),
        width=1.0,
    )

    page.draw_line(
        fitz.Point(50.0, 110.0),
        fitz.Point(350.0, 110.0),
        color=(0.0, 0.0, 0.0),
        width=1.0,
    )

    page.insert_text(
        fitz.Point(60.0, 90.0),
        "Header A",
    )

    page.insert_text(
        fitz.Point(210.0, 90.0),
        "Header B",
    )

    page.insert_text(
        fitz.Point(60.0, 140.0),
        "Value A",
    )

    page.insert_text(
        fitz.Point(210.0, 140.0),
        "Value B",
    )

    return document, page


def make_table(
    *,
    disposition: EditableTableDisposition = (
        EditableTableDisposition.EDITABLE
    ),
    table_id: str = "table:1:1",
) -> EditableTable:
    table = EditableTable(
        page_number=1,
        table_id=table_id,
        bbox=Rectangle(
            left=50.0,
            top=60.0,
            right=350.0,
            bottom=160.0,
        ),
        row_count=2,
        column_count=2,
        disposition=disposition,
        confidence=0.95,
    )

    for row_index in range(2):
        table.add_row(
            EditableTableRow(
                row_index=row_index,
                top=(
                    60.0
                    + row_index * 50.0
                ),
                bottom=(
                    60.0
                    + (row_index + 1)
                    * 50.0
                ),
                is_header=(
                    row_index == 0
                ),
                confidence=0.95,
            )
        )

    for column_index in range(2):
        table.add_column(
            EditableTableColumn(
                column_index=column_index,
                left=(
                    50.0
                    + column_index
                    * 150.0
                ),
                right=(
                    50.0
                    + (column_index + 1)
                    * 150.0
                ),
                confidence=0.95,
            )
        )

    values = [
        [
            "Header A",
            "Header B",
        ],
        [
            "Value A",
            "Value B",
        ],
    ]

    for row_index in range(2):
        for column_index in range(2):
            table.add_cell(
                EditableTableCell(
                    row_index=row_index,
                    column_index=column_index,
                    bbox=Rectangle(
                        left=(
                            50.0
                            + column_index
                            * 150.0
                        ),
                        top=(
                            60.0
                            + row_index
                            * 50.0
                        ),
                        right=(
                            50.0
                            + (column_index + 1)
                            * 150.0
                        ),
                        bottom=(
                            60.0
                            + (row_index + 1)
                            * 50.0
                        ),
                    ),
                    text=values[
                        row_index
                    ][
                        column_index
                    ],
                    confidence=0.95,
                )
            )

    return table


def make_page(
    source_page,
):
    return SimpleNamespace(
        number=1,
        source_pdf_page=source_page,
        profile=None,
        editable_table_export_results={},
        editable_table_validation_reports={},
    )


class EditableTableFallbackExportTests(
    unittest.TestCase
):

    def test_region_fallback_renders_inline_image(
        self,
    ) -> None:
        pdf_document, source_page = (
            make_pdf_page()
        )

        self.addCleanup(
            pdf_document.close
        )

        word_document = WordDocument()

        (
            _,
            crop,
            rendered_width,
            rendered_height,
        ) = (
            EditableTableRegionFallbackRenderer
            .render(
                container=word_document,
                page=make_page(
                    source_page
                ),
                table=make_table(),
                available_width=250.0,
            )
        )

        self.assertEqual(
            len(
                word_document.inline_shapes
            ),
            1,
        )

        self.assertLessEqual(
            rendered_width,
            250.0,
        )

        self.assertGreater(
            rendered_height,
            0.0,
        )

        self.assertLess(
            crop.left,
            50.0,
        )

        self.assertGreater(
            crop.right,
            350.0,
        )

    def test_crop_is_clipped_to_source_page(
        self,
    ) -> None:
        pdf_document, source_page = (
            make_pdf_page()
        )

        self.addCleanup(
            pdf_document.close
        )

        word_document = WordDocument()

        table = make_table()

        table.bbox = Rectangle(
            left=-20.0,
            top=-10.0,
            right=420.0,
            bottom=310.0,
        )

        _, crop, _, _ = (
            EditableTableRegionFallbackRenderer
            .render(
                container=word_document,
                page=make_page(
                    source_page
                ),
                table=table,
                available_width=300.0,
            )
        )

        self.assertEqual(
            crop.left,
            0.0,
        )

        self.assertEqual(
            crop.top,
            0.0,
        )

        self.assertEqual(
            crop.right,
            400.0,
        )

        self.assertEqual(
            crop.bottom,
            300.0,
        )

    def test_validation_selected_fallback_is_recorded(
        self,
    ) -> None:
        pdf_document, source_page = (
            make_pdf_page()
        )

        self.addCleanup(
            pdf_document.close
        )

        word_document = WordDocument()
        page = make_page(
            source_page
        )
        table = make_table(
            disposition=(
                EditableTableDisposition
                .VISUAL_FALLBACK
            )
        )

        report = EditableTableValidationReport(
            table_id=table.table_id,
            page_number=1,
            decision=(
                EditableTableRenderDecision
                .VISUAL_FALLBACK
            ),
            native_confidence=0.40,
        )

        result = (
            EditableTableExportCoordinator
            .render(
                container=word_document,
                page=page,
                table=table,
                available_width=300.0,
                prefer_native=False,
                validation_report=report,
            )
        )

        self.assertTrue(
            result.success
        )

        self.assertFalse(
            result.native_attempted
        )

        self.assertTrue(
            result.fallback_attempted
        )

        self.assertEqual(
            result.final_mode,
            EditableTableExportMode
            .VISUAL_FALLBACK,
        )

        self.assertIs(
            page.editable_table_export_results[
                table.table_id
            ],
            result,
        )

    def test_native_failure_rolls_back_partial_table(
        self,
    ) -> None:
        pdf_document, source_page = (
            make_pdf_page()
        )

        self.addCleanup(
            pdf_document.close
        )

        word_document = WordDocument()
        page = make_page(
            source_page
        )
        table = make_table()

        report = EditableTableValidationReport(
            table_id=table.table_id,
            page_number=1,
            decision=(
                EditableTableRenderDecision
                .NATIVE_SAFE
            ),
            native_confidence=0.95,
        )

        def insert_then_fail(
            *,
            container,
            **_,
        ):
            container.add_table(
                rows=1,
                cols=1,
            )

            raise RuntimeError(
                "simulated native failure"
            )

        with patch.object(
            EditableWordTableRenderer,
            "render",
            side_effect=insert_then_fail,
        ):
            result = (
                EditableTableExportCoordinator
                .render(
                    container=word_document,
                    page=page,
                    table=table,
                    available_width=300.0,
                    prefer_native=True,
                    validation_report=report,
                )
            )

        self.assertTrue(
            result.success
        )

        self.assertTrue(
            result.native_attempted
        )

        self.assertTrue(
            result.fallback_attempted
        )

        self.assertEqual(
            len(
                word_document.tables
            ),
            0,
        )

        self.assertEqual(
            len(
                word_document.inline_shapes
            ),
            1,
        )

        self.assertEqual(
            report.decision,
            EditableTableRenderDecision
            .VISUAL_FALLBACK,
        )

        self.assertIn(
            "NATIVE_RENDER_EXCEPTION",
            {
                issue.code
                for issue in report.issues
            },
        )

    def test_native_success_does_not_render_fallback(
        self,
    ) -> None:
        word_document = WordDocument()
        page = make_page(
            None
        )
        table = make_table()

        report = EditableTableValidationReport(
            table_id=table.table_id,
            page_number=1,
            decision=(
                EditableTableRenderDecision
                .NATIVE_SAFE
            ),
            native_confidence=0.95,
        )

        result = (
            EditableTableExportCoordinator
            .render(
                container=word_document,
                page=page,
                table=table,
                available_width=300.0,
                prefer_native=True,
                validation_report=report,
            )
        )

        self.assertTrue(
            result.success
        )

        self.assertEqual(
            result.final_mode,
            EditableTableExportMode.NATIVE,
        )

        self.assertFalse(
            result.fallback_attempted
        )

        self.assertEqual(
            len(
                word_document.tables
            ),
            1,
        )

        self.assertEqual(
            len(
                word_document.inline_shapes
            ),
            0,
        )

    def test_missing_source_page_does_not_stop_export(
        self,
    ) -> None:
        word_document = WordDocument()
        page = make_page(
            None
        )
        table = make_table(
            disposition=(
                EditableTableDisposition
                .VISUAL_FALLBACK
            )
        )

        report = EditableTableValidationReport(
            table_id=table.table_id,
            page_number=1,
            decision=(
                EditableTableRenderDecision
                .VISUAL_FALLBACK
            ),
            native_confidence=0.30,
        )

        result = (
            EditableTableExportCoordinator
            .render(
                container=word_document,
                page=page,
                table=table,
                available_width=300.0,
                prefer_native=False,
                validation_report=report,
            )
        )

        self.assertFalse(
            result.success
        )

        self.assertEqual(
            result.final_mode,
            EditableTableExportMode.FAILED,
        )

        self.assertEqual(
            len(
                word_document.inline_shapes
            ),
            0,
        )

        self.assertIn(
            "TABLE_FALLBACK_RENDER_EXCEPTION",
            {
                issue.code
                for issue in report.issues
            },
        )

    def test_failed_table_does_not_block_later_native_table(
        self,
    ) -> None:
        word_document = WordDocument()
        page = make_page(
            None
        )

        failed_table = make_table(
            disposition=(
                EditableTableDisposition
                .VISUAL_FALLBACK
            ),
            table_id="table:1:failed",
        )

        native_table = make_table(
            table_id="table:1:native",
        )

        failed_result = (
            EditableTableExportCoordinator
            .render(
                container=word_document,
                page=page,
                table=failed_table,
                available_width=300.0,
                prefer_native=False,
            )
        )

        native_result = (
            EditableTableExportCoordinator
            .render(
                container=word_document,
                page=page,
                table=native_table,
                available_width=300.0,
                prefer_native=True,
            )
        )

        self.assertFalse(
            failed_result.success
        )

        self.assertTrue(
            native_result.success
        )

        self.assertEqual(
            len(
                word_document.tables
            ),
            1,
        )

    def test_docx_exporter_preserves_fallback_flow_order(
        self,
    ) -> None:
        pdf_document, source_page = (
            make_pdf_page()
        )

        self.addCleanup(
            pdf_document.close
        )

        word_document = WordDocument()
        page = make_page(
            source_page
        )

        first = SimpleNamespace(
            text="Before fallback",
            list_type=None,
            left=50.0,
            right=350.0,
        )

        second = SimpleNamespace(
            text="After fallback",
            list_type=None,
            left=50.0,
            right=350.0,
        )

        table = make_table(
            disposition=(
                EditableTableDisposition
                .VISUAL_FALLBACK
            )
        )

        page.editable_table_validation_reports = {
            table.table_id: (
                EditableTableValidationReport(
                    table_id=table.table_id,
                    page_number=1,
                    decision=(
                        EditableTableRenderDecision
                        .VISUAL_FALLBACK
                    ),
                    native_confidence=0.40,
                )
            )
        }

        page_plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .RENDER_PARAGRAPH
                    ),
                    source=first,
                    layout_item=(
                        SimpleNamespace()
                    ),
                ),
                EditableRenderInstruction(
                    order=2,
                    action=(
                        EditableRenderAction
                        .RENDER_TABLE_FALLBACK
                    ),
                    source=table,
                ),
                EditableRenderInstruction(
                    order=3,
                    action=(
                        EditableRenderAction
                        .RENDER_PARAGRAPH
                    ),
                    source=second,
                    layout_item=(
                        SimpleNamespace()
                    ),
                ),
            ],
        )

        def render_runs(
            *,
            word_paragraph,
            paragraph,
            **_,
        ):
            word_paragraph.add_run(
                paragraph.text
            )

        with (
            patch.object(
                DocxExporter,
                "_build_editable_render_plan",
                return_value=page_plan,
            ),
            patch.object(
                DocxExporter,
                "_region_is_heading",
                return_value=False,
            ),
            patch.object(
                DocxExporter,
                "_apply_region_layout",
            ),
            patch.object(
                DocxExporter,
                "_render_paragraph_runs",
                side_effect=render_runs,
            ),
            patch.object(
                DocxExporter,
                "_normalize_alignment_indentation",
            ),
            patch(
                (
                    "src.exporter."
                    "docx_exporter."
                    "EditableLayoutResolver."
                    "apply_alignment"
                )
            ),
        ):
            DocxExporter._render_page(
                word_document=word_document,
                page=page,
                numbering_manager=Mock(),
                list_sequence_resolver=Mock(),
            )

        body_paragraphs = [
            element
            for element in (
                word_document
                ._element
                .body
            )
            if element.tag
            == qn(
                "w:p"
            )
        ]

        self.assertEqual(
            len(body_paragraphs),
            3,
        )

        self.assertEqual(
            word_document.paragraphs[0].text,
            "Before fallback",
        )

        self.assertEqual(
            word_document.paragraphs[2].text,
            "After fallback",
        )

        drawing_count = len(
            body_paragraphs[1].xpath(
                ".//w:drawing"
            )
        )

        self.assertEqual(
            drawing_count,
            1,
        )

    def test_export_results_replace_stale_page_results(
        self,
    ) -> None:
        word_document = WordDocument()
        page = make_page(
            None
        )
        page.editable_table_export_results = {
            "stale": object()
        }

        page_plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[],
        )

        with patch.object(
            DocxExporter,
            "_build_editable_render_plan",
            return_value=page_plan,
        ):
            DocxExporter._render_page(
                word_document=word_document,
                page=page,
                numbering_manager=Mock(),
                list_sequence_resolver=Mock(),
            )

        self.assertEqual(
            page.editable_table_export_results,
            {},
        )


if __name__ == "__main__":
    unittest.main()
