from __future__ import annotations

import unittest

from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

from src.exporter.docx_exporter import (
    DocxExporter,
)
from src.exporter.editable_page_render_resolver import (
    EditablePageRenderPlan,
    EditableRenderAction,
    EditableRenderInstruction,
)
from src.models.editable_image import (
    EditableImage,
    EditableImageDisposition,
    EditableImageExtractionMode,
    EditableImageHorizontalAlignment,
    EditableImagePayloadStatus,
    EditableImagePlacement,
)
from src.models.geometry.rectangle import (
    Rectangle,
)
from src.models.page_render_plan import (
    PageRenderItem,
    RenderDisposition,
    RenderItemKind,
    RenderItemRole,
    RenderPlacement,
)


def create_test_png() -> bytes:
    try:
        import pymupdf

    except ImportError:
        import fitz as pymupdf

    pixmap = pymupdf.Pixmap(
        pymupdf.csRGB,
        pymupdf.IRect(
            0,
            0,
            4,
            2,
        ),
        False,
    )

    pixmap.clear_with(
        255
    )

    return pixmap.tobytes(
        "png"
    )


PNG_BYTES = create_test_png()


def make_image(
    *,
    width: float = 200.0,
    height: float = 100.0,
    payload: bytes = PNG_BYTES,
    status: EditableImagePayloadStatus = (
        EditableImagePayloadStatus.READY
    ),
    disposition: EditableImageDisposition = (
        EditableImageDisposition.NATIVE
    ),
    placement: EditableImagePlacement = (
        EditableImagePlacement.INLINE
    ),
) -> EditableImage:
    return EditableImage(
        page_number=1,
        image_id="image:1:1",
        bbox=Rectangle(
            left=0.0,
            top=100.0,
            right=width,
            bottom=100.0 + height,
        ),
        extraction_mode=(
            EditableImageExtractionMode
            .DIRECT_BYTES
        ),
        disposition=disposition,
        placement=placement,
        horizontal_alignment=(
            EditableImageHorizontalAlignment
            .CENTER
        ),
        payload=payload,
        payload_status=status,
        payload_mime_type="image/png",
        extension="png",
        pixel_width=4,
        pixel_height=2,
        confidence=0.95,
        payload_confidence=0.95,
    )


def make_image_render_item(
    image,
    *,
    order: int = 1,
) -> PageRenderItem:
    return PageRenderItem(
        order=order,
        page_number=1,
        item_id=(
            f"image:{order}"
        ),
        kind=RenderItemKind.IMAGE,
        placement=RenderPlacement.FLOW,
        disposition=(
            RenderDisposition.VISUAL
        ),
        role=RenderItemRole.BODY,
        bbox=image.bbox,
        source=image.source_image,
        source_index=order - 1,
        confidence=0.95,
    )


def make_page():
    return SimpleNamespace(
        number=1,
        profile=None,
        editable_table_export_results={},
        editable_table_validation_reports={},
    )


def make_paragraph_region(
    *,
    text: str,
    top: float,
    bottom: float,
):
    return SimpleNamespace(
        text=text,
        left=50.0,
        right=500.0,
        top=top,
        bottom=bottom,
        list_type=None,
    )


def make_paragraph_instruction(
    *,
    order: int,
    text: str,
    top: float,
):
    region = make_paragraph_region(
        text=text,
        top=top,
        bottom=top + 20.0,
    )

    render_item = PageRenderItem(
        order=order,
        page_number=1,
        item_id=(
            f"paragraph:{order}"
        ),
        kind=RenderItemKind.PARAGRAPH,
        placement=RenderPlacement.FLOW,
        disposition=(
            RenderDisposition.EDITABLE
        ),
        role=RenderItemRole.BODY,
        bbox=Rectangle(
            left=region.left,
            top=region.top,
            right=region.right,
            bottom=region.bottom,
        ),
        source=region,
        source_index=order - 1,
        confidence=0.95,
    )

    layout_item = SimpleNamespace(
        paragraph=region,
        apply_alignment=False,
    )

    return EditableRenderInstruction(
        order=order,
        action=(
            EditableRenderAction
            .RENDER_PARAGRAPH
        ),
        source=region,
        render_item=render_item,
        layout_item=layout_item,
    )


class DocxInlineImageIntegrationTests(
    unittest.TestCase
):

    def render_plan(
        self,
        plan,
    ):
        document = Document()

        page = make_page()

        def render_text(
            *,
            word_paragraph,
            paragraph,
            **_,
        ):
            word_paragraph.add_run(
                paragraph.text
            )

        with (
            patch.object(
                DocxExporter,
                "_build_editable_render_plan",
                return_value=plan,
            ),
            patch.object(
                DocxExporter,
                "_region_is_heading",
                return_value=False,
            ),
            patch.object(
                DocxExporter,
                "_apply_region_layout",
            ),
            patch.object(
                DocxExporter,
                "_render_paragraph_runs",
                side_effect=render_text,
            ),
            patch.object(
                DocxExporter,
                "_normalize_alignment_indentation",
            ),
            patch(
                (
                    "src.exporter.docx_exporter."
                    "EditableLayoutResolver."
                    "apply_alignment"
                )
            ),
        ):
            DocxExporter._render_page(
                word_document=document,
                page=page,
                numbering_manager=(
                    SimpleNamespace()
                ),
                list_sequence_resolver=(
                    SimpleNamespace()
                ),
                validation_report=None,
            )

        return document

    def test_image_only_page_renders_inline_shape(
        self,
    ) -> None:
        image = make_image()

        plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .RENDER_INLINE_IMAGE
                    ),
                    source=image,
                    render_item=(
                        make_image_render_item(
                            image
                        )
                    ),
                )
            ],
        )

        document = self.render_plan(
            plan
        )

        self.assertEqual(
            len(
                document.inline_shapes
            ),
            1,
        )

    def test_paragraph_image_paragraph_order(
        self,
    ) -> None:
        image = make_image()

        plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                make_paragraph_instruction(
                    order=1,
                    text="Before",
                    top=50.0,
                ),
                EditableRenderInstruction(
                    order=2,
                    action=(
                        EditableRenderAction
                        .RENDER_INLINE_IMAGE
                    ),
                    source=image,
                    render_item=(
                        make_image_render_item(
                            image,
                            order=2,
                        )
                    ),
                ),
                make_paragraph_instruction(
                    order=3,
                    text="After",
                    top=250.0,
                ),
            ],
        )

        document = self.render_plan(
            plan
        )

        self.assertEqual(
            [
                paragraph.text
                for paragraph
                in document.paragraphs
            ],
            [
                "Before",
                "",
                "After",
            ],
        )

        self.assertIn(
            "w:drawing",
            document.paragraphs[1]._p.xml,
        )

    def test_wide_image_scales_to_section_width(
        self,
    ) -> None:
        image = make_image(
            width=1000.0,
            height=500.0,
        )

        plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .RENDER_INLINE_IMAGE
                    ),
                    source=image,
                    render_item=(
                        make_image_render_item(
                            image
                        )
                    ),
                )
            ],
        )

        document = self.render_plan(
            plan
        )

        section = document.sections[-1]

        available_width = (
            section.page_width.pt
            - section.left_margin.pt
            - section.right_margin.pt
        )

        rendered_width = (
            document.inline_shapes[0]
            .width
            .pt
        )

        self.assertLessEqual(
            rendered_width,
            available_width + 0.5,
        )

    def test_region_rendered_payload_is_rendered(
        self,
    ) -> None:
        image = make_image(
            status=(
                EditableImagePayloadStatus
                .REGION_RENDERED
            ),
            disposition=(
                EditableImageDisposition
                .REGION_FALLBACK
            ),
        )

        plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .RENDER_INLINE_IMAGE
                    ),
                    source=image,
                    render_item=(
                        make_image_render_item(
                            image
                        )
                    ),
                )
            ],
        )

        document = self.render_plan(
            plan
        )

        self.assertEqual(
            len(
                document.inline_shapes
            ),
            1,
        )

    def test_deferred_floating_image_is_not_rendered(
        self,
    ) -> None:
        image = make_image(
            placement=(
                EditableImagePlacement.FLOATING
            ),
        )

        plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .DEFER_IMAGE
                    ),
                    source=image,
                )
            ],
        )

        document = self.render_plan(
            plan
        )

        self.assertEqual(
            len(
                document.inline_shapes
            ),
            0,
        )

    def test_deferred_background_image_is_not_rendered(
        self,
    ) -> None:
        image = make_image(
            placement=(
                EditableImagePlacement
                .BACKGROUND
            ),
        )

        plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .DEFER_IMAGE
                    ),
                    source=image,
                )
            ],
        )

        document = self.render_plan(
            plan
        )

        self.assertEqual(
            len(
                document.inline_shapes
            ),
            0,
        )

    def test_invalid_image_does_not_block_later_paragraph(
        self,
    ) -> None:
        image = make_image(
            payload=b"invalid-image",
        )

        plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .RENDER_INLINE_IMAGE
                    ),
                    source=image,
                    render_item=(
                        make_image_render_item(
                            image
                        )
                    ),
                ),
                make_paragraph_instruction(
                    order=2,
                    text="Still rendered",
                    top=250.0,
                ),
            ],
        )

        document = self.render_plan(
            plan
        )

        self.assertEqual(
            len(
                document.inline_shapes
            ),
            0,
        )

        self.assertEqual(
            [
                paragraph.text
                for paragraph
                in document.paragraphs
            ],
            [
                "Still rendered",
            ],
        )

        self.assertTrue(
            any(
                warning.startswith(
                    "[inline-export]"
                )
                for warning
                in image.warnings
            )
        )

    def test_deferred_image_only_page_is_safe(
        self,
    ) -> None:
        image = make_image(
            placement=(
                EditableImagePlacement.OVERLAY
            ),
        )

        plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .DEFER_IMAGE
                    ),
                    source=image,
                )
            ],
        )

        document = self.render_plan(
            plan
        )

        self.assertEqual(
            document.paragraphs,
            [],
        )

        self.assertEqual(
            len(
                document.inline_shapes
            ),
            0,
        )


if __name__ == "__main__":
    unittest.main()