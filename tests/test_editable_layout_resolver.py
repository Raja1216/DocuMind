from __future__ import annotations

import unittest

from types import SimpleNamespace

from docx import Document as WordDocument
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
)

from src.exporter.editable_layout_resolver import (
    EditableLayoutResolver,
)
from src.models.alignment_validation import (
    AlignmentValidationCode,
    AlignmentValidationIssue,
    AlignmentValidationReport,
    AlignmentValidationSeverity,
)
from src.models.geometry.rectangle import (
    Rectangle,
)
from src.models.page import Page
from src.models.paragraph_alignment import (
    AlignmentReferenceType,
    ParagraphAlignment,
    ParagraphAlignmentResult,
)
from src.models.reading_order import (
    ReadingOrderEntry,
    ReadingOrderRole,
)


def make_page() -> Page:
    return Page(
        number=1,

        bbox=Rectangle(
            left=0.0,
            top=0.0,
            right=600.0,
            bottom=800.0,
        ),

        rotation=0,
    )


def add_paragraph(
    page: Page,
    number: int,
    text: str,
    top: float,
    detected_alignment: ParagraphAlignment = (
        ParagraphAlignment.LEFT
    ),
    confidence: float = 0.85,
):
    paragraph = SimpleNamespace(
        region_number=number,

        text=text,

        left=50.0,
        top=top,
        right=550.0,
        bottom=top + 30.0,

        reading_order=None,

        detected_alignment=(
            detected_alignment
        ),

        alignment_confidence=confidence,
    )

    page.paragraph_regions.append(
        paragraph
    )

    return paragraph


def add_alignment_result(
    page: Page,
    number: int,
    alignment: ParagraphAlignment,
    confidence: float = 0.85,
):
    result = ParagraphAlignmentResult(
        page_number=page.number,

        paragraph_region_number=number,

        alignment=alignment,

        confidence=confidence,

        reference_type=(
            AlignmentReferenceType.PAGE_BODY
        ),

        reference_id=1,

        paragraph_bbox=Rectangle(
            left=50.0,
            top=100.0,
            right=550.0,
            bottom=150.0,
        ),

        reference_bbox=Rectangle(
            left=50.0,
            top=50.0,
            right=550.0,
            bottom=750.0,
        ),
    )

    page.paragraph_alignment_results.append(
        result
    )

    return result


class EditableLayoutResolverTests(
    unittest.TestCase
):

    def test_reading_order_controls_export_order(
        self,
    ) -> None:
        page = make_page()

        add_paragraph(
            page,
            1,
            "First physical paragraph",
            100.0,
        )

        add_paragraph(
            page,
            2,
            "Second physical paragraph",
            200.0,
        )

        page.reading_order_entries.extend([
            ReadingOrderEntry(
                order=1,
                page_number=1,
                paragraph_region_number=2,
                role=ReadingOrderRole.BODY,
                bbox=Rectangle(
                    left=50.0,
                    top=200.0,
                    right=550.0,
                    bottom=230.0,
                ),
            ),

            ReadingOrderEntry(
                order=2,
                page_number=1,
                paragraph_region_number=1,
                role=ReadingOrderRole.BODY,
                bbox=Rectangle(
                    left=50.0,
                    top=100.0,
                    right=550.0,
                    bottom=130.0,
                ),
            ),
        ])

        plan = (
            EditableLayoutResolver
            .build_page_plan(
                page
            )
        )

        self.assertEqual(
            [
                item.paragraph_region_number
                for item in plan
            ],
            [
                2,
                1,
            ],
        )

    def test_center_alignment_maps_to_word(
        self,
    ) -> None:
        page = make_page()

        add_paragraph(
            page,
            1,
            "Centered title",
            100.0,
        )

        add_alignment_result(
            page=page,
            number=1,
            alignment=(
                ParagraphAlignment.CENTER
            ),
            confidence=0.90,
        )

        plan = (
            EditableLayoutResolver
            .build_page_plan(
                page
            )
        )

        self.assertTrue(
            plan[0].apply_alignment
        )

        self.assertEqual(
            plan[0].word_alignment,
            WD_ALIGN_PARAGRAPH.CENTER,
        )

    def test_justify_alignment_maps_to_word(
        self,
    ) -> None:
        page = make_page()

        add_paragraph(
            page,
            1,
            "Justified paragraph",
            100.0,
        )

        add_alignment_result(
            page=page,
            number=1,
            alignment=(
                ParagraphAlignment.JUSTIFY
            ),
            confidence=0.88,
        )

        plan = (
            EditableLayoutResolver
            .build_page_plan(
                page
            )
        )

        self.assertEqual(
            plan[0].word_alignment,
            WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    def test_low_confidence_alignment_is_not_applied(
        self,
    ) -> None:
        page = make_page()

        add_paragraph(
            page,
            1,
            "Uncertain paragraph",
            100.0,
        )

        add_alignment_result(
            page=page,
            number=1,
            alignment=(
                ParagraphAlignment.CENTER
            ),
            confidence=0.40,
        )

        plan = (
            EditableLayoutResolver
            .build_page_plan(
                page
            )
        )

        self.assertFalse(
            plan[0].apply_alignment
        )

        self.assertIsNone(
            plan[0].word_alignment
        )

    def test_validation_conflict_blocks_alignment(
        self,
    ) -> None:
        page = make_page()

        add_paragraph(
            page,
            1,
            "Invalid centered paragraph",
            100.0,
        )

        add_alignment_result(
            page=page,
            number=1,
            alignment=(
                ParagraphAlignment.CENTER
            ),
            confidence=0.95,
        )

        report = AlignmentValidationReport()

        report.add_issue(
            AlignmentValidationIssue(
                page_number=1,

                paragraph_region_number=1,

                code=(
                    AlignmentValidationCode
                    .CENTER_GEOMETRY_CONFLICT
                ),

                severity=(
                    AlignmentValidationSeverity
                    .WARNING
                ),

                message=(
                    "Centered geometry conflicts."
                ),
            )
        )

        plan = (
            EditableLayoutResolver
            .build_page_plan(
                page=page,
                validation_report=report,
            )
        )

        self.assertFalse(
            plan[0].apply_alignment
        )

        self.assertIsNone(
            plan[0].word_alignment
        )

    def test_unassigned_paragraph_is_preserved(
        self,
    ) -> None:
        page = make_page()

        add_paragraph(
            page,
            1,
            "Ordered paragraph",
            100.0,
        )

        add_paragraph(
            page,
            2,
            "Unassigned paragraph",
            200.0,
        )

        page.reading_order_entries.append(
            ReadingOrderEntry(
                order=1,

                page_number=1,

                paragraph_region_number=1,

                role=ReadingOrderRole.BODY,

                bbox=Rectangle(
                    left=50.0,
                    top=100.0,
                    right=550.0,
                    bottom=130.0,
                ),
            )
        )

        plan = (
            EditableLayoutResolver
            .build_page_plan(
                page
            )
        )

        self.assertEqual(
            len(plan),
            2,
        )

        self.assertEqual(
            plan[1].paragraph_region_number,
            2,
        )

        self.assertEqual(
            plan[1].role,
            ReadingOrderRole.UNASSIGNED,
        )

    def test_apply_alignment_updates_word_paragraph(
        self,
    ) -> None:
        page = make_page()

        add_paragraph(
            page,
            1,
            "Centered paragraph",
            100.0,
        )

        add_alignment_result(
            page=page,
            number=1,
            alignment=(
                ParagraphAlignment.CENTER
            ),
            confidence=0.90,
        )

        plan = (
            EditableLayoutResolver
            .build_page_plan(
                page
            )
        )

        word_document = WordDocument()

        word_paragraph = (
            word_document.add_paragraph(
                "Centered paragraph"
            )
        )

        EditableLayoutResolver.apply_alignment(
            word_paragraph=word_paragraph,
            plan=plan[0],
        )

        self.assertEqual(
            word_paragraph.alignment,
            WD_ALIGN_PARAGRAPH.CENTER,
        )


if __name__ == "__main__":
    unittest.main()