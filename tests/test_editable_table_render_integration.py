from __future__ import annotations

import unittest

from types import SimpleNamespace
from unittest.mock import Mock, patch

from docx import Document as WordDocument
from docx.oxml.ns import qn
from docx.shared import Pt

from src.exporter.docx_exporter import (
    DocxExporter,
)
from src.exporter.editable_page_render_resolver import (
    EditablePageRenderPlan,
    EditablePageRenderResolver,
    EditableRenderAction,
    EditableRenderInstruction,
)
from src.models.editable_table import (
    EditableTable,
    EditableTableCell,
    EditableTableColumn,
    EditableTableDisposition,
    EditableTableRow,
)
from src.models.geometry.rectangle import (
    Rectangle,
)
from src.models.page_render_plan import (
    PageRenderItem,
    PageRenderPlan,
    RenderDisposition,
    RenderItemKind,
    RenderItemRole,
    RenderPlacement,
)


def make_source_table():
    return SimpleNamespace(
        table_number=1,
    )


def make_editable_table(
    *,
    source_table=None,
    disposition=(
        EditableTableDisposition.EDITABLE
    ),
) -> EditableTable:
    table = EditableTable(
        page_number=1,
        table_id="table:1:1",
        bbox=Rectangle(
            left=50.0,
            top=100.0,
            right=350.0,
            bottom=180.0,
        ),
        row_count=2,
        column_count=2,
        disposition=disposition,
        confidence=0.95,
        source_table=source_table,
    )

    table.add_row(
        EditableTableRow(
            row_index=0,
            top=100.0,
            bottom=140.0,
            is_header=True,
            confidence=0.95,
        )
    )

    table.add_row(
        EditableTableRow(
            row_index=1,
            top=140.0,
            bottom=180.0,
            confidence=0.95,
        )
    )

    table.add_column(
        EditableTableColumn(
            column_index=0,
            left=50.0,
            right=200.0,
            confidence=0.95,
        )
    )

    table.add_column(
        EditableTableColumn(
            column_index=1,
            left=200.0,
            right=350.0,
            confidence=0.95,
        )
    )

    values = [
        [
            "Filename",
            "Description",
        ],
        [
            "sample.pdf",
            "Sample file",
        ],
    ]

    for row_index in range(2):
        for column_index in range(2):
            table.add_cell(
                EditableTableCell(
                    row_index=row_index,
                    column_index=column_index,
                    bbox=Rectangle(
                        left=(
                            50.0
                            + column_index
                            * 150.0
                        ),
                        top=(
                            100.0
                            + row_index
                            * 40.0
                        ),
                        right=(
                            50.0
                            + (
                                column_index
                                + 1
                            )
                            * 150.0
                        ),
                        bottom=(
                            100.0
                            + (
                                row_index
                                + 1
                            )
                            * 40.0
                        ),
                    ),
                    text=values[
                        row_index
                    ][
                        column_index
                    ],
                    confidence=0.95,
                )
            )

    return table


def make_table_render_item(
    source_table,
    *,
    disposition=(
        RenderDisposition.EDITABLE
    ),
    placement=(
        RenderPlacement.FLOW
    ),
) -> PageRenderItem:
    return PageRenderItem(
        order=1,
        page_number=1,
        item_id="table:1",
        kind=RenderItemKind.TABLE,
        placement=placement,
        disposition=disposition,
        role=RenderItemRole.BODY,
        bbox=Rectangle(
            left=50.0,
            top=100.0,
            right=350.0,
            bottom=180.0,
        ),
        source=source_table,
        source_index=0,
        confidence=0.95,
    )


class EditableTableRenderIntegrationTests(
    unittest.TestCase
):

    def test_matching_editable_table_becomes_render_instruction(
        self,
    ) -> None:
        source_table = (
            make_source_table()
        )

        editable_table = (
            make_editable_table(
                source_table=source_table
            )
        )

        render_plan = PageRenderPlan(
            page_number=1
        )

        render_plan.add_item(
            make_table_render_item(
                source_table
            )
        )

        page = SimpleNamespace(
            number=1,
            render_plan=render_plan,
            editable_tables=[
                editable_table
            ],
        )

        with patch(
            (
                "src.exporter."
                "editable_page_render_resolver."
                "EditableLayoutResolver."
                "build_page_plan"
            ),
            return_value=[],
        ):
            result = (
                EditablePageRenderResolver
                .build_page_plan(
                    page
                )
            )

        self.assertEqual(
            len(
                result.table_instructions
            ),
            1,
        )

        instruction = (
            result.table_instructions[0]
        )

        self.assertEqual(
            instruction.action,
            EditableRenderAction
            .RENDER_TABLE,
        )

        self.assertIs(
            instruction.source,
            editable_table,
        )

    def test_visual_fallback_table_remains_deferred(
        self,
    ) -> None:
        source_table = (
            make_source_table()
        )

        editable_table = (
            make_editable_table(
                source_table=source_table,
                disposition=(
                    EditableTableDisposition
                    .VISUAL_FALLBACK
                ),
            )
        )

        render_plan = PageRenderPlan(
            page_number=1
        )

        render_plan.add_item(
            make_table_render_item(
                source_table
            )
        )

        page = SimpleNamespace(
            number=1,
            render_plan=render_plan,
            editable_tables=[
                editable_table
            ],
        )

        with patch(
            (
                "src.exporter."
                "editable_page_render_resolver."
                "EditableLayoutResolver."
                "build_page_plan"
            ),
            return_value=[],
        ):
            result = (
                EditablePageRenderResolver
                .build_page_plan(
                    page
                )
            )

        self.assertEqual(
            result.instructions[0].action,
            EditableRenderAction
            .DEFER_TABLE,
        )

        self.assertEqual(
            result.table_instructions,
            [],
        )

    def test_table_only_page_is_rendered(
        self,
    ) -> None:
        word_document = WordDocument()

        editable_table = (
            make_editable_table()
        )

        page_plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .RENDER_TABLE
                    ),
                    source=editable_table,
                )
            ],
        )

        with patch.object(
            DocxExporter,
            "_build_editable_render_plan",
            return_value=page_plan,
        ):
            DocxExporter._render_page(
                word_document=word_document,
                page=SimpleNamespace(
                    number=1,
                    profile=None,
                ),
                numbering_manager=Mock(),
                list_sequence_resolver=Mock(),
            )

        self.assertEqual(
            len(
                word_document.tables
            ),
            1,
        )

        self.assertEqual(
            word_document.tables[0]
            .cell(
                1,
                0,
            )
            .text,
            "sample.pdf",
        )

    def test_paragraph_table_paragraph_order_is_preserved(
        self,
    ) -> None:
        word_document = WordDocument()

        first = SimpleNamespace(
            text="Before table",
            list_type=None,
            left=50.0,
            right=350.0,
        )

        second = SimpleNamespace(
            text="After table",
            list_type=None,
            left=50.0,
            right=350.0,
        )

        first_layout = (
            SimpleNamespace()
        )

        second_layout = (
            SimpleNamespace()
        )

        editable_table = (
            make_editable_table()
        )

        page_plan = EditablePageRenderPlan(
            page_number=1,
            instructions=[
                EditableRenderInstruction(
                    order=1,
                    action=(
                        EditableRenderAction
                        .RENDER_PARAGRAPH
                    ),
                    source=first,
                    layout_item=first_layout,
                ),
                EditableRenderInstruction(
                    order=2,
                    action=(
                        EditableRenderAction
                        .RENDER_TABLE
                    ),
                    source=editable_table,
                ),
                EditableRenderInstruction(
                    order=3,
                    action=(
                        EditableRenderAction
                        .RENDER_PARAGRAPH
                    ),
                    source=second,
                    layout_item=second_layout,
                ),
            ],
        )

        def render_runs(
            *,
            word_paragraph,
            paragraph,
            **_,
        ):
            word_paragraph.add_run(
                paragraph.text
            )

        with (
            patch.object(
                DocxExporter,
                "_build_editable_render_plan",
                return_value=page_plan,
            ),
            patch.object(
                DocxExporter,
                "_region_is_heading",
                return_value=False,
            ),
            patch.object(
                DocxExporter,
                "_apply_region_layout",
            ),
            patch.object(
                DocxExporter,
                "_render_paragraph_runs",
                side_effect=render_runs,
            ),
            patch.object(
                DocxExporter,
                "_normalize_alignment_indentation",
            ),
            patch(
                (
                    "src.exporter."
                    "docx_exporter."
                    "EditableLayoutResolver."
                    "apply_alignment"
                )
            ),
        ):
            DocxExporter._render_page(
                word_document=word_document,
                page=SimpleNamespace(
                    number=1,
                    profile=None,
                ),
                numbering_manager=Mock(),
                list_sequence_resolver=Mock(),
            )

        body_types = [
            element.tag
            for element in (
                word_document
                ._element
                .body
            )
            if element.tag
            in {
                qn(
                    "w:p"
                ),
                qn(
                    "w:tbl"
                ),
            }
        ]

        self.assertEqual(
            body_types,
            [
                qn(
                    "w:p"
                ),
                qn(
                    "w:tbl"
                ),
                qn(
                    "w:p"
                ),
            ],
        )

        self.assertEqual(
            [
                paragraph.text
                for paragraph
                in word_document.paragraphs
            ],
            [
                "Before table",
                "After table",
            ],
        )

    def test_available_width_uses_active_section_margins(
        self,
    ) -> None:
        word_document = WordDocument()

        section = (
            word_document.sections[-1]
        )

        section.page_width = Pt(
            600.0
        )

        section.left_margin = Pt(
            50.0
        )

        section.right_margin = Pt(
            70.0
        )

        self.assertAlmostEqual(
            DocxExporter
            ._resolve_available_table_width(
                word_document
            ),
            480.0,
            places=2,
        )


if __name__ == "__main__":
    unittest.main()
