from __future__ import annotations

import unittest

from docx import Document as WordDocument
from docx.oxml.ns import qn

from src.exporter.docx_exporter import (
    DocxExporter,
)
from src.exporter.editable_list_sequence_resolver import (
    EditableListSequenceResolver,
)
from src.exporter.word_numbering_manager import (
    WordNumberingManager,
)
from src.models.geometry.rectangle import (
    Rectangle,
)
from src.models.list_item import (
    ListMarkerKind,
    ListMarkerSource,
)
from src.models.list_sequence import (
    ListContainerType,
    ListSequence,
    ListSequenceItem,
)
from src.models.page import Page
from types import SimpleNamespace

def make_page_and_sequence():
    page = Page(
        number=1,
        bbox=Rectangle(
            left=0.0,
            top=0.0,
            right=600.0,
            bottom=800.0,
        ),
        rotation=0,
    )

    sequence = ListSequence(
        sequence_id=1,
        page_number=1,
        list_type="bullet",
        container_type=(
            ListContainerType.PAGE_BODY
        ),
        container_id=1,
        container_left=50.0,
        container_right=550.0,
        maximum_level=1,
    )

    sequence.items.extend([
        ListSequenceItem(
            page_number=1,
            paragraph_region_number=1,
            item_index=0,
            level=0,
            marker="•",
            marker_kind=(
                ListMarkerKind.BULLET
            ),
            marker_source=(
                ListMarkerSource.TEXT
            ),
            indent=80.0,
        ),
        ListSequenceItem(
            page_number=1,
            paragraph_region_number=2,
            item_index=1,
            level=1,
            marker="◦",
            marker_kind=(
                ListMarkerKind.BULLET
            ),
            marker_source=(
                ListMarkerSource.TEXT
            ),
            indent=115.0,
        ),
    ])

    page.list_sequences.append(
        sequence
    )

    return page, sequence


class DocxSequenceListExportTests(
    unittest.TestCase
):

    def test_same_num_id_and_different_levels(
        self,
    ) -> None:
        word_document = WordDocument()

        numbering_manager = (
            WordNumberingManager(
                word_document
            )
        )

        resolver = (
            EditableListSequenceResolver(
                numbering_manager
            )
        )

        page, _ = (
            make_page_and_sequence()
        )

        first_source = type(
            "Paragraph",
            (),
            {
                "list_sequence_id": 1,
                "list_level": 0,
                "list_marker_source": (
                    ListMarkerSource.TEXT
                ),
            },
        )()

        second_source = type(
            "Paragraph",
            (),
            {
                "list_sequence_id": 1,
                "list_level": 1,
                "list_marker_source": (
                    ListMarkerSource.TEXT
                ),
            },
        )()

        first_binding = resolver.resolve(
            page=page,
            paragraph=first_source,
            marker_font_name="Arial",
            marker_font_size=11.0,
        )

        second_binding = resolver.resolve(
            page=page,
            paragraph=second_source,
            marker_font_name="Arial",
            marker_font_size=11.0,
        )

        first_word_paragraph = (
            word_document.add_paragraph(
                "Parent item"
            )
        )

        second_word_paragraph = (
            word_document.add_paragraph(
                "Nested item"
            )
        )

        DocxExporter._apply_list_numbering(
            word_paragraph=(
                first_word_paragraph
            ),
            number_id=(
                first_binding.number_id
            ),
            level=(
                first_binding.level
            ),
        )

        DocxExporter._apply_list_numbering(
            word_paragraph=(
                second_word_paragraph
            ),
            number_id=(
                second_binding.number_id
            ),
            level=(
                second_binding.level
            ),
        )

        first_num_properties = (
            first_word_paragraph
            ._p
            .get_or_add_pPr()
            .find(
                qn("w:numPr")
            )
        )

        second_num_properties = (
            second_word_paragraph
            ._p
            .get_or_add_pPr()
            .find(
                qn("w:numPr")
            )
        )

        first_number_id = (
            first_num_properties
            .find(
                qn("w:numId")
            )
            .get(
                qn("w:val")
            )
        )

        second_number_id = (
            second_num_properties
            .find(
                qn("w:numId")
            )
            .get(
                qn("w:val")
            )
        )

        first_level = (
            first_num_properties
            .find(
                qn("w:ilvl")
            )
            .get(
                qn("w:val")
            )
        )

        second_level = (
            second_num_properties
            .find(
                qn("w:ilvl")
            )
            .get(
                qn("w:val")
            )
        )

        self.assertEqual(
            first_number_id,
            second_number_id,
        )

        self.assertEqual(
            first_level,
            "0",
        )

        self.assertEqual(
            second_level,
            "1",
        )

    def test_sequence_list_clears_direct_indentation(
        self,
    ) -> None:
        word_document = WordDocument()

        word_paragraph = (
            word_document.add_paragraph(
                "List item"
            )
        )

        from docx.shared import Pt

        word_paragraph.paragraph_format.left_indent = (
            Pt(100.0)
        )

        word_paragraph.paragraph_format.first_line_indent = (
            Pt(-20.0)
        )

        DocxExporter._clear_direct_list_indentation(
            word_paragraph
        )

        self.assertIsNone(
            word_paragraph
            .paragraph_format
            .left_indent
        )

        self.assertIsNone(
            word_paragraph
            .paragraph_format
            .first_line_indent
        )

    def test_list_marker_style_uses_content_font(
        self,
    ) -> None:
        region = SimpleNamespace(
            list_marker="•",
    
            lines=[
                SimpleNamespace(
                    spans=[
                        SimpleNamespace(
                            text="• ",
                            font="Symbol",
                            font_size=11.0,
                            left=50.0,
                        ),
    
                        SimpleNamespace(
                            text="List item text",
                            font="Arial",
                            font_size=11.0,
                            left=70.0,
                        ),
                    ]
                )
            ],
        )
    
        (
            marker_font_name,
            marker_font_size,
        ) = (
            DocxExporter
            ._resolve_list_marker_style(
                region
            )
        )
    
        self.assertEqual(
            marker_font_name,
            "Arial",
        )
    
        self.assertEqual(
            marker_font_size,
            11.0,
        )

if __name__ == "__main__":
    unittest.main()