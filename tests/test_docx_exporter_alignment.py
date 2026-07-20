from __future__ import annotations

import unittest

from types import SimpleNamespace

from docx import Document as WordDocument
from docx.shared import Pt

from src.exporter.docx_exporter import (
    DocxExporter,
)
from src.models.paragraph_alignment import (
    ParagraphAlignment,
)


class DocxExporterAlignmentTests(
    unittest.TestCase
):

    def test_centered_paragraph_removes_pdf_left_indent(
        self,
    ) -> None:
        word_document = WordDocument()

        word_paragraph = (
            word_document.add_paragraph(
                "Centered title"
            )
        )

        word_paragraph.paragraph_format.left_indent = (
            Pt(150.0)
        )

        word_paragraph.paragraph_format.first_line_indent = (
            Pt(20.0)
        )

        paragraph_plan = SimpleNamespace(
            apply_alignment=True,
            detected_alignment=(
                ParagraphAlignment.CENTER
            ),
        )

        DocxExporter._normalize_alignment_indentation(
            word_paragraph=word_paragraph,
            paragraph_plan=paragraph_plan,
        )

        self.assertEqual(
            word_paragraph
            .paragraph_format
            .left_indent
            .pt,
            0.0,
        )

        self.assertEqual(
            word_paragraph
            .paragraph_format
            .right_indent
            .pt,
            0.0,
        )

        self.assertEqual(
            word_paragraph
            .paragraph_format
            .first_line_indent
            .pt,
            0.0,
        )

    def test_left_paragraph_preserves_indent(
        self,
    ) -> None:
        word_document = WordDocument()

        word_paragraph = (
            word_document.add_paragraph(
                "Indented paragraph"
            )
        )

        word_paragraph.paragraph_format.left_indent = (
            Pt(30.0)
        )

        paragraph_plan = SimpleNamespace(
            apply_alignment=True,
            detected_alignment=(
                ParagraphAlignment.LEFT
            ),
        )

        DocxExporter._normalize_alignment_indentation(
            word_paragraph=word_paragraph,
            paragraph_plan=paragraph_plan,
        )

        self.assertEqual(
            word_paragraph
            .paragraph_format
            .left_indent
            .pt,
            30.0,
        )


if __name__ == "__main__":
    unittest.main()