from __future__ import annotations

import math

from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
)
from docx.shared import (
    Pt,
)

from src.models.editable_image import (
    EditableImage,
    EditableImageDisposition,
    EditableImageHorizontalAlignment,
    EditableImagePayloadStatus,
    EditableImagePlacement,
)


@dataclass(slots=True)
class EditableInlineImageRenderResult:
    """
    Runtime result of rendering one native inline Word image.
    """

    paragraph: Any

    inline_shape: Any

    rendered_width: float

    rendered_height: float


class EditableInlineImageRenderer:
    """
    Render one normalized EditableImage as a native inline
    python-docx image.

    This renderer is intentionally independent from DocxExporter.
    Render-plan integration happens in Step 64G.1.5.
    """

    MINIMUM_RENDER_DIMENSION = 0.10

    MAXIMUM_RENDER_DIMENSION = 1440.0

    ROTATION_TOLERANCE = 0.01

    DEFAULT_ALT_TEXT = (
        "Image imported from PDF"
    )

    @classmethod
    def render(
        cls,
        *,
        container,
        image: EditableImage,
        available_width: float,
        available_height: float | None = None,
    ) -> EditableInlineImageRenderResult:
        """
        Add one paragraph containing a native inline image.

        ``container`` can be a python-docx Document or table cell
        exposing ``add_paragraph()``.

        Width and height are measured in points.
        """

        cls._validate_container(
            container
        )

        cls._validate_image(
            image
        )

        (
            rendered_width,
            rendered_height,
        ) = cls._resolve_dimensions(
            image=image,
            available_width=available_width,
            available_height=(
                available_height
            ),
        )

        paragraph = container.add_paragraph()

        try:
            cls._prepare_paragraph(
                paragraph=paragraph,
                image=image,
            )

            run = paragraph.add_run()

            inline_shape = run.add_picture(
                BytesIO(
                    image.payload
                    or b""
                ),
                width=Pt(
                    rendered_width
                ),
                height=Pt(
                    rendered_height
                ),
            )

            cls._apply_image_metadata(
                inline_shape=inline_shape,
                image=image,
            )

        except Exception as error:
            cls._remove_paragraph(
                paragraph
            )

            raise ValueError(
                (
                    "Unable to render inline Word image "
                    f"{image.image_id!r}: {error}"
                )
            ) from error

        return EditableInlineImageRenderResult(
            paragraph=paragraph,
            inline_shape=inline_shape,
            rendered_width=rendered_width,
            rendered_height=rendered_height,
        )

    # ---------------------------------------------------------
    # Validation
    # ---------------------------------------------------------

    @staticmethod
    def _validate_container(
        container,
    ) -> None:
        if not callable(
            getattr(
                container,
                "add_paragraph",
                None,
            )
        ):
            raise TypeError(
                (
                    "Inline image container must expose "
                    "an add_paragraph() method."
                )
            )

    @classmethod
    def _validate_image(
        cls,
        image: EditableImage,
    ) -> None:
        if (
            image.disposition
            == EditableImageDisposition.SKIP
        ):
            raise ValueError(
                (
                    "Skipped images cannot be rendered "
                    "as native Word images."
                )
            )

        if (
            image.placement
            != EditableImagePlacement.INLINE
        ):
            raise ValueError(
                (
                    "EditableInlineImageRenderer only "
                    "supports INLINE image placement."
                )
            )

        if not image.has_valid_geometry:
            raise ValueError(
                (
                    "Inline image has invalid source "
                    "geometry."
                )
            )

        if not image.has_resolved_payload:
            raise ValueError(
                (
                    "Inline image has no resolved "
                    "renderable payload."
                )
            )

        if (
            image.payload_status
            not in {
                EditableImagePayloadStatus.READY,
                EditableImagePayloadStatus
                .REGION_RENDERED,
            }
        ):
            raise ValueError(
                (
                    "Inline image payload status is "
                    "not renderable."
                )
            )

        if (
            image.payload is None
            or len(
                image.payload
            ) == 0
        ):
            raise ValueError(
                "Inline image payload is empty."
            )

        normalized_rotation = float(
            image.rotation
        ) % 360.0

        if (
            abs(
                normalized_rotation
            )
            > cls.ROTATION_TOLERANCE
            and abs(
                normalized_rotation - 360.0
            )
            > cls.ROTATION_TOLERANCE
        ):
            raise ValueError(
                (
                    "Non-zero image rotation is not "
                    "supported by the inline renderer."
                )
            )

    # ---------------------------------------------------------
    # Dimensions
    # ---------------------------------------------------------

    @classmethod
    def _resolve_dimensions(
        cls,
        *,
        image: EditableImage,
        available_width: float,
        available_height: float | None,
    ) -> tuple[
        float,
        float,
    ]:
        resolved_available_width = (
            cls._positive_finite_number(
                available_width,
                label="available_width",
            )
        )

        resolved_available_height = None

        if available_height is not None:
            resolved_available_height = (
                cls._positive_finite_number(
                    available_height,
                    label="available_height",
                )
            )

        source_width = float(
            image.width
        )

        source_height = float(
            image.height
        )

        if (
            not math.isfinite(
                source_width
            )
            or not math.isfinite(
                source_height
            )
            or source_width
            <= cls.MINIMUM_RENDER_DIMENSION
            or source_height
            <= cls.MINIMUM_RENDER_DIMENSION
        ):
            raise ValueError(
                (
                    "Inline image source dimensions "
                    "are invalid."
                )
            )

        scale_candidates = [
            1.0,
            (
                resolved_available_width
                / source_width
            ),
            (
                cls.MAXIMUM_RENDER_DIMENSION
                / source_width
            ),
            (
                cls.MAXIMUM_RENDER_DIMENSION
                / source_height
            ),
        ]

        if (
            resolved_available_height
            is not None
        ):
            scale_candidates.append(
                resolved_available_height
                / source_height
            )

        scale = min(
            scale_candidates
        )

        if (
            not math.isfinite(
                scale
            )
            or scale <= 0.0
        ):
            raise ValueError(
                (
                    "Unable to calculate a valid "
                    "inline image scale."
                )
            )

        rendered_width = (
            source_width
            * scale
        )

        rendered_height = (
            source_height
            * scale
        )

        if (
            rendered_width
            <= cls.MINIMUM_RENDER_DIMENSION
            or rendered_height
            <= cls.MINIMUM_RENDER_DIMENSION
        ):
            raise ValueError(
                (
                    "Scaled inline image dimensions "
                    "are too small to render."
                )
            )

        return (
            rendered_width,
            rendered_height,
        )

    @staticmethod
    def _positive_finite_number(
        value,
        *,
        label: str,
    ) -> float:
        try:
            normalized = float(
                value
            )

        except (
            TypeError,
            ValueError,
            OverflowError,
        ) as error:
            raise ValueError(
                f"{label} must be a positive number."
            ) from error

        if (
            not math.isfinite(
                normalized
            )
            or normalized <= 0.0
        ):
            raise ValueError(
                f"{label} must be a positive number."
            )

        return normalized

    # ---------------------------------------------------------
    # Word paragraph formatting
    # ---------------------------------------------------------

    @classmethod
    def _prepare_paragraph(
        cls,
        *,
        paragraph,
        image: EditableImage,
    ) -> None:
        paragraph.alignment = (
            cls._resolve_word_alignment(
                image.horizontal_alignment
            )
        )

        paragraph_format = (
            paragraph.paragraph_format
        )

        paragraph_format.space_before = Pt(
            0.0
        )

        paragraph_format.space_after = Pt(
            0.0
        )

        paragraph_format.left_indent = Pt(
            0.0
        )

        paragraph_format.right_indent = Pt(
            0.0
        )

        paragraph_format.first_line_indent = Pt(
            0.0
        )

        paragraph_format.keep_together = True

        paragraph_format.keep_with_next = (
            False
        )

    @staticmethod
    def _resolve_word_alignment(
        alignment: (
            EditableImageHorizontalAlignment
        ),
    ):
        mapping = {
            EditableImageHorizontalAlignment.LEFT: (
                WD_ALIGN_PARAGRAPH.LEFT
            ),
            EditableImageHorizontalAlignment.CENTER: (
                WD_ALIGN_PARAGRAPH.CENTER
            ),
            EditableImageHorizontalAlignment.RIGHT: (
                WD_ALIGN_PARAGRAPH.RIGHT
            ),
            EditableImageHorizontalAlignment.ABSOLUTE: (
                WD_ALIGN_PARAGRAPH.LEFT
            ),
        }

        return mapping.get(
            alignment,
            WD_ALIGN_PARAGRAPH.LEFT,
        )

    # ---------------------------------------------------------
    # Image properties
    # ---------------------------------------------------------

    @classmethod
    def _apply_image_metadata(
        cls,
        *,
        inline_shape,
        image: EditableImage,
    ) -> None:
        inline = getattr(
            inline_shape,
            "_inline",
            None,
        )

        if inline is None:
            return

        document_properties = getattr(
            inline,
            "docPr",
            None,
        )

        if document_properties is None:
            return

        image_name = str(
            image.image_id
        ).strip()

        if not image_name:
            image_name = (
                cls.DEFAULT_ALT_TEXT
            )

        document_properties.set(
            "name",
            image_name,
        )

        document_properties.set(
            "title",
            image_name,
        )

        document_properties.set(
            "descr",
            cls.DEFAULT_ALT_TEXT,
        )

    # ---------------------------------------------------------
    # Rollback
    # ---------------------------------------------------------

    @staticmethod
    def _remove_paragraph(
        paragraph,
    ) -> None:
        paragraph_element = getattr(
            paragraph,
            "_p",
            None,
        )

        if paragraph_element is None:
            return

        parent = (
            paragraph_element.getparent()
        )

        if parent is not None:
            parent.remove(
                paragraph_element
            )