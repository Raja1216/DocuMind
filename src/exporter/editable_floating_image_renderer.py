from __future__ import annotations

import math

from dataclasses import dataclass
from enum import Enum
from io import BytesIO
from typing import Any

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
)
from docx.oxml import (
    OxmlElement,
)
from docx.oxml.ns import (
    qn,
)
from docx.shared import (
    Pt,
)

from src.models.editable_image import (
    EditableImage,
    EditableImageDisposition,
    EditableImagePayloadStatus,
    EditableImagePlacement,
    EditableImageRole,
)
from src.models.geometry.rectangle import (
    Rectangle,
)


class EditableFloatingImageWrapMode(
    str,
    Enum,
):
    """
    Word text-wrapping strategy for a floating image.
    """

    AUTO = "auto"

    NONE = "none"

    SQUARE = "square"


@dataclass(slots=True)
class EditableFloatingImageRenderResult:
    """
    Runtime result of rendering one floating Word image.
    """

    paragraph: Any

    anchor: Any

    rendered_width: float

    rendered_height: float

    horizontal_offset: float

    vertical_offset: float

    wrap_mode: EditableFloatingImageWrapMode


class EditableFloatingImageRenderer:
    """
    Render one resolved EditableImage as a floating Word drawing.

    PDF and Word page coordinates both use points. The image's
    source rectangle is therefore converted into page-relative
    Word DrawingML offsets without sample-specific coordinates.

    Render-plan integration occurs in Step 64G.1.7.
    """

    POINT_TO_EMU = 12_700

    MINIMUM_RENDER_DIMENSION = 0.10

    MAXIMUM_RENDER_DIMENSION = 1_440.0

    DEFAULT_WRAP_DISTANCE = 3.0

    ROTATION_TOLERANCE = 0.01

    DEFAULT_ALT_TEXT = (
        "Floating image imported from PDF"
    )

    BASE_RELATIVE_HEIGHT = 251_658_240

    @classmethod
    def render(
        cls,
        *,
        container,
        image: EditableImage,
        page_bbox: Rectangle,
        wrap_mode: (
            EditableFloatingImageWrapMode
        ) = EditableFloatingImageWrapMode.AUTO,
    ) -> EditableFloatingImageRenderResult:
        """
        Insert a floating Word image positioned relative to page.

        ``container`` must expose add_paragraph(), normally a
        python-docx Document body.
        """

        cls._validate_container(
            container
        )

        cls._validate_image(
            image
        )

        cls._validate_page_bbox(
            page_bbox
        )

        resolved_wrap_mode = (
            cls._resolve_wrap_mode(
                image=image,
                requested_mode=wrap_mode,
            )
        )

        (
            horizontal_offset,
            vertical_offset,
            rendered_width,
            rendered_height,
        ) = cls._resolve_geometry(
            image=image,
            page_bbox=page_bbox,
        )

        paragraph = container.add_paragraph()

        try:
            cls._prepare_anchor_paragraph(
                paragraph
            )

            run = paragraph.add_run()

            inline_shape = run.add_picture(
                BytesIO(
                    image.payload
                    or b""
                ),
                width=Pt(
                    rendered_width
                ),
                height=Pt(
                    rendered_height
                ),
            )

            anchor = cls._convert_inline_to_anchor(
                inline_shape=inline_shape,
                image=image,
                horizontal_offset=(
                    horizontal_offset
                ),
                vertical_offset=(
                    vertical_offset
                ),
                wrap_mode=(
                    resolved_wrap_mode
                ),
            )

        except Exception as error:
            cls._remove_paragraph(
                paragraph
            )

            raise ValueError(
                (
                    "Unable to render floating Word "
                    f"image {image.image_id!r}: "
                    f"{error}"
                )
            ) from error

        return EditableFloatingImageRenderResult(
            paragraph=paragraph,
            anchor=anchor,
            rendered_width=rendered_width,
            rendered_height=rendered_height,
            horizontal_offset=(
                horizontal_offset
            ),
            vertical_offset=(
                vertical_offset
            ),
            wrap_mode=resolved_wrap_mode,
        )

    # ---------------------------------------------------------
    # Validation
    # ---------------------------------------------------------

    @staticmethod
    def _validate_container(
        container,
    ) -> None:
        if not callable(
            getattr(
                container,
                "add_paragraph",
                None,
            )
        ):
            raise TypeError(
                (
                    "Floating image container must "
                    "expose add_paragraph()."
                )
            )

    @classmethod
    def _validate_image(
        cls,
        image: EditableImage,
    ) -> None:
        if (
            image.disposition
            == EditableImageDisposition.SKIP
        ):
            raise ValueError(
                "Skipped images cannot be rendered."
            )

        if (
            image.placement
            != EditableImagePlacement.FLOATING
        ):
            raise ValueError(
                (
                    "EditableFloatingImageRenderer "
                    "only supports FLOATING placement."
                )
            )

        if not image.has_valid_geometry:
            raise ValueError(
                "Floating image geometry is invalid."
            )

        if not image.has_resolved_payload:
            raise ValueError(
                (
                    "Floating image has no resolved "
                    "renderable payload."
                )
            )

        if (
            image.payload_status
            not in {
                EditableImagePayloadStatus.READY,
                EditableImagePayloadStatus
                .REGION_RENDERED,
            }
        ):
            raise ValueError(
                (
                    "Floating image payload status "
                    "is not renderable."
                )
            )

        if not image.payload:
            raise ValueError(
                "Floating image payload is empty."
            )

        try:
            normalized_rotation = (
                float(
                    image.rotation
                )
                % 360.0
            )

        except (
            TypeError,
            ValueError,
            OverflowError,
        ):
            normalized_rotation = 0.0

        rotation_distance = min(
            abs(
                normalized_rotation
            ),
            abs(
                360.0
                - normalized_rotation
            ),
        )

        if (
            rotation_distance
            > cls.ROTATION_TOLERANCE
        ):
            raise ValueError(
                (
                    "Rotated floating images are "
                    "deferred to a later transform step."
                )
            )

    @classmethod
    def _validate_page_bbox(
        cls,
        page_bbox: Rectangle,
    ) -> None:
        page_width = (
            float(
                page_bbox.right
            )
            - float(
                page_bbox.left
            )
        )

        page_height = (
            float(
                page_bbox.bottom
            )
            - float(
                page_bbox.top
            )
        )

        if (
            not math.isfinite(
                page_width
            )
            or not math.isfinite(
                page_height
            )
            or page_width
            <= cls.MINIMUM_RENDER_DIMENSION
            or page_height
            <= cls.MINIMUM_RENDER_DIMENSION
        ):
            raise ValueError(
                "Source page geometry is invalid."
            )

    # ---------------------------------------------------------
    # Source-page geometry
    # ---------------------------------------------------------

    @classmethod
    def _resolve_geometry(
        cls,
        *,
        image: EditableImage,
        page_bbox: Rectangle,
    ) -> tuple[
        float,
        float,
        float,
        float,
    ]:
        page_width = max(
            float(
                page_bbox.right
            )
            - float(
                page_bbox.left
            ),
            cls.MINIMUM_RENDER_DIMENSION,
        )

        page_height = max(
            float(
                page_bbox.bottom
            )
            - float(
                page_bbox.top
            ),
            cls.MINIMUM_RENDER_DIMENSION,
        )

        horizontal_offset = max(
            float(
                image.bbox.left
            )
            - float(
                page_bbox.left
            ),
            0.0,
        )

        vertical_offset = max(
            float(
                image.bbox.top
            )
            - float(
                page_bbox.top
            ),
            0.0,
        )

        source_width = float(
            image.width
        )

        source_height = float(
            image.height
        )

        available_width = max(
            page_width
            - horizontal_offset,
            cls.MINIMUM_RENDER_DIMENSION,
        )

        available_height = max(
            page_height
            - vertical_offset,
            cls.MINIMUM_RENDER_DIMENSION,
        )

        scale = min(
            1.0,
            available_width
            / source_width,
            available_height
            / source_height,
            cls.MAXIMUM_RENDER_DIMENSION
            / source_width,
            cls.MAXIMUM_RENDER_DIMENSION
            / source_height,
        )

        if (
            not math.isfinite(
                scale
            )
            or scale <= 0.0
        ):
            raise ValueError(
                (
                    "Unable to calculate valid "
                    "floating-image dimensions."
                )
            )

        rendered_width = (
            source_width
            * scale
        )

        rendered_height = (
            source_height
            * scale
        )

        if (
            rendered_width
            <= cls.MINIMUM_RENDER_DIMENSION
            or rendered_height
            <= cls.MINIMUM_RENDER_DIMENSION
        ):
            raise ValueError(
                (
                    "Scaled floating-image dimensions "
                    "are too small."
                )
            )

        return (
            horizontal_offset,
            vertical_offset,
            rendered_width,
            rendered_height,
        )

    # ---------------------------------------------------------
    # Wrapping
    # ---------------------------------------------------------

    @staticmethod
    def _resolve_wrap_mode(
        *,
        image: EditableImage,
        requested_mode: (
            EditableFloatingImageWrapMode
        ),
    ) -> EditableFloatingImageWrapMode:
        if (
            requested_mode
            != EditableFloatingImageWrapMode.AUTO
        ):
            return requested_mode

        if image.role == EditableImageRole.CONTENT:
            return (
                EditableFloatingImageWrapMode
                .SQUARE
            )

        return (
            EditableFloatingImageWrapMode.NONE
        )

    # ---------------------------------------------------------
    # DrawingML conversion
    # ---------------------------------------------------------

    @classmethod
    def _convert_inline_to_anchor(
        cls,
        *,
        inline_shape,
        image: EditableImage,
        horizontal_offset: float,
        vertical_offset: float,
        wrap_mode: (
            EditableFloatingImageWrapMode
        ),
    ):
        inline = getattr(
            inline_shape,
            "_inline",
            None,
        )

        if inline is None:
            raise ValueError(
                (
                    "python-docx did not expose the "
                    "generated inline drawing."
                )
            )

        drawing = inline.getparent()

        if drawing is None:
            raise ValueError(
                (
                    "Generated inline drawing has "
                    "no DrawingML parent."
                )
            )

        extent = inline.find(
            qn(
                "wp:extent"
            )
        )

        effect_extent = inline.find(
            qn(
                "wp:effectExtent"
            )
        )

        document_properties = inline.find(
            qn(
                "wp:docPr"
            )
        )

        frame_properties = inline.find(
            qn(
                "wp:cNvGraphicFramePr"
            )
        )

        graphic = inline.find(
            qn(
                "a:graphic"
            )
        )

        required_children = {
            "extent": extent,
            "document properties": (
                document_properties
            ),
            "graphic frame properties": (
                frame_properties
            ),
            "graphic": graphic,
        }

        for label, child in (
            required_children.items()
        ):
            if child is None:
                raise ValueError(
                    (
                        "Generated image is missing "
                        f"{label}."
                    )
                )

        anchor = OxmlElement(
            "wp:anchor"
        )

        wrap_distance_emu = (
            cls._points_to_emu(
                cls.DEFAULT_WRAP_DISTANCE
            )
            if (
                wrap_mode
                == EditableFloatingImageWrapMode
                .SQUARE
            )
            else 0
        )

        anchor.set(
            "distT",
            str(
                wrap_distance_emu
            ),
        )

        anchor.set(
            "distB",
            str(
                wrap_distance_emu
            ),
        )

        anchor.set(
            "distL",
            str(
                wrap_distance_emu
            ),
        )

        anchor.set(
            "distR",
            str(
                wrap_distance_emu
            ),
        )

        anchor.set(
            "simplePos",
            "0",
        )

        doc_property_id = cls._resolve_doc_property_id(
            document_properties
        )

        anchor.set(
            "relativeHeight",
            str(
                cls.BASE_RELATIVE_HEIGHT
                + doc_property_id
            ),
        )

        anchor.set(
            "behindDoc",
            "0",
        )

        anchor.set(
            "locked",
            "0",
        )

        anchor.set(
            "layoutInCell",
            "1",
        )

        anchor.set(
            "allowOverlap",
            "1",
        )

        simple_position = OxmlElement(
            "wp:simplePos"
        )

        simple_position.set(
            "x",
            "0",
        )

        simple_position.set(
            "y",
            "0",
        )

        horizontal_position = (
            OxmlElement(
                "wp:positionH"
            )
        )

        horizontal_position.set(
            "relativeFrom",
            "page",
        )

        horizontal_position_offset = (
            OxmlElement(
                "wp:posOffset"
            )
        )

        horizontal_position_offset.text = str(
            cls._points_to_emu(
                horizontal_offset
            )
        )

        horizontal_position.append(
            horizontal_position_offset
        )

        vertical_position = OxmlElement(
            "wp:positionV"
        )

        vertical_position.set(
            "relativeFrom",
            "page",
        )

        vertical_position_offset = (
            OxmlElement(
                "wp:posOffset"
            )
        )

        vertical_position_offset.text = str(
            cls._points_to_emu(
                vertical_offset
            )
        )

        vertical_position.append(
            vertical_position_offset
        )

        cls._detach_child(
            extent
        )

        if effect_extent is not None:
            cls._detach_child(
                effect_extent
            )

        cls._detach_child(
            document_properties
        )

        cls._detach_child(
            frame_properties
        )

        cls._detach_child(
            graphic
        )

        anchor.append(
            simple_position
        )

        anchor.append(
            horizontal_position
        )

        anchor.append(
            vertical_position
        )

        anchor.append(
            extent
        )

        if effect_extent is not None:
            anchor.append(
                effect_extent
            )

        if (
            wrap_mode
            == EditableFloatingImageWrapMode
            .SQUARE
        ):
            wrap_element = OxmlElement(
                "wp:wrapSquare"
            )

            wrap_element.set(
                "wrapText",
                "bothSides",
            )

        else:
            wrap_element = OxmlElement(
                "wp:wrapNone"
            )

        anchor.append(
            wrap_element
        )

        cls._apply_image_metadata(
            document_properties=(
                document_properties
            ),
            image=image,
        )

        anchor.append(
            document_properties
        )

        anchor.append(
            frame_properties
        )

        anchor.append(
            graphic
        )

        drawing.replace(
            inline,
            anchor,
        )

        return anchor

    @classmethod
    def _apply_image_metadata(
        cls,
        *,
        document_properties,
        image: EditableImage,
    ) -> None:
        image_name = str(
            image.image_id
            or ""
        ).strip()

        if not image_name:
            image_name = (
                cls.DEFAULT_ALT_TEXT
            )

        document_properties.set(
            "name",
            image_name,
        )

        document_properties.set(
            "title",
            image_name,
        )

        document_properties.set(
            "descr",
            cls.DEFAULT_ALT_TEXT,
        )

    # ---------------------------------------------------------
    # Word paragraph
    # ---------------------------------------------------------

    @staticmethod
    def _prepare_anchor_paragraph(
        paragraph,
    ) -> None:
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        paragraph_format = (
            paragraph.paragraph_format
        )

        paragraph_format.space_before = Pt(
            0.0
        )

        paragraph_format.space_after = Pt(
            0.0
        )

        paragraph_format.left_indent = Pt(
            0.0
        )

        paragraph_format.right_indent = Pt(
            0.0
        )

        paragraph_format.first_line_indent = Pt(
            0.0
        )

        paragraph_format.keep_together = True

    # ---------------------------------------------------------
    # Primitive helpers
    # ---------------------------------------------------------

    @classmethod
    def _points_to_emu(
        cls,
        value: float,
    ) -> int:
        return int(
            round(
                float(
                    value
                )
                * cls.POINT_TO_EMU
            )
        )

    @staticmethod
    def _resolve_doc_property_id(
        document_properties,
    ) -> int:
        try:
            return max(
                int(
                    document_properties.get(
                        "id"
                    )
                    or 1
                ),
                1,
            )

        except (
            TypeError,
            ValueError,
            OverflowError,
        ):
            return 1

    @staticmethod
    def _detach_child(
        child,
    ) -> None:
        parent = child.getparent()

        if parent is not None:
            parent.remove(
                child
            )

    @staticmethod
    def _remove_paragraph(
        paragraph,
    ) -> None:
        paragraph_element = getattr(
            paragraph,
            "_p",
            None,
        )

        if paragraph_element is None:
            return

        parent = (
            paragraph_element.getparent()
        )

        if parent is not None:
            parent.remove(
                paragraph_element
            )