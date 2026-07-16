from __future__ import annotations

from docx import Document as WordDocument
from docx.dml.color import RGBColor as WordRGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree
from io import BytesIO

from src.exporter.font_name_resolver import (
    FontNameResolver,
)
from src.exporter.text_normalizer import TextNormalizer


class FixedLayoutDocxExporter:
    """
    Export a PDF document model as a fixed-layout DOCX.

    Every visible PDF text span is placed inside an
    absolutely positioned Word text box.

    This mode prioritizes visual fidelity and original
    pagination over normal Word paragraph reflow.
    """

    VML_NAMESPACE = (
        "urn:schemas-microsoft-com:vml"
    )

    OFFICE_NAMESPACE = (
        "urn:schemas-microsoft-com:office:office"
    )

    WORD_2003_NAMESPACE = (
        "urn:schemas-microsoft-com:office:word"
    )

    WORD_NAMESPACE = (
        "http://schemas.openxmlformats.org/"
        "wordprocessingml/2006/main"
    )
    
    RELATIONSHIPS_NAMESPACE = (
        "http://schemas.openxmlformats.org/"
        "officeDocument/2006/relationships"
    )

    VML_NAMESPACE_MAP = {
        "v": VML_NAMESPACE,
        "o": OFFICE_NAMESPACE,
        "w10": WORD_2003_NAMESPACE,
        "w": WORD_NAMESPACE,
        "r": RELATIONSHIPS_NAMESPACE,
    }
    
    TEXTBOX_WIDTH_PADDING = 3.0
    TEXTBOX_HEIGHT_PADDING = 6.0

    TEXTBOX_TOP_ADJUSTMENT = 1.5
    LINE_HEIGHT_FACTOR = 1.25
    
    BULLET_FONT_NAME = "Arial"
    BULLET_SIZE_FACTOR = 1.0

    ANCHOR_FONT_SIZE = 1.0
    ANCHOR_LINE_SPACING = 1.0
    
    TABLE_BORDER_COLOR = "#B7B7B7"
    TABLE_BORDER_THICKNESS = 0.50

    TABLE_GRID_TOLERANCE = 0.5
    MINIMUM_VISIBLE_TABLE_BORDER = 0.75
    TABLE_GRID_Z_INDEX = 110
    TABLE_CELL_FILL_Z_INDEX = 80

    TABLE_CELL_HORIZONTAL_PADDING = 2.0
    TABLE_CELL_VERTICAL_PADDING = 1.0

    DEFAULT_TABLE_FONT = "Arial"
    DEFAULT_TABLE_FONT_SIZE = 9.0
    
    TABLE_LINE_TOLERANCE = 2.0
    TABLE_ALIGNMENT_TOLERANCE = 4.0

    TABLE_CENTER_MIN_MARGIN = 4.0
    TABLE_RIGHT_ALIGNMENT_RATIO = 0.65
    TABLE_ALIGNMENT_REQUIRED_RATIO = 0.75
    TABLE_VERTICAL_ALIGNMENT_TOLERANCE = 3.0

    TABLE_VERTICAL_TOP = "top"
    TABLE_VERTICAL_CENTER = "middle"
    TABLE_VERTICAL_BOTTOM = "bottom"


    TABLE_Z_INDEX = 100
    TEXT_Z_INDEX = 251659264
    IMAGE_Z_INDEX = 1

    _shape_id = 1

    @staticmethod
    def export(
        document,
        output_path: str,
    ) -> None:
        word_document = WordDocument()

        FixedLayoutDocxExporter._shape_id = 1

        for page_index, page in enumerate(
            document.pages
        ):
            section = (
                FixedLayoutDocxExporter
                ._prepare_page_section(
                    word_document=word_document,
                    page_index=page_index,
                )
            )

            FixedLayoutDocxExporter._configure_section(
                section=section,
                page=page,
            )

            anchor_paragraph = (
                FixedLayoutDocxExporter
                ._create_anchor_paragraph(
                    word_document
                )
            )

            FixedLayoutDocxExporter._render_page(
                word_document=word_document,
                anchor_paragraph=anchor_paragraph,
                page=page,
            )

        word_document.save(
            output_path
        )

    @staticmethod
    def _prepare_page_section(
        word_document,
        page_index: int,
    ):
        """
        Use the original first section for page 1.

        Every later PDF page gets a new Word section
        beginning on a new page.
        """

        if page_index == 0:
            return word_document.sections[0]

        return word_document.add_section(
            WD_SECTION.NEW_PAGE
        )

    @staticmethod
    def _configure_section(
        section,
        page,
    ) -> None:
        """
        Match the PDF page dimensions.

        Fixed-layout positioning is relative to the physical
        page, so all Word margins are set to zero.
        """

        section.page_width = Pt(
            page.bbox.width
        )

        section.page_height = Pt(
            page.bbox.height
        )

        zero = Pt(0)

        section.left_margin = zero
        section.right_margin = zero
        section.top_margin = zero
        section.bottom_margin = zero

        section.header_distance = zero
        section.footer_distance = zero
        section.gutter = zero

    @staticmethod
    def _create_anchor_paragraph(
        word_document,
    ):
        """
        Create a minimal paragraph that owns all floating
        text boxes on the current Word page.
        """

        paragraph = word_document.add_paragraph()

        paragraph_format = paragraph.paragraph_format

        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        paragraph_format.line_spacing_rule = (
            WD_LINE_SPACING.EXACTLY
        )

        paragraph_format.line_spacing = Pt(
            FixedLayoutDocxExporter
            .ANCHOR_LINE_SPACING
        )

        anchor_run = paragraph.add_run("")

        anchor_run.font.size = Pt(
            FixedLayoutDocxExporter
            .ANCHOR_FONT_SIZE
        )

        return paragraph

    @staticmethod
    def _render_page(
        word_document,
        anchor_paragraph,
        page,
    ) -> None:
        """
        Render all supported page elements in fixed layout.

        Layer order:
        1. Images
        2. Tables
        3. Normal text
        """

        for image in page.images:
            FixedLayoutDocxExporter._add_image_shape(
                document_part=word_document.part,
                anchor_paragraph=anchor_paragraph,
                image=image,
            )

        for table in page.tables:
            FixedLayoutDocxExporter._render_table(
                anchor_paragraph=anchor_paragraph,
                page=page,
                table=table,
            )

        for block in page.blocks:
            for line in block.lines:
                for span in line.spans:

                    if not span.text.strip():
                        continue

                    if (
                        FixedLayoutDocxExporter
                        ._span_is_inside_any_table(
                            span=span,
                            tables=page.tables,
                        )
                    ):
                        continue

                    FixedLayoutDocxExporter._add_span_textbox(
                        anchor_paragraph=anchor_paragraph,
                        span=span,
                    )

    @staticmethod
    def _render_table(
        anchor_paragraph,
        page,
        table,
    ) -> None:
        """
        Render one detected table.

        Grid lines are rendered only once. Cell text boxes are
        borderless and placed above the grid.
        """

        FixedLayoutDocxExporter._render_table_grid(
            anchor_paragraph=anchor_paragraph,
            table=table,
        )

        for cell in table.cells:

            cell_spans = (
                FixedLayoutDocxExporter
                ._find_spans_for_cell(
                    page=page,
                    cell=cell,
                )
            )

            FixedLayoutDocxExporter._add_table_cell_shape(
                anchor_paragraph=anchor_paragraph,
                cell=cell,
                cell_spans=cell_spans,
            )

    @staticmethod
    def _render_table_grid(
        anchor_paragraph,
        table,
    ) -> None:
        """
        Render every unique horizontal and vertical table grid
        segment exactly once.
        """

        horizontal_segments = (
            FixedLayoutDocxExporter
            ._collect_horizontal_table_segments(
                table
            )
        )

        vertical_segments = (
            FixedLayoutDocxExporter
            ._collect_vertical_table_segments(
                table
            )
        )

        for left, right, y_position in horizontal_segments:
            FixedLayoutDocxExporter._add_table_grid_line(
                anchor_paragraph=anchor_paragraph,
                left=left,
                top=y_position,
                right=right,
                bottom=y_position,
                border_color=table.border_style.color,
                border_thickness=table.border_style.thickness,
            )

        for x_position, top, bottom in vertical_segments:
            FixedLayoutDocxExporter._add_table_grid_line(
                anchor_paragraph=anchor_paragraph,
                left=x_position,
                top=top,
                right=x_position,
                bottom=bottom,
                border_color=table.border_style.color,
                border_thickness=table.border_style.thickness,
            )

    @staticmethod
    def _collect_horizontal_table_segments(
        table,
    ) -> list[tuple[float, float, float]]:
        """
        Collect and merge unique horizontal cell-border
        segments.

        Returned tuple:
            left, right, y
        """

        raw_segments = []

        for cell in table.cells:
            raw_segments.append(
                (
                    cell.left,
                    cell.right,
                    cell.top,
                )
            )

            raw_segments.append(
                (
                    cell.left,
                    cell.right,
                    cell.bottom,
                )
            )

        return (
            FixedLayoutDocxExporter
            ._merge_horizontal_segments(
                raw_segments
            )
        )

    @staticmethod
    def _collect_vertical_table_segments(
        table,
    ) -> list[tuple[float, float, float]]:
        """
        Collect and merge unique vertical cell-border segments.

        Returned tuple:
            x, top, bottom
        """

        raw_segments = []

        for cell in table.cells:
            raw_segments.append(
                (
                    cell.left,
                    cell.top,
                    cell.bottom,
                )
            )

            raw_segments.append(
                (
                    cell.right,
                    cell.top,
                    cell.bottom,
                )
            )

        return (
            FixedLayoutDocxExporter
            ._merge_vertical_segments(
                raw_segments
            )
        )

    @staticmethod
    def _merge_horizontal_segments(
        segments: list[tuple[float, float, float]],
    ) -> list[tuple[float, float, float]]:
        """
        Merge touching or overlapping horizontal segments that
        share approximately the same Y coordinate.
        """

        tolerance = (
            FixedLayoutDocxExporter
            .TABLE_GRID_TOLERANCE
        )

        grouped: list[list[float]] = []

        for left, right, y_position in sorted(
            segments,
            key=lambda item: (
                item[2],
                item[0],
            ),
        ):
            matching_group = None

            for group in grouped:
                if abs(group[2] - y_position) <= tolerance:
                    matching_group = group
                    break

            if matching_group is None:
                grouped.append(
                    [
                        left,
                        right,
                        y_position,
                    ]
                )
                continue

            if left <= matching_group[1] + tolerance:
                matching_group[0] = min(
                    matching_group[0],
                    left,
                )

                matching_group[1] = max(
                    matching_group[1],
                    right,
                )
            else:
                grouped.append(
                    [
                        left,
                        right,
                        y_position,
                    ]
                )

        return [
            (
                float(group[0]),
                float(group[1]),
                float(group[2]),
            )
            for group in grouped
        ]
        
    @staticmethod
    def _merge_vertical_segments(
        segments: list[tuple[float, float, float]],
    ) -> list[tuple[float, float, float]]:
        """
        Merge touching or overlapping vertical segments that
        share approximately the same X coordinate.
        """

        tolerance = (
            FixedLayoutDocxExporter
            .TABLE_GRID_TOLERANCE
        )

        grouped: list[list[float]] = []

        for x_position, top, bottom in sorted(
            segments,
            key=lambda item: (
                item[0],
                item[1],
            ),
        ):
            matching_group = None

            for group in grouped:
                if abs(group[0] - x_position) <= tolerance:
                    matching_group = group
                    break

            if matching_group is None:
                grouped.append(
                    [
                        x_position,
                        top,
                        bottom,
                    ]
                )
                continue

            if top <= matching_group[2] + tolerance:
                matching_group[1] = min(
                    matching_group[1],
                    top,
                )

                matching_group[2] = max(
                    matching_group[2],
                    bottom,
                )
            else:
                grouped.append(
                    [
                        x_position,
                        top,
                        bottom,
                    ]
                )

        return [
            (
                float(group[0]),
                float(group[1]),
                float(group[2]),
            )
            for group in grouped
        ]    

    @staticmethod
    def _add_table_grid_line(
        anchor_paragraph,
        left: float,
        top: float,
        right: float,
        bottom: float,
        border_color: str,
        border_thickness: float,
    ) -> None:
        """
        Render one table border as a very thin filled VML
        rectangle.

        Rectangle shapes are more reliably rendered by Word
        than absolutely positioned VML line elements.
        """

        thickness = max(
            border_thickness,
            FixedLayoutDocxExporter
            .MINIMUM_VISIBLE_TABLE_BORDER,
        )

        is_horizontal = abs(
            bottom - top
        ) <= (
            FixedLayoutDocxExporter
            .TABLE_GRID_TOLERANCE
        )

        if is_horizontal:
            width = max(
                right - left,
                thickness,
            )

            height = thickness

            adjusted_left = left
            adjusted_top = top - (thickness / 2)

        else:
            width = thickness

            height = max(
                bottom - top,
                thickness,
            )

            adjusted_left = left - (thickness / 2)
            adjusted_top = top
            
        adjusted_left = round(
            adjusted_left,
            2,
        )

        adjusted_top = round(
            adjusted_top,
            2,
        )

        width = round(
            width,
            2,
        )

        height = round(
            height,
            2,
        )    

        shape_id = (
            FixedLayoutDocxExporter._next_shape_id()
        )

        shape = (
            FixedLayoutDocxExporter
            ._create_table_border_shape(
                shape_id=shape_id,
                left=adjusted_left,
                top=adjusted_top,
                width=width,
                height=height,
                border_color=border_color,
            )
        )

        pict = OxmlElement(
            "w:pict"
        )

        pict.append(
            shape
        )

        anchor_run = (
            anchor_paragraph.add_run()
        )

        anchor_run._r.append(
            pict
        )
        
    @staticmethod
    def _create_table_border_shape(
        shape_id: int,
        left: float,
        top: float,
        width: float,
        height: float,
        border_color: str,
    ):
        """
        Create one thin, filled VML rectangle used as a
        table border.
        """

        shape = (
            FixedLayoutDocxExporter
            ._create_vml_element(
                "rect"
            )
        )

        shape.set(
            "id",
            f"DocuMindTableBorder{shape_id}",
        )

        shape.set(
            "filled",
            "t",
        )

        shape.set(
            "fillcolor",
            border_color,
        )

        shape.set(
            "stroked",
            "f",
        )

        shape.set(
            (
                f"{{"
                f"{FixedLayoutDocxExporter.OFFICE_NAMESPACE}"
                f"}}allowincell"
            ),
            "f",
        )

        style = (
            "position:absolute;"
            f"margin-left:{left:.3f}pt;"
            f"margin-top:{top:.3f}pt;"
            f"width:{width:.3f}pt;"
            f"height:{height:.3f}pt;"
            f"z-index:{FixedLayoutDocxExporter.TABLE_GRID_Z_INDEX};"
            "mso-position-horizontal-relative:page;"
            "mso-position-vertical-relative:page;"
            "mso-wrap-style:none;"
        )

        shape.set(
            "style",
            style,
        )

        wrap = (
            FixedLayoutDocxExporter
            ._create_word_2003_element(
                "wrap"
            )
        )

        wrap.set(
            "type",
            "none",
        )

        shape.append(
            wrap
        )

        return shape    

    @staticmethod
    def _add_table_cell_shape(
        anchor_paragraph,
        cell,
        cell_spans: list,
    ) -> None:
        """
        Add one positioned table cell with border and text.
        """

        width = max(
            cell.right - cell.left,
            1.0,
        )

        height = max(
            cell.bottom - cell.top,
            1.0,
        )
        
        vertical_alignment = (
            FixedLayoutDocxExporter
            ._detect_table_cell_vertical_alignment(
                cell=cell,
                cell_spans=cell_spans,
            )
        )

        shape_id = (
            FixedLayoutDocxExporter._next_shape_id()
        )

        shape = (
            FixedLayoutDocxExporter
            ._create_table_cell_shape(
                shape_id=shape_id,
                left=cell.left,
                top=cell.top,
                width=width,
                height=height,
                vertical_alignment=vertical_alignment,
                fill_color=cell.fill_color,
            )
        )

        textbox = (
            FixedLayoutDocxExporter
            ._create_vml_element(
                "textbox"
            )
        )

        textbox.set(
            "inset",
            (
                f"{FixedLayoutDocxExporter.TABLE_CELL_HORIZONTAL_PADDING}pt,"
                f"{FixedLayoutDocxExporter.TABLE_CELL_VERTICAL_PADDING}pt,"
                f"{FixedLayoutDocxExporter.TABLE_CELL_HORIZONTAL_PADDING}pt,"
                f"{FixedLayoutDocxExporter.TABLE_CELL_VERTICAL_PADDING}pt"
            ),
        )

        textbox_content = OxmlElement(
            "w:txbxContent"
        )

        paragraph_element = OxmlElement(
            "w:p"
        )

        textbox_content.append(
            paragraph_element
        )

        textbox.append(
            textbox_content
        )

        shape.append(
            textbox
        )

        pict = OxmlElement(
            "w:pict"
        )

        pict.append(
            shape
        )

        anchor_run = anchor_paragraph.add_run()

        anchor_run._r.append(
            pict
        )

        text_paragraph = Paragraph(
            paragraph_element,
            anchor_paragraph._parent,
        )

        FixedLayoutDocxExporter._configure_table_cell_paragraph(
            text_paragraph=text_paragraph,
            cell=cell,
            cell_spans=cell_spans,
        )

        FixedLayoutDocxExporter._add_table_cell_content(
            text_paragraph=text_paragraph,
            cell=cell,
            cell_spans=cell_spans,
        )

    @staticmethod
    def _create_table_cell_shape(
        shape_id: int,
        left: float,
        top: float,
        width: float,
        height: float,
        vertical_alignment: str,
        fill_color: str | None,
    ):
        """
        Create one bordered VML rectangle representing
        a detected PDF table cell.
        """

        shape = (
            FixedLayoutDocxExporter
            ._create_vml_element(
                "shape"
            )
        )

        shape.set(
            "id",
            f"DocuMindTableCell{shape_id}",
        )

        shape.set(
            "type",
            "#_x0000_t202",
        )

        if fill_color is None:
            shape.set(
                "filled",
                "f",
            )
        else:
            shape.set(
                "filled",
                "t",
            )

            shape.set(
                "fillcolor",
                fill_color,
            )

        shape.set(
            "stroked",
            "f",
        )

        shape.set(
            (
                f"{{"
                f"{FixedLayoutDocxExporter.OFFICE_NAMESPACE}"
                f"}}allowincell"
            ),
            "f",
        )

        style = (
            "position:absolute;"
            f"margin-left:{left:.3f}pt;"
            f"margin-top:{top:.3f}pt;"
            f"width:{width:.3f}pt;"
            f"height:{height:.3f}pt;"
            f"z-index:{FixedLayoutDocxExporter.TABLE_Z_INDEX};"
            "mso-position-horizontal-relative:page;"
            "mso-position-vertical-relative:page;"
            "mso-wrap-style:none;"
            f"v-text-anchor:{vertical_alignment};"
        )

        shape.set(
            "style",
            style,
        )

        wrap = (
            FixedLayoutDocxExporter
            ._create_word_2003_element(
                "wrap"
            )
        )

        wrap.set(
            "type",
            "none",
        )

        shape.append(
            wrap
        )

        return shape

    @staticmethod
    def _configure_table_cell_paragraph(
        text_paragraph,
        cell,
        cell_spans: list,
    ) -> None:
        """
        Remove Word defaults and apply the alignment detected
        from the original PDF cell text.
        """

        paragraph_format = (
            text_paragraph.paragraph_format
        )

        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        paragraph_format.left_indent = Pt(0)
        paragraph_format.right_indent = Pt(0)
        paragraph_format.first_line_indent = Pt(0)

        paragraph_format.line_spacing_rule = (
            WD_LINE_SPACING.SINGLE
        )

        alignment = (
            FixedLayoutDocxExporter
            ._detect_table_cell_alignment(
                cell=cell,
                cell_spans=cell_spans,
            )
        )

        text_paragraph.alignment = alignment

    @staticmethod
    def _add_table_cell_content(
        text_paragraph,
        cell,
        cell_spans: list,
    ) -> None:
        """
        Render original PDF spans inside the detected cell.

        When no original spans are found, use the extracted
        plain cell text as a fallback.
        """

        if not cell_spans:
            FixedLayoutDocxExporter._add_fallback_cell_text(
                text_paragraph=text_paragraph,
                text=cell.text,
            )
            return

        span_lines = (
            FixedLayoutDocxExporter
            ._group_table_spans_into_lines(
                cell_spans
            )
        )

        for line_index, span_line in enumerate(
            span_lines
        ):
            for span in span_line:
                run = text_paragraph.add_run(
                    TextNormalizer.normalize(
                        span.text
                    )
                )

                FixedLayoutDocxExporter._apply_span_format(
                    run=run,
                    span=span,
                )

            if line_index < len(span_lines) - 1:
                break_run = text_paragraph.add_run()

                break_run.add_break(
                    WD_BREAK.LINE
                )

    @staticmethod
    def _add_fallback_cell_text(
        text_paragraph,
        text: str,
    ) -> None:
        """
        Render extracted table text when no original spans
        can be matched to the cell.
        """

        lines = text.splitlines() or [""]

        for line_index, line_text in enumerate(lines):

            normalized_line_text = (
                TextNormalizer.normalize(
                    line_text
                )
            )

            run = text_paragraph.add_run(
                normalized_line_text
            )

            run.font.size = Pt(
                FixedLayoutDocxExporter
                .DEFAULT_TABLE_FONT_SIZE
            )

            fallback_font = (
                FixedLayoutDocxExporter.BULLET_FONT_NAME
                if TextNormalizer.contains_bullet(
                    normalized_line_text
                )
                else FixedLayoutDocxExporter.DEFAULT_TABLE_FONT
            )
            
            FixedLayoutDocxExporter._apply_font_name(
                run=run,
                pdf_font_name=fallback_font,
            )

            if line_index < len(lines) - 1:
                run.add_break(
                    WD_BREAK.LINE
                )

    @staticmethod
    def _apply_span_format(
        run,
        span,
    ) -> None:
        """
        Apply original PDF typography to a Word run.

        Bullet glyphs use a reliable Unicode-capable font.
        """

        run.font.size = Pt(
            span.font_size
        )

        normalized_text = TextNormalizer.normalize(
            span.text
        )

        if TextNormalizer.contains_bullet(
            normalized_text
        ):
            FixedLayoutDocxExporter._apply_font_name(
                run=run,
                pdf_font_name=(
                    FixedLayoutDocxExporter
                    .BULLET_FONT_NAME
                ),
            )
        else:
            FixedLayoutDocxExporter._apply_font_name(
                run=run,
                pdf_font_name=span.font,
            )

        run.font.color.rgb = WordRGBColor(
            span.color.red,
            span.color.green,
            span.color.blue,
        )

        run.bold = (
            FixedLayoutDocxExporter
            ._is_bold(span.flags)
        )

        run.italic = (
            FixedLayoutDocxExporter
            ._is_italic(span.flags)
        )
    
    @staticmethod
    def _find_spans_for_cell(
        page,
        cell,
    ) -> list:
        """
        Return all visible PDF spans whose centers lie inside
        the detected table cell.

        Spans are sorted by vertical and horizontal position.
        """

        spans = []

        for block in page.blocks:
            for line in block.lines:
                for span in line.spans:

                    if not span.text.strip():
                        continue

                    center_x = (
                        span.left + span.right
                    ) / 2

                    center_y = (
                        span.top + span.bottom
                    ) / 2

                    if (
                        cell.left <= center_x <= cell.right
                        and cell.top <= center_y <= cell.bottom
                    ):
                        spans.append(span)

        spans.sort(
            key=lambda span: (
                span.origin_y,
                span.left,
            )
        )

        return spans

    @staticmethod
    def _group_table_spans_into_lines(
        spans: list,
    ) -> list[list]:
        """
        Group cell spans into visual PDF lines using their
        baseline coordinates.
        """

        if not spans:
            return []

        lines: list[list] = []

        current_line = [spans[0]]
        current_baseline = spans[0].origin_y

        for span in spans[1:]:

            if (
                abs(
                    span.origin_y - current_baseline
                )
                <= FixedLayoutDocxExporter
                .TABLE_LINE_TOLERANCE
            ):
                current_line.append(span)
                continue

            current_line.sort(
                key=lambda item: item.left
            )

            lines.append(
                current_line
            )

            current_line = [span]
            current_baseline = span.origin_y

        current_line.sort(
            key=lambda item: item.left
        )

        lines.append(
            current_line
        )

        return lines

    @staticmethod
    def _detect_table_cell_alignment(
        cell,
        cell_spans: list,
    ):
        """
        Detect horizontal table-cell alignment by analyzing
        each visual text line separately.

        The method intentionally defaults to left alignment.
        Center or right alignment is used only when most visible
        lines clearly support it.
        """

        if not cell_spans:
            return WD_ALIGN_PARAGRAPH.LEFT

        span_lines = (
            FixedLayoutDocxExporter
            ._group_table_spans_into_lines(
                cell_spans
            )
        )

        if not span_lines:
            return WD_ALIGN_PARAGRAPH.LEFT

        detected_alignments = [
            FixedLayoutDocxExporter
            ._detect_single_table_line_alignment(
                cell=cell,
                span_line=span_line,
            )
            for span_line in span_lines
        ]

        center_count = sum(
            alignment == WD_ALIGN_PARAGRAPH.CENTER
            for alignment in detected_alignments
        )

        right_count = sum(
            alignment == WD_ALIGN_PARAGRAPH.RIGHT
            for alignment in detected_alignments
        )

        total_lines = len(
            detected_alignments
        )

        required_ratio = (
            FixedLayoutDocxExporter
            .TABLE_ALIGNMENT_REQUIRED_RATIO
        )

        if (
            center_count / total_lines
            >= required_ratio
        ):
            return WD_ALIGN_PARAGRAPH.CENTER

        if (
            right_count / total_lines
            >= required_ratio
        ):
            return WD_ALIGN_PARAGRAPH.RIGHT

        return WD_ALIGN_PARAGRAPH.LEFT
    
    @staticmethod
    def _detect_single_table_line_alignment(
        cell,
        span_line: list,
    ):
        """
        Detect alignment for one visual line inside a table cell.

        Left alignment is preferred unless the line clearly
        appears centered or right-aligned.
        """

        if not span_line:
            return WD_ALIGN_PARAGRAPH.LEFT

        text_left = min(
            span.left
            for span in span_line
        )

        text_right = max(
            span.right
            for span in span_line
        )

        cell_width = max(
            cell.right - cell.left,
            1.0,
        )

        text_width = max(
            text_right - text_left,
            0.0,
        )

        left_gap = max(
            text_left - cell.left,
            0.0,
        )

        right_gap = max(
            cell.right - text_right,
            0.0,
        )

        tolerance = (
            FixedLayoutDocxExporter
            .TABLE_ALIGNMENT_TOLERANCE
        )

        minimum_center_margin = (
            FixedLayoutDocxExporter
            .TABLE_CENTER_MIN_MARGIN
        )

        # A line is considered centered only when both sides have
        # meaningful whitespace and those margins are nearly equal.
        is_centered = (
            left_gap >= minimum_center_margin
            and right_gap >= minimum_center_margin
            and abs(left_gap - right_gap) <= tolerance
        )

        if is_centered:
            return WD_ALIGN_PARAGRAPH.CENTER

        # Right alignment requires the text to sit substantially
        # closer to the right boundary than the left boundary.
        if left_gap > 0:
            right_alignment_ratio = (
                left_gap
                / max(
                    left_gap + right_gap,
                    1.0,
                )
            )
        else:
            right_alignment_ratio = 0.0

        is_right_aligned = (
            right_gap <= tolerance
            and right_alignment_ratio
            >= FixedLayoutDocxExporter
            .TABLE_RIGHT_ALIGNMENT_RATIO
        )

        if is_right_aligned:
            return WD_ALIGN_PARAGRAPH.RIGHT

        # Wide text usually represents normal left-aligned content,
        # even when its outer bounds appear visually balanced.
        if text_width >= cell_width * 0.60:
            return WD_ALIGN_PARAGRAPH.LEFT

        return WD_ALIGN_PARAGRAPH.LEFT
    
    @staticmethod
    def _detect_table_cell_vertical_alignment(
        cell,
        cell_spans: list,
    ) -> str:
        """
        Detect whether cell text is aligned to the top,
        vertical center, or bottom.

        The decision is based on the free PDF space above and
        below the visible text.
        """

        if not cell_spans:
            return (
                FixedLayoutDocxExporter
                .TABLE_VERTICAL_TOP
            )

        text_top = min(
            span.top
            for span in cell_spans
        )

        text_bottom = max(
            span.bottom
            for span in cell_spans
        )

        top_gap = max(
            text_top - cell.top,
            0.0,
        )

        bottom_gap = max(
            cell.bottom - text_bottom,
            0.0,
        )

        tolerance = (
            FixedLayoutDocxExporter
            .TABLE_VERTICAL_ALIGNMENT_TOLERANCE
        )

        if abs(top_gap - bottom_gap) <= tolerance:
            return (
                FixedLayoutDocxExporter
                .TABLE_VERTICAL_CENTER
            )

        if top_gap > bottom_gap + tolerance:
            return (
                FixedLayoutDocxExporter
                .TABLE_VERTICAL_BOTTOM
            )

        return (
            FixedLayoutDocxExporter
            .TABLE_VERTICAL_TOP
        )

    @staticmethod
    def _span_is_inside_any_table(
        span,
        tables,
    ) -> bool:
        """
        Return True when the center of a PDF span lies inside
        any detected table.

        These spans are skipped from normal text rendering
        because the table renderer recreates their content.
        """

        center_x = (
            span.left + span.right
        ) / 2

        center_y = (
            span.top + span.bottom
        ) / 2

        for table in tables:
            if (
                table.left <= center_x <= table.right
                and table.top <= center_y <= table.bottom
            ):
                return True

        return False

    @staticmethod
    def _add_image_shape(
        document_part,
        anchor_paragraph,
        image,
    ) -> None:
        """
        Add one extracted PDF image as an absolutely positioned
        Word VML image shape.
        """

        relationship_id, _ = (
            document_part.get_or_add_image(
                BytesIO(image.image_bytes)
            )
        )

        width = max(
            image.displayed_width,
            1.0,
        )

        height = max(
            image.displayed_height,
            1.0,
        )

        shape_id = (
            FixedLayoutDocxExporter._next_shape_id()
        )

        shape = (
            FixedLayoutDocxExporter
            ._create_image_shape(
                shape_id=shape_id,
                left=image.left,
                top=image.top,
                width=width,
                height=height,
            )
        )

        image_data = (
            FixedLayoutDocxExporter
            ._create_vml_element(
                "imagedata"
            )
        )

        image_data.set(
            (
                f"{{"
                f"{FixedLayoutDocxExporter.RELATIONSHIPS_NAMESPACE}"
                f"}}id"
            ),
            relationship_id,
        )

        image_data.set(
            (
                f"{{"
                f"{FixedLayoutDocxExporter.OFFICE_NAMESPACE}"
                f"}}title"
            ),
            "",
        )

        shape.append(
            image_data
        )

        pict = OxmlElement(
            "w:pict"
        )

        pict.append(
            shape
        )

        anchor_run = (
            anchor_paragraph.add_run()
        )

        anchor_run._r.append(
            pict
        )

    @staticmethod
    def _create_image_shape(
        shape_id: int,
        left: float,
        top: float,
        width: float,
        height: float,
    ):
        """
        Create an absolutely positioned VML image shape.
        """
    
        shape = (
            FixedLayoutDocxExporter
            ._create_vml_element(
                "shape"
            )
        )
    
        shape.set(
            "id",
            f"DocuMindImage{shape_id}",
        )
    
        shape.set(
            "type",
            "#_x0000_t75",
        )
    
        shape.set(
            "stroked",
            "f",
        )
    
        shape.set(
            (
                f"{{"
                f"{FixedLayoutDocxExporter.OFFICE_NAMESPACE}"
                f"}}allowincell"
            ),
            "f",
        )
    
        style = (
            "position:absolute;"
            f"margin-left:{left:.3f}pt;"
            f"margin-top:{top:.3f}pt;"
            f"width:{width:.3f}pt;"
            f"height:{height:.3f}pt;"
            f"z-index:{FixedLayoutDocxExporter.IMAGE_Z_INDEX};"
            "mso-position-horizontal-relative:page;"
            "mso-position-vertical-relative:page;"
            "mso-wrap-style:none;"
        )
    
        shape.set(
            "style",
            style,
        )
    
        wrap = (
            FixedLayoutDocxExporter
            ._create_word_2003_element(
                "wrap"
            )
        )
    
        wrap.set(
            "type",
            "none",
        )
    
        shape.append(
            wrap
        )
    
        return shape

    @staticmethod
    def _add_span_textbox(
        anchor_paragraph,
        span,
    ) -> None:
        """
        Create one absolutely positioned Word text box for
        one formatted PDF span.
        """

        left = span.left

        top = max(
            span.top
            - FixedLayoutDocxExporter.TEXTBOX_TOP_ADJUSTMENT,
            0.0,
        )

        width = max(
            span.right
            - span.left
            + FixedLayoutDocxExporter.TEXTBOX_WIDTH_PADDING,
            1.0,
        )

        pdf_span_height = max(
            span.bottom - span.top,
            1.0,
        )

        word_line_height = max(
            span.font_size
            * FixedLayoutDocxExporter.LINE_HEIGHT_FACTOR,
            pdf_span_height,
        )

        height = (
            word_line_height
            + FixedLayoutDocxExporter.TEXTBOX_HEIGHT_PADDING
        )

        shape_id = (
            FixedLayoutDocxExporter
            ._next_shape_id()
        )

        shape = (
            FixedLayoutDocxExporter
            ._create_textbox_shape(
                shape_id=shape_id,
                left=left,
                top=top,
                width=width,
                height=height,
            )
        )

        textbox = FixedLayoutDocxExporter._create_vml_element(
            "textbox"
        )

        textbox.set(
            "inset",
            "0,0,0,0",
        )

        textbox_content = OxmlElement(
            "w:txbxContent"
        )

        text_paragraph_element = OxmlElement(
            "w:p"
        )

        textbox_content.append(
            text_paragraph_element
        )

        textbox.append(
            textbox_content
        )

        shape.append(
            textbox
        )

        pict = OxmlElement(
            "w:pict"
        )

        pict.append(
            shape
        )

        anchor_run = anchor_paragraph.add_run()

        anchor_run._r.append(
            pict
        )

        text_paragraph = Paragraph(
            text_paragraph_element,
            anchor_paragraph._parent,
        )

        FixedLayoutDocxExporter._configure_text_paragraph(
            text_paragraph=text_paragraph,
            span=span,
        )

        FixedLayoutDocxExporter._add_formatted_run(
            text_paragraph=text_paragraph,
            span=span,
        )

    @staticmethod
    def _create_textbox_shape(
        shape_id: int,
        left: float,
        top: float,
        width: float,
        height: float,
    ):
        """
        Create an absolutely positioned VML text-box shape.

        VML elements are created using explicit namespace URLs
        because python-docx does not register the `v`, `o`, and
        `w10` prefixes in OxmlElement.
        """

        shape = FixedLayoutDocxExporter._create_vml_element(
            "shape"
        )

        shape.set(
            "id",
            f"DocuMindTextBox{shape_id}",
        )

        shape.set(
            "type",
            "#_x0000_t202",
        )

        shape.set(
            "stroked",
            "f",
        )

        shape.set(
            "filled",
            "f",
        )

        shape.set(
            (
                f"{{{FixedLayoutDocxExporter.OFFICE_NAMESPACE}}}"
                "allowincell"
            ),
            "f",
        )

        style = (
            "position:absolute;"
            f"margin-left:{left:.3f}pt;"
            f"margin-top:{top:.3f}pt;"
            f"width:{width:.3f}pt;"
            f"height:{height:.3f}pt;"
            f"z-index:{FixedLayoutDocxExporter.TEXT_Z_INDEX};"
            "mso-position-horizontal-relative:page;"
            "mso-position-vertical-relative:page;"
            "mso-wrap-style:none;"
            "mso-fit-shape-to-text:t;"
        )

        shape.set(
            "style",
            style,
        )

        wrap = (
            FixedLayoutDocxExporter
            ._create_word_2003_element(
                "wrap"
            )
        )

        wrap.set(
            "type",
            "none",
        )

        shape.append(
            wrap
        )

        return shape

    @staticmethod
    def _create_vml_element(
        local_name: str,
    ):
        """
        Create a VML element with all namespaces required by
        Microsoft Word's legacy text-box format.
        """

        return etree.Element(
            (
                f"{{{FixedLayoutDocxExporter.VML_NAMESPACE}}}"
                f"{local_name}"
            ),
            nsmap=FixedLayoutDocxExporter.VML_NAMESPACE_MAP,
        )

    @staticmethod
    def _create_word_2003_element(
        local_name: str,
    ):
        """
        Create an element from the legacy Microsoft Word 2003
        namespace, such as w10:wrap.
        """

        return etree.Element(
            (
                f"{{"
                f"{FixedLayoutDocxExporter.WORD_2003_NAMESPACE}"
                f"}}{local_name}"
            )
        )

    @staticmethod
    def _configure_text_paragraph(
        text_paragraph,
        span,
    ) -> None:
        """
        Remove Word paragraph spacing inside the text box.
        """

        paragraph_format = (
            text_paragraph.paragraph_format
        )

        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        paragraph_format.left_indent = Pt(0)
        paragraph_format.right_indent = Pt(0)
        paragraph_format.first_line_indent = Pt(0)

        paragraph_format.line_spacing_rule = (
            WD_LINE_SPACING.EXACTLY
        )

        paragraph_format.line_spacing = Pt(
            max(
                span.font_size
                * FixedLayoutDocxExporter.LINE_HEIGHT_FACTOR,
                span.bottom - span.top,
            )
        )

    @staticmethod
    def _add_formatted_run(
        text_paragraph,
        span,
    ) -> None:
        """
        Add editable text while preserving typography.

        PDF-specific bullet glyphs are normalized into reliable
        Unicode characters before being written to Word.
        """

        normalized_text = TextNormalizer.normalize(
            span.text
        )

        run = text_paragraph.add_run(
            normalized_text
        )

        font = run.font

        font.size = Pt(
            span.font_size
        )

        if TextNormalizer.contains_bullet(
            normalized_text
        ):
            FixedLayoutDocxExporter._apply_font_name(
                run=run,
                pdf_font_name=(
                    FixedLayoutDocxExporter
                    .BULLET_FONT_NAME
                ),
            )

            font.size = Pt(
                span.font_size
                * FixedLayoutDocxExporter
                .BULLET_SIZE_FACTOR
            )
        else:
            FixedLayoutDocxExporter._apply_font_name(
                run=run,
                pdf_font_name=span.font,
            )

        font.color.rgb = WordRGBColor(
            span.color.red,
            span.color.green,
            span.color.blue,
        )

        run.bold = (
            FixedLayoutDocxExporter
            ._is_bold(span.flags)
        )

        run.italic = (
            FixedLayoutDocxExporter
            ._is_italic(span.flags)
        )

    @staticmethod
    def _apply_font_name(
        run,
        pdf_font_name: str,
    ) -> None:
        """
        Resolve and apply the PDF font family to every
        Word font slot.
        """

        word_font_name = FontNameResolver.resolve(
            pdf_font_name
        )

        run.font.name = word_font_name

        run_properties = (
            run._element.get_or_add_rPr()
        )

        font_properties = (
            run_properties.get_or_add_rFonts()
        )

        font_properties.set(
            qn("w:ascii"),
            word_font_name,
        )

        font_properties.set(
            qn("w:hAnsi"),
            word_font_name,
        )

        font_properties.set(
            qn("w:eastAsia"),
            word_font_name,
        )

        font_properties.set(
            qn("w:cs"),
            word_font_name,
        )

    @staticmethod
    def _is_bold(flags: int) -> bool:
        """
        PyMuPDF flag bit 4 represents bold text.
        """

        return bool(
            flags & (1 << 4)
        )

    @staticmethod
    def _is_italic(flags: int) -> bool:
        """
        PyMuPDF flag bit 1 represents italic text.
        """

        return bool(
            flags & (1 << 1)
        )

    @staticmethod
    def _next_shape_id() -> int:
        """
        Return a unique ID for every positioned text box.
        """

        shape_id = (
            FixedLayoutDocxExporter._shape_id
        )

        FixedLayoutDocxExporter._shape_id += 1

        return shape_id