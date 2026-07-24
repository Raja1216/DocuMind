from __future__ import annotations

from io import BytesIO
import math

import fitz

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.models.geometry.rectangle import (
    Rectangle,
)


class EditableTableRegionFallbackRenderer:
    """
    Render one source-PDF table region as an inline image.

    The crop is derived only from the detected table rectangle,
    source page geometry and normalized border widths. No filename,
    page number, text, row count or sample-specific rule is used.
    """

    DEFAULT_DPI = 200

    MINIMUM_DPI = 96

    MAXIMUM_DPI = 400

    DEFAULT_SAFETY_MARGIN = 1.0

    BORDER_MARGIN_EXTRA = 0.75

    MAXIMUM_SAFETY_MARGIN = 6.0

    MINIMUM_REGION_SIZE = 0.5

    MINIMUM_WORD_WIDTH = 6.0

    @classmethod
    def render(
        cls,
        *,
        container,
        page,
        table,
        available_width: float,
        dpi: int | None = None,
    ):
        """
        Add an inline cropped region and return metadata.

        Returns:
            (inline_shape, crop_rectangle,
             rendered_width_points, rendered_height_points)
        """

        source_page = getattr(
            page,
            "source_pdf_page",
            None,
        )

        if source_page is None:
            raise ValueError(
                "The source PDF page is unavailable for table-region fallback."
            )

        source_bbox = cls._resolve_table_bbox(
            table
        )

        margin = cls._resolve_safety_margin(
            table
        )

        page_rectangle = cls._resolve_page_rectangle(
            source_page
        )

        crop_rectangle = cls._expand_and_clip(
            rectangle=source_bbox,
            page_rectangle=page_rectangle,
            margin=margin,
        )

        if (
            crop_rectangle.width
            < cls.MINIMUM_REGION_SIZE
            or crop_rectangle.height
            < cls.MINIMUM_REGION_SIZE
        ):
            raise ValueError(
                "The table fallback crop has no usable area."
            )

        resolved_dpi = cls._normalize_dpi(
            dpi
        )

        scale = (
            float(resolved_dpi)
            / 72.0
        )

        pixmap = source_page.get_pixmap(
            matrix=fitz.Matrix(
                scale,
                scale,
            ),
            clip=fitz.Rect(
                crop_rectangle.left,
                crop_rectangle.top,
                crop_rectangle.right,
                crop_rectangle.bottom,
            ),
            alpha=False,
            annots=True,
        )

        image_bytes = pixmap.tobytes(
            "png"
        )

        if not image_bytes:
            raise ValueError(
                "The source page produced an empty table fallback image."
            )

        target_width = min(
            max(
                float(
                    crop_rectangle.width
                ),
                cls.MINIMUM_WORD_WIDTH,
            ),
            max(
                float(
                    available_width
                ),
                cls.MINIMUM_WORD_WIDTH,
            ),
        )

        aspect_ratio = (
            float(
                crop_rectangle.height
            )
            / max(
                float(
                    crop_rectangle.width
                ),
                cls.MINIMUM_REGION_SIZE,
            )
        )

        target_height = (
            target_width
            * aspect_ratio
        )

        paragraph = container.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)

        run = paragraph.add_run()

        inline_shape = run.add_picture(
            BytesIO(
                image_bytes
            ),
            width=Pt(
                target_width
            ),
        )

        return (
            inline_shape,
            crop_rectangle,
            target_width,
            target_height,
        )

    @classmethod
    def _resolve_table_bbox(
        cls,
        table,
    ) -> Rectangle:
        bbox = getattr(
            table,
            "bbox",
            None,
        )

        if bbox is not None:
            geometry_source = bbox
        else:
            source_table = getattr(
                table,
                "source_table",
                None,
            )

            geometry_source = (
                source_table
                if source_table is not None
                else table
            )

        try:
            left = float(
                getattr(
                    geometry_source,
                    "left"
                )
            )
            top = float(
                getattr(
                    geometry_source,
                    "top"
                )
            )
            right = float(
                getattr(
                    geometry_source,
                    "right"
                )
            )
            bottom = float(
                getattr(
                    geometry_source,
                    "bottom"
                )
            )

        except (
            AttributeError,
            TypeError,
            ValueError,
        ) as error:
            raise ValueError(
                "The table does not expose a valid fallback rectangle."
            ) from error

        if not all(
            math.isfinite(
                value
            )
            for value in (
                left,
                top,
                right,
                bottom,
            )
        ):
            raise ValueError(
                "The table fallback rectangle contains non-finite coordinates."
            )

        if right < left:
            left, right = right, left

        if bottom < top:
            top, bottom = bottom, top

        return Rectangle(
            left=left,
            top=top,
            right=right,
            bottom=bottom,
        )

    @staticmethod
    def _resolve_page_rectangle(
        source_page,
    ) -> Rectangle:
        source_rectangle = source_page.rect

        return Rectangle(
            left=float(
                source_rectangle.x0
            ),
            top=float(
                source_rectangle.y0
            ),
            right=float(
                source_rectangle.x1
            ),
            bottom=float(
                source_rectangle.y1
            ),
        )

    @classmethod
    def _resolve_safety_margin(
        cls,
        table,
    ) -> float:
        border_widths: list[float] = []

        for cell in getattr(
            table,
            "cells",
            [],
        ) or []:
            borders = getattr(
                cell,
                "borders",
                None,
            )

            if borders is None:
                continue

            for edge_name in (
                "top",
                "right",
                "bottom",
                "left",
            ):
                border = getattr(
                    borders,
                    edge_name,
                    None,
                )

                try:
                    width = float(
                        getattr(
                            border,
                            "width",
                            0.0,
                        )
                    )
                except (
                    TypeError,
                    ValueError,
                ):
                    continue

                if math.isfinite(width):
                    border_widths.append(
                        max(
                            width,
                            0.0,
                        )
                    )

        maximum_border_width = max(
            border_widths,
            default=0.0,
        )

        return min(
            max(
                cls.DEFAULT_SAFETY_MARGIN,
                maximum_border_width
                + cls.BORDER_MARGIN_EXTRA,
            ),
            cls.MAXIMUM_SAFETY_MARGIN,
        )

    @staticmethod
    def _expand_and_clip(
        *,
        rectangle: Rectangle,
        page_rectangle: Rectangle,
        margin: float,
    ) -> Rectangle:
        return Rectangle(
            left=max(
                float(
                    page_rectangle.left
                ),
                float(
                    rectangle.left
                )
                - float(margin),
            ),
            top=max(
                float(
                    page_rectangle.top
                ),
                float(
                    rectangle.top
                )
                - float(margin),
            ),
            right=min(
                float(
                    page_rectangle.right
                ),
                float(
                    rectangle.right
                )
                + float(margin),
            ),
            bottom=min(
                float(
                    page_rectangle.bottom
                ),
                float(
                    rectangle.bottom
                )
                + float(margin),
            ),
        )

    @classmethod
    def _normalize_dpi(
        cls,
        dpi: int | None,
    ) -> int:
        try:
            resolved = int(
                dpi
                if dpi is not None
                else cls.DEFAULT_DPI
            )
        except (
            TypeError,
            ValueError,
        ):
            resolved = cls.DEFAULT_DPI

        return max(
            cls.MINIMUM_DPI,
            min(
                resolved,
                cls.MAXIMUM_DPI,
            ),
        )
