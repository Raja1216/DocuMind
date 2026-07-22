from __future__ import annotations

import math
from typing import Any

from docx.dml.color import RGBColor as WordRGBColor
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.exporter.font_name_resolver import (
    FontNameResolver,
)
from src.models.editable_table import (
    EditableBorderLineStyle,
    EditableCellHorizontalAlignment,
    EditableCellVerticalAlignment,
    EditableTable,
    EditableTableBorder,
    EditableTableCell,
    EditableTableCellParagraph,
    EditableTableDisposition,
)


class EditableWordTableRenderer:
    """
    Render an EditableTable as a native, editable Word table.

    This renderer is deliberately independent from DocxExporter.
    Step 63G.1.8 will connect it to the unified page render plan.

    It handles:
        - fixed table and column widths;
        - row heights and repeated header rows;
        - rectangular horizontal/vertical merged cells;
        - formatted runs and fallback cell text;
        - cell alignment, padding, fill and individual borders.
    """

    DEFAULT_AVAILABLE_WIDTH = 468.0
    MINIMUM_TABLE_WIDTH = 36.0
    MINIMUM_COLUMN_WIDTH = 6.0
    MINIMUM_ROW_HEIGHT = 1.0

    DEFAULT_FONT_NAME = "Arial"
    DEFAULT_FONT_SIZE = 10.0

    @classmethod
    def render(
        cls,
        *,
        container,
        table: EditableTable,
        available_width: float | None = None,
    ):
        """
        Add and return one native python-docx Table.

        ``container`` may be a Word Document, a table cell, or any
        python-docx object exposing ``add_table(rows, cols)``.
        """

        cls._validate_renderable_table(
            table
        )

        resolved_available_width = max(
            float(
                available_width
                if available_width is not None
                else cls.DEFAULT_AVAILABLE_WIDTH
            ),
            cls.MINIMUM_TABLE_WIDTH,
        )

        column_widths = (
            cls._resolve_column_widths(
                table=table,
                available_width=(
                    resolved_available_width
                ),
            )
        )

        word_table = container.add_table(
            rows=table.row_count,
            cols=table.column_count,
        )

        word_table.alignment = (
            WD_TABLE_ALIGNMENT.LEFT
        )

        word_table.autofit = False

        cls._set_fixed_table_layout(
            word_table
        )

        cls._set_table_width(
            word_table=word_table,
            width=sum(
                column_widths
            ),
        )

        cls._apply_column_widths(
            word_table=word_table,
            column_widths=column_widths,
        )

        cls._apply_rows(
            word_table=word_table,
            table=table,
        )

        merged_word_cells = (
            cls._apply_merges(
                word_table=word_table,
                table=table,
            )
        )

        for editable_cell in sorted(
            table.cells,
            key=lambda cell: (
                cell.row_index,
                cell.column_index,
            ),
        ):
            word_cell = (
                merged_word_cells.get(
                    id(
                        editable_cell
                    )
                )
                or word_table.cell(
                    editable_cell.row_index,
                    editable_cell.column_index,
                )
            )

            cls._apply_cell_style(
                word_cell=word_cell,
                editable_cell=editable_cell,
            )

            cls._render_cell_content(
                word_cell=word_cell,
                editable_cell=editable_cell,
            )

        return word_table

    # ---------------------------------------------------------
    # Validation and geometry
    # ---------------------------------------------------------

    @staticmethod
    def _validate_renderable_table(
        table: EditableTable,
    ) -> None:
        if (
            table.disposition
            != EditableTableDisposition.EDITABLE
        ):
            raise ValueError(
                (
                    "Only editable table models can be rendered "
                    "as native Word tables."
                )
            )

        structure_errors = (
            table.validate_structure()
        )

        if structure_errors:
            raise ValueError(
                (
                    "Cannot render a structurally invalid table: "
                    + "; ".join(
                        structure_errors
                    )
                )
            )

        if (
            len(table.rows)
            != table.row_count
        ):
            raise ValueError(
                (
                    "Cannot render a table without complete "
                    "row definitions."
                )
            )

        if (
            len(table.columns)
            != table.column_count
        ):
            raise ValueError(
                (
                    "Cannot render a table without complete "
                    "column definitions."
                )
            )

    @classmethod
    def _resolve_column_widths(
        cls,
        *,
        table: EditableTable,
        available_width: float,
    ) -> list[float]:
        source_widths = [
            max(
                float(
                    column.width
                ),
                cls.MINIMUM_COLUMN_WIDTH,
            )
            for column in sorted(
                table.columns,
                key=lambda column: (
                    column.column_index
                ),
            )
        ]

        source_total = sum(
            source_widths
        )

        if source_total <= 0.0:
            return [
                available_width
                / max(
                    table.column_count,
                    1,
                )
                for _ in range(
                    table.column_count
                )
            ]

        target_width = min(
            max(
                source_total,
                cls.MINIMUM_TABLE_WIDTH,
            ),
            available_width,
        )

        scale = (
            target_width
            / source_total
        )

        widths = [
            max(
                width * scale,
                cls.MINIMUM_COLUMN_WIDTH,
            )
            for width in source_widths
        ]

        # Re-normalize after enforcing the minimum column width.
        normalized_total = sum(
            widths
        )

        if (
            normalized_total
            > available_width
            and normalized_total > 0.0
        ):
            secondary_scale = (
                available_width
                / normalized_total
            )

            widths = [
                max(
                    width
                    * secondary_scale,
                    1.0,
                )
                for width in widths
            ]

        return widths

    @staticmethod
    def _set_fixed_table_layout(
        word_table,
    ) -> None:
        table_properties = (
            word_table._tbl.tblPr
        )

        existing = (
            table_properties
            .first_child_found_in(
                "w:tblLayout"
            )
        )

        if existing is None:
            existing = OxmlElement(
                "w:tblLayout"
            )

            table_properties.append(
                existing
            )

        existing.set(
            qn(
                "w:type"
            ),
            "fixed",
        )

    @staticmethod
    def _set_table_width(
        *,
        word_table,
        width: float,
    ) -> None:
        table_properties = (
            word_table._tbl.tblPr
        )

        table_width = (
            table_properties
            .first_child_found_in(
                "w:tblW"
            )
        )

        if table_width is None:
            table_width = OxmlElement(
                "w:tblW"
            )

            table_properties.append(
                table_width
            )

        table_width.set(
            qn(
                "w:type"
            ),
            "dxa",
        )

        table_width.set(
            qn(
                "w:w"
            ),
            str(
                cls_points_to_twips(
                    width
                )
            ),
        )

    @staticmethod
    def _apply_column_widths(
        *,
        word_table,
        column_widths: list[float],
    ) -> None:
        for column_index, width in enumerate(
            column_widths
        ):
            word_table.columns[
                column_index
            ].width = Pt(
                width
            )

            for row in word_table.rows:
                word_cell = row.cells[
                    column_index
                ]

                word_cell.width = Pt(
                    width
                )

                EditableWordTableRenderer._set_cell_width(
                    word_cell=word_cell,
                    width=width,
                )

    @staticmethod
    def _set_cell_width(
        *,
        word_cell,
        width: float,
    ) -> None:
        cell_properties = (
            word_cell._tc
            .get_or_add_tcPr()
        )

        cell_width = (
            cell_properties
            .first_child_found_in(
                "w:tcW"
            )
        )

        if cell_width is None:
            cell_width = OxmlElement(
                "w:tcW"
            )

            cell_properties.append(
                cell_width
            )

        cell_width.set(
            qn(
                "w:type"
            ),
            "dxa",
        )

        cell_width.set(
            qn(
                "w:w"
            ),
            str(
                cls_points_to_twips(
                    width
                )
            ),
        )

    # ---------------------------------------------------------
    # Rows and merged cells
    # ---------------------------------------------------------

    @classmethod
    def _apply_rows(
        cls,
        *,
        word_table,
        table: EditableTable,
    ) -> None:
        rows_by_index = {
            row.row_index: row
            for row in table.rows
        }

        for row_index, word_row in enumerate(
            word_table.rows
        ):
            editable_row = (
                rows_by_index[
                    row_index
                ]
            )

            word_row.height = Pt(
                max(
                    float(
                        editable_row.height
                    ),
                    cls.MINIMUM_ROW_HEIGHT,
                )
            )

            word_row.height_rule = (
                WD_ROW_HEIGHT_RULE
                .AT_LEAST
            )

            cls._prevent_row_split(
                word_row
            )

            if editable_row.is_header:
                cls._repeat_row_as_header(
                    word_row
                )

    @staticmethod
    def _prevent_row_split(
        word_row,
    ) -> None:
        row_properties = (
            word_row._tr
            .get_or_add_trPr()
        )

        existing = (
            row_properties.find(
                qn(
                    "w:cantSplit"
                )
            )
        )

        if existing is None:
            row_properties.append(
                OxmlElement(
                    "w:cantSplit"
                )
            )

    @staticmethod
    def _repeat_row_as_header(
        word_row,
    ) -> None:
        row_properties = (
            word_row._tr
            .get_or_add_trPr()
        )

        existing = (
            row_properties.find(
                qn(
                    "w:tblHeader"
                )
            )
        )

        if existing is None:
            existing = OxmlElement(
                "w:tblHeader"
            )

            row_properties.append(
                existing
            )

        existing.set(
            qn(
                "w:val"
            ),
            "true",
        )

    @staticmethod
    def _apply_merges(
        *,
        word_table,
        table: EditableTable,
    ) -> dict[int, Any]:
        merged_word_cells: dict[
            int,
            Any,
        ] = {}

        for editable_cell in sorted(
            table.cells,
            key=lambda cell: (
                cell.row_index,
                cell.column_index,
            ),
        ):
            start_cell = (
                word_table.cell(
                    editable_cell.row_index,
                    editable_cell.column_index,
                )
            )

            if not editable_cell.is_merged:
                merged_word_cells[
                    id(
                        editable_cell
                    )
                ] = start_cell

                continue

            end_cell = (
                word_table.cell(
                    editable_cell.row_index
                    + editable_cell.row_span
                    - 1,
                    editable_cell.column_index
                    + editable_cell.column_span
                    - 1,
                )
            )

            merged_word_cells[
                id(
                    editable_cell
                )
            ] = start_cell.merge(
                end_cell
            )

        return merged_word_cells

    # ---------------------------------------------------------
    # Cell styling
    # ---------------------------------------------------------

    @classmethod
    def _apply_cell_style(
        cls,
        *,
        word_cell,
        editable_cell: EditableTableCell,
    ) -> None:
        cls._apply_vertical_alignment(
            word_cell=word_cell,
            alignment=(
                editable_cell
                .vertical_alignment
            ),
        )

        cls._apply_cell_margins(
            word_cell=word_cell,
            editable_cell=editable_cell,
        )

        cls._apply_cell_fill(
            word_cell=word_cell,
            fill_color=(
                editable_cell
                .fill_color
            ),
        )

        cls._apply_cell_borders(
            word_cell=word_cell,
            editable_cell=editable_cell,
        )

    @staticmethod
    def _apply_vertical_alignment(
        *,
        word_cell,
        alignment: EditableCellVerticalAlignment,
    ) -> None:
        mapping = {
            EditableCellVerticalAlignment.TOP: (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            ),
            EditableCellVerticalAlignment.CENTER: (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            ),
            EditableCellVerticalAlignment.BOTTOM: (
                WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            ),
            EditableCellVerticalAlignment.UNKNOWN: (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            ),
        }

        word_cell.vertical_alignment = (
            mapping.get(
                alignment,
                WD_CELL_VERTICAL_ALIGNMENT.TOP,
            )
        )

    @staticmethod
    def _apply_cell_margins(
        *,
        word_cell,
        editable_cell: EditableTableCell,
    ) -> None:
        cell_properties = (
            word_cell._tc
            .get_or_add_tcPr()
        )

        margins = (
            cell_properties
            .first_child_found_in(
                "w:tcMar"
            )
        )

        if margins is None:
            margins = OxmlElement(
                "w:tcMar"
            )

            cell_properties.append(
                margins
            )

        values = {
            "top": (
                editable_cell.padding.top
            ),
            "right": (
                editable_cell.padding.right
            ),
            "bottom": (
                editable_cell.padding.bottom
            ),
            "left": (
                editable_cell.padding.left
            ),
        }

        for edge_name, point_value in (
            values.items()
        ):
            existing = margins.find(
                qn(
                    f"w:{edge_name}"
                )
            )

            if existing is None:
                existing = OxmlElement(
                    f"w:{edge_name}"
                )

                margins.append(
                    existing
                )

            existing.set(
                qn(
                    "w:w"
                ),
                str(
                    cls_points_to_twips(
                        point_value
                    )
                ),
            )

            existing.set(
                qn(
                    "w:type"
                ),
                "dxa",
            )

    @staticmethod
    def _apply_cell_fill(
        *,
        word_cell,
        fill_color: str | None,
    ) -> None:
        if not fill_color:
            return

        cell_properties = (
            word_cell._tc
            .get_or_add_tcPr()
        )

        shading = (
            cell_properties
            .first_child_found_in(
                "w:shd"
            )
        )

        if shading is None:
            shading = OxmlElement(
                "w:shd"
            )

            cell_properties.append(
                shading
            )

        shading.set(
            qn(
                "w:val"
            ),
            "clear",
        )

        shading.set(
            qn(
                "w:color"
            ),
            "auto",
        )

        shading.set(
            qn(
                "w:fill"
            ),
            str(
                fill_color
            ).lstrip(
                "#"
            ).upper(),
        )

    @classmethod
    def _apply_cell_borders(
        cls,
        *,
        word_cell,
        editable_cell: EditableTableCell,
    ) -> None:
        cell_properties = (
            word_cell._tc
            .get_or_add_tcPr()
        )

        borders = (
            cell_properties
            .first_child_found_in(
                "w:tcBorders"
            )
        )

        if borders is None:
            borders = OxmlElement(
                "w:tcBorders"
            )

            cell_properties.append(
                borders
            )

        border_values = {
            "top": (
                editable_cell.borders.top
            ),
            "right": (
                editable_cell.borders.right
            ),
            "bottom": (
                editable_cell.borders.bottom
            ),
            "left": (
                editable_cell.borders.left
            ),
        }

        for edge_name, border in (
            border_values.items()
        ):
            existing = borders.find(
                qn(
                    f"w:{edge_name}"
                )
            )

            if existing is None:
                existing = OxmlElement(
                    f"w:{edge_name}"
                )

                borders.append(
                    existing
                )

            cls._write_border(
                element=existing,
                border=border,
            )

    @staticmethod
    def _write_border(
        *,
        element,
        border: EditableTableBorder,
    ) -> None:
        style_mapping = {
            EditableBorderLineStyle.NONE: (
                "nil"
            ),
            EditableBorderLineStyle.SINGLE: (
                "single"
            ),
            EditableBorderLineStyle.DOUBLE: (
                "double"
            ),
            EditableBorderLineStyle.DASHED: (
                "dashed"
            ),
            EditableBorderLineStyle.DOTTED: (
                "dotted"
            ),
        }

        word_style = style_mapping.get(
            border.style,
            "single",
        )

        element.set(
            qn(
                "w:val"
            ),
            word_style,
        )

        element.set(
            qn(
                "w:color"
            ),
            str(
                border.color
                or "000000"
            ).lstrip(
                "#"
            ).upper(),
        )

        border_size = (
            0
            if word_style == "nil"
            else max(
                int(
                    round(
                        max(
                            float(
                                border.width
                            ),
                            0.25,
                        )
                        * 8.0
                    )
                ),
                2,
            )
        )

        element.set(
            qn(
                "w:sz"
            ),
            str(
                border_size
            ),
        )

        element.set(
            qn(
                "w:space"
            ),
            "0",
        )

    # ---------------------------------------------------------
    # Cell content
    # ---------------------------------------------------------

    @classmethod
    def _render_cell_content(
        cls,
        *,
        word_cell,
        editable_cell: EditableTableCell,
    ) -> None:
        cls._remove_cell_paragraphs(
            word_cell
        )

        content_paragraphs = list(
            editable_cell
            .content_paragraphs
            or []
        )

        if not content_paragraphs:
            content_paragraphs = [
                EditableTableCellParagraph(
                    text=str(
                        editable_cell.text
                        or ""
                    ),
                    confidence=(
                        editable_cell
                        .confidence
                    ),
                )
            ]

        for paragraph_index, paragraph_plan in enumerate(
            content_paragraphs
        ):
            word_paragraph = (
                word_cell.add_paragraph()
            )

            cls._apply_paragraph_formatting(
                word_paragraph=word_paragraph,
                editable_cell=editable_cell,
            )

            cls._render_paragraph_content(
                word_paragraph=word_paragraph,
                paragraph_plan=paragraph_plan,
            )

        # Word requires at least one paragraph inside every cell.
        if not word_cell.paragraphs:
            word_cell.add_paragraph()

    @staticmethod
    def _remove_cell_paragraphs(
        word_cell,
    ) -> None:
        for paragraph in list(
            word_cell.paragraphs
        ):
            paragraph._element.getparent().remove(
                paragraph._element
            )

    @staticmethod
    def _apply_paragraph_formatting(
        *,
        word_paragraph,
        editable_cell: EditableTableCell,
    ) -> None:
        alignment_mapping = {
            EditableCellHorizontalAlignment.LEFT: (
                WD_ALIGN_PARAGRAPH.LEFT
            ),
            EditableCellHorizontalAlignment.CENTER: (
                WD_ALIGN_PARAGRAPH.CENTER
            ),
            EditableCellHorizontalAlignment.RIGHT: (
                WD_ALIGN_PARAGRAPH.RIGHT
            ),
            EditableCellHorizontalAlignment.JUSTIFY: (
                WD_ALIGN_PARAGRAPH.JUSTIFY
            ),
            EditableCellHorizontalAlignment.UNKNOWN: (
                WD_ALIGN_PARAGRAPH.LEFT
            ),
        }

        word_paragraph.alignment = (
            alignment_mapping.get(
                editable_cell
                .horizontal_alignment,
                WD_ALIGN_PARAGRAPH.LEFT,
            )
        )

        paragraph_format = (
            word_paragraph
            .paragraph_format
        )

        paragraph_format.space_before = Pt(
            0
        )

        paragraph_format.space_after = Pt(
            0
        )

        paragraph_format.left_indent = Pt(
            0
        )

        paragraph_format.right_indent = Pt(
            0
        )

        paragraph_format.first_line_indent = (
            Pt(
                0
            )
        )

    @classmethod
    def _render_paragraph_content(
        cls,
        *,
        word_paragraph,
        paragraph_plan: EditableTableCellParagraph,
    ) -> None:
        runs = list(
            paragraph_plan.runs
            or []
        )

        if runs:
            for text_run in runs:
                if not str(
                    text_run.text
                    or ""
                ):
                    continue

                word_run = (
                    word_paragraph.add_run(
                        str(
                            text_run.text
                        )
                    )
                )

                cls._apply_text_run_style(
                    word_run=word_run,
                    text_run=text_run,
                )

            if word_paragraph.runs:
                return

        fallback_text = str(
            paragraph_plan.text
            or ""
        )

        if fallback_text:
            word_run = (
                word_paragraph.add_run(
                    fallback_text
                )
            )

            word_run.font.size = Pt(
                cls.DEFAULT_FONT_SIZE
            )

            cls._apply_font_name(
                word_run=word_run,
                pdf_font_name=(
                    cls.DEFAULT_FONT_NAME
                ),
            )

    @classmethod
    def _apply_text_run_style(
        cls,
        *,
        word_run,
        text_run,
    ) -> None:
        font_size = cls._round_font_size(
            getattr(
                text_run,
                "font_size",
                cls.DEFAULT_FONT_SIZE,
            )
        )

        word_run.font.size = Pt(
            font_size
        )

        cls._apply_font_name(
            word_run=word_run,
            pdf_font_name=str(
                getattr(
                    text_run,
                    "font_name",
                    cls.DEFAULT_FONT_NAME,
                )
                or cls.DEFAULT_FONT_NAME
            ),
        )

        color = getattr(
            text_run,
            "color",
            None,
        )

        if (
            color is not None
            and all(
                hasattr(
                    color,
                    channel
                )
                for channel in (
                    "red",
                    "green",
                    "blue",
                )
            )
        ):
            word_run.font.color.rgb = (
                WordRGBColor(
                    int(
                        color.red
                    ),
                    int(
                        color.green
                    ),
                    int(
                        color.blue
                    ),
                )
            )

        word_run.bold = bool(
            getattr(
                text_run,
                "bold",
                False,
            )
        )

        word_run.italic = bool(
            getattr(
                text_run,
                "italic",
                False,
            )
        )

    @staticmethod
    def _round_font_size(
        value: float,
    ) -> float:
        size = max(
            float(
                value
            ),
            0.5,
        )

        return (
            math.floor(
                size * 2.0 + 0.5
            )
            / 2.0
        )

    @staticmethod
    def _apply_font_name(
        *,
        word_run,
        pdf_font_name: str,
    ) -> None:
        word_font_name = (
            FontNameResolver.resolve(
                pdf_font_name
            )
        )

        word_run.font.name = (
            word_font_name
        )

        run_properties = (
            word_run._element
            .get_or_add_rPr()
        )

        font_properties = (
            run_properties
            .get_or_add_rFonts()
        )

        for slot in (
            "ascii",
            "hAnsi",
            "eastAsia",
            "cs",
        ):
            font_properties.set(
                qn(
                    f"w:{slot}"
                ),
                word_font_name,
            )


def cls_points_to_twips(
    value: float,
) -> int:
    """Convert point units to Word twentieths of a point."""

    return max(
        int(
            round(
                float(
                    value
                )
                * 20.0
            )
        ),
        0,
    )
