from __future__ import annotations

from docx import Document as WordDocument
from docx.dml.color import RGBColor as WordRGBColor
from docx.enum.section import WD_SECTION
from docx.shared import Pt

from src.exporter.builders.run_builder import RunBuilder
from src.models.enums.block_type import BlockType

from docx.oxml.ns import qn
from src.exporter.font_name_resolver import (
    FontNameResolver,
)

from collections import Counter
from statistics import median

import re

from docx.oxml import OxmlElement

from src.exporter.word_numbering_manager import (
    WordNumberingManager,
)
from src.models.text_run import TextRun
import math
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.text.paragraph import (
    Paragraph,
)

from src.exporter.editable_layout_resolver import (
    EditableLayoutResolver,
    EditableParagraphPlan,
)

from src.models.page_profile import (
    PageType,
)
from src.models.layout_region import (
    LayoutRegionType,
)

from src.models.paragraph_alignment import (
    ParagraphAlignment,
)

from src.exporter.header_footer_resolver import (
    HeaderFooterResolver,
    PageNumberFieldPlan,
)

from src.models.list_item import (
    ListMarkerSource,
)
from src.exporter.editable_list_sequence_resolver import (
    EditableListSequenceResolver,
)
from src.exporter.editable_page_render_resolver import (
    EditablePageRenderResolver,
    EditableRenderAction,
)

class DocxExporter:
    """
    Exports the analyzed DocuMind document model to DOCX.
    """

    DEFAULT_MARGIN = 36.0
    MINIMUM_MARGIN = 18.0
    
    INLINE_SPACE_MINIMUM_GAP = 0.5
    INLINE_SPACE_GAP_FACTOR = 0.04

    BROKEN_WORD_MINIMUM_PREVIOUS_LENGTH = 3
    BROKEN_WORD_LINE_FILL_RATIO = 0.80

    COMMON_COMPLETE_WORDS = {
        "a",
        "an",
        "and",
        "are",
        "as",
        "at",
        "be",
        "by",
        "for",
        "from",
        "has",
        "he",
        "in",
        "is",
        "it",
        "of",
        "on",
        "or",
        "that",
        "the",
        "this",
        "to",
        "was",
        "were",
        "with",
    }
    
    EDITABLE_RIGHT_MARGIN = 36.0
    EDITABLE_BOTTOM_MARGIN = 36.0

    MAXIMUM_BODY_SPACING_BEFORE = 42.0
    MAXIMUM_COVER_SPACING_BEFORE = 144.0

    CAPTION_SPACING_BEFORE = 12.0

    @staticmethod
    def export(
        document,
        output_path: str,
    ) -> None:
        word_document = WordDocument()

        numbering_manager = (
            WordNumberingManager(
                word_document
            )
        )

        list_sequence_resolver = (
            EditableListSequenceResolver(
                numbering_manager
            )
        )

        pages = list(
            document.pages
        )

        if not pages:
            word_document.save(
                output_path
            )
            return

        page_groups = (
            DocxExporter
            ._group_pages_by_section_template(
                pages
            )
        )

        for group_index, page_group in enumerate(
            page_groups
        ):
            if group_index == 0:
                section = (
                    word_document.sections[0]
                )

            else:
                section = (
                    word_document.add_section(
                        WD_SECTION.NEW_PAGE
                    )
                )

            DocxExporter._configure_page_geometry(
                section=section,
                page=page_group[0],
            )

            DocxExporter._configure_page_group_margins(
                section=section,
                pages=page_group,
            )
            
            DocxExporter._render_section_header_footer(
                section=section,
                page=page_group[0],
                validation_report=getattr(
                    document,
                    "alignment_validation_report",
                    None,
                ),
            )

            for page in page_group:
                DocxExporter._render_page(
                    word_document=word_document,
                    page=page,
                    numbering_manager=(
                        numbering_manager
                    ),
                    list_sequence_resolver=(
                        list_sequence_resolver
                    ),
                    validation_report=getattr(
                        document,
                        "alignment_validation_report",
                        None,
                    ),
                )

        word_document.save(
            output_path
        )

    @staticmethod
    def _render_section_header_footer(
        section,
        page,
        validation_report=None,
    ) -> None:
        """
        Export detected PDF header/footer paragraphs into real
        Word section headers and footers.
        """

        header_footer_plan = (
            HeaderFooterResolver.build(
                page=page,
                validation_report=(
                    validation_report
                ),
            )
        )

        section.header.is_linked_to_previous = (
            False
        )

        section.footer.is_linked_to_previous = (
            False
        )

        DocxExporter._render_story_plans(
            story=section.header,
            paragraph_plans=(
                header_footer_plan.header_plans
            ),
            page_number_fields=(
                header_footer_plan
                .page_number_fields
            ),
        )

        DocxExporter._render_story_plans(
            story=section.footer,
            paragraph_plans=(
                header_footer_plan.footer_plans
            ),
            page_number_fields=(
                header_footer_plan
                .page_number_fields
            ),
        )


    @staticmethod
    def _render_story_plans(
        story,
        paragraph_plans,
        page_number_fields,
    ) -> None:
        """
        Render paragraphs into a Word header or footer story.
        """

        first_paragraph = (
            DocxExporter
            ._reset_story_paragraphs(
                story
            )
        )

        if not paragraph_plans:
            return

        for index, paragraph_plan in enumerate(
            paragraph_plans
        ):
            if index == 0:
                word_paragraph = (
                    first_paragraph
                )

            else:
                word_paragraph = (
                    story.add_paragraph()
                )

            field_plan = (
                page_number_fields.get(
                    paragraph_plan
                    .paragraph_region_number
                )
            )

            if field_plan is not None:
                DocxExporter._render_page_field(
                    word_paragraph=(
                        word_paragraph
                    ),
                    field_plan=field_plan,
                )

            else:
                DocxExporter._render_paragraph_runs(
                    word_paragraph=(
                        word_paragraph
                    ),
                    paragraph=(
                        paragraph_plan.paragraph
                    ),
                    strip_textual_list_marker=False,
                    preserve_line_breaks=False,
                )

            DocxExporter._normalize_alignment_indentation(
                word_paragraph=(
                    word_paragraph
                ),
                paragraph_plan=(
                    paragraph_plan
                ),
            )

            EditableLayoutResolver.apply_alignment(
                word_paragraph=(
                    word_paragraph
                ),
                plan=paragraph_plan,
            )


    @staticmethod
    def _reset_story_paragraphs(
        story,
    ):
        """
        Keep one empty paragraph and remove stale paragraphs from
        a Word header/footer story.
        """

        paragraphs = list(
            story.paragraphs
        )

        if not paragraphs:
            return story.add_paragraph()

        first_paragraph = (
            paragraphs[0]
        )

        first_paragraph.text = ""

        for paragraph in paragraphs[1:]:
            paragraph_element = (
                paragraph._element
            )

            parent = (
                paragraph_element
                .getparent()
            )

            if parent is not None:
                parent.remove(
                    paragraph_element
                )

        return first_paragraph

    @staticmethod
    def _render_page_field(
        word_paragraph,
        field_plan: PageNumberFieldPlan,
    ) -> None:
        """
        Render PAGE and optionally NUMPAGES Word fields.
    
        Examples:
    
            1
            Page 1
            Page 1 of 10
            1 / 10
        """
    
        if field_plan.prefix:
            word_paragraph.add_run(
                field_plan.prefix
            )
    
        DocxExporter._append_word_field(
            word_paragraph=word_paragraph,
            instruction="PAGE",
            placeholder="1",
        )
    
        if field_plan.include_total_pages:
            if field_plan.separator:
                word_paragraph.add_run(
                    field_plan.separator
                )
    
            DocxExporter._append_word_field(
                word_paragraph=word_paragraph,
                instruction="NUMPAGES",
                placeholder="1",
            )
    
        if field_plan.suffix:
            word_paragraph.add_run(
                field_plan.suffix
            )
    
    
    @staticmethod
    def _append_word_field(
        word_paragraph,
        instruction: str,
        placeholder: str,
    ) -> None:
        """
        Append a Word field such as PAGE or NUMPAGES.
        """
    
        run = word_paragraph.add_run()
    
        begin = OxmlElement(
            "w:fldChar"
        )
    
        begin.set(
            qn(
                "w:fldCharType"
            ),
            "begin",
        )
    
        instruction_text = OxmlElement(
            "w:instrText"
        )
    
        instruction_text.set(
            qn(
                "xml:space"
            ),
            "preserve",
        )
    
        instruction_text.text = (
            f" {instruction} "
        )
    
        separate = OxmlElement(
            "w:fldChar"
        )
    
        separate.set(
            qn(
                "w:fldCharType"
            ),
            "separate",
        )
    
        displayed_text = OxmlElement(
            "w:t"
        )
    
        displayed_text.text = (
            placeholder
        )
    
        end = OxmlElement(
            "w:fldChar"
        )
    
        end.set(
            qn(
                "w:fldCharType"
            ),
            "end",
        )
    
        run._r.append(
            begin
        )
    
        run._r.append(
            instruction_text
        )
    
        run._r.append(
            separate
        )
    
        run._r.append(
            displayed_text
        )
    
        run._r.append(
            end
        )

    @staticmethod
    def _group_pages_by_section_template(
        pages,
    ) -> list[list]:
        """
        Group consecutive pages that can share one Word section.

        Pages must have:

            the same size;
            the same rotation;
            the same normalized header;
            the same normalized footer.
        """

        groups: list[list] = []

        previous_key = None

        for page in pages:
            geometry_key = (
                round(
                    float(
                        page.bbox.width
                    ),
                    2,
                ),
                round(
                    float(
                        page.bbox.height
                    ),
                    2,
                ),
                int(
                    getattr(
                        page,
                        "rotation",
                        0,
                    )
                )
                % 360,
            )

            section_key = (
                geometry_key,
                HeaderFooterResolver
                .section_signature(
                    page
                ),
            )

            if (
                previous_key is None
                or section_key
                != previous_key
            ):
                groups.append(
                    [page]
                )

            else:
                groups[-1].append(
                    page
                )

            previous_key = section_key

        return groups

    @staticmethod
    def _configure_page_group_margins(
        section,
        pages,
    ) -> None:
        """
        Configure one editable section using all consecutive
        pages that share the same geometry.

        Page-body containers are preferred because they exclude
        repeated headers and footers.
        """

        body_regions = [
            region
            for page in pages
            for region in getattr(
                page,
                "layout_regions",
                [],
            )
            if (
                region.region_type
                == LayoutRegionType.PAGE_BODY
            )
        ]

        if body_regions:
            left_edge = min(
                region.left
                for region in body_regions
            )

            top_edge = min(
                region.top
                for region in body_regions
            )

        else:
            paragraph_regions = [
                region
                for page in pages
                for region in getattr(
                    page,
                    "paragraph_regions",
                    [],
                )
                if str(
                    getattr(
                        region,
                        "text",
                        "",
                    )
                ).strip()
            ]

            if not paragraph_regions:
                DocxExporter._apply_default_margins(
                    section
                )
                return

            left_edge = min(
                region.left
                for region in paragraph_regions
            )

            top_edge = min(
                region.top
                for region in paragraph_regions
            )

        section.left_margin = Pt(
            max(
                left_edge,
                DocxExporter.MINIMUM_MARGIN,
            )
        )

        section.top_margin = Pt(
            max(
                top_edge,
                DocxExporter.MINIMUM_MARGIN,
            )
        )

        section.right_margin = Pt(
            DocxExporter
            .EDITABLE_RIGHT_MARGIN
        )

        section.bottom_margin = Pt(
            DocxExporter
            .EDITABLE_BOTTOM_MARGIN
        )

        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)
        section.gutter = Pt(0)

    @staticmethod
    def _prepare_page_section(
        word_document,
        page_index: int,
    ):
        """
        Use the existing first section for PDF page 1.

        For each later PDF page, create a new Word section
        beginning on a new page.
        """

        if page_index == 0:
            return word_document.sections[0]

        return word_document.add_section(
            WD_SECTION.NEW_PAGE
        )

    @staticmethod
    def _configure_page_geometry(
        section,
        page,
    ) -> None:
        """
        Apply the PDF page width and height to the Word section.

        PDF coordinates and Word points both use 1/72 inch.
        """

        section.page_width = Pt(
            page.bbox.width
        )

        section.page_height = Pt(
            page.bbox.height
        )

    @staticmethod
    def _normalize_alignment_indentation(
        word_paragraph,
        paragraph_plan,
    ) -> None:
        """
        Remove PDF-position indentation when Word alignment itself
        should control the horizontal placement.

        Without this correction, a centered title can retain a
        large PDF-derived left indent and appear shifted to the
        right even though w:jc is set to center.
        """

        if not paragraph_plan.apply_alignment:
            return

        if (
            paragraph_plan.detected_alignment
            not in {
                ParagraphAlignment.CENTER,
                ParagraphAlignment.RIGHT,
            }
        ):
            return

        paragraph_format = (
            word_paragraph.paragraph_format
        )

        paragraph_format.left_indent = Pt(0)
        paragraph_format.right_indent = Pt(0)
        paragraph_format.first_line_indent = Pt(0)

    @staticmethod
    def _configure_page_margins(
        section,
        page,
    ) -> None:
        """
        Configure margins suitable for a reflowing editable Word
        document.

        Left and top margins follow the PDF content origin.
        Right and bottom margins remain stable so extracted text
        widths do not restrict Word paragraph wrapping.
        """

        regions = [
            region
            for region in page.paragraph_regions
            if region.text.strip()
        ]

        if not regions:
            DocxExporter._apply_default_margins(
                section
            )
            return

        left_edge = min(
            region.left
            for region in regions
        )

        top_edge = min(
            region.top
            for region in regions
        )

        section.left_margin = Pt(
            max(
                left_edge,
                DocxExporter.MINIMUM_MARGIN,
            )
        )

        section.top_margin = Pt(
            max(
                top_edge,
                DocxExporter.MINIMUM_MARGIN,
            )
        )

        # Do not derive these margins from the last extracted
        # character. Editable paragraphs need room to reflow.
        section.right_margin = Pt(
            DocxExporter.EDITABLE_RIGHT_MARGIN
        )

        section.bottom_margin = Pt(
            DocxExporter.EDITABLE_BOTTOM_MARGIN
        )

        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)
        section.gutter = Pt(0)

    @staticmethod
    def _apply_default_margins(section) -> None:
        """
        Apply fallback margins to pages without text blocks.
        """

        margin = Pt(
            DocxExporter.DEFAULT_MARGIN
        )

        section.left_margin = margin
        section.top_margin = margin
        section.right_margin = margin
        section.bottom_margin = margin

    @staticmethod
    def _build_editable_render_plan(
        page,
        validation_report=None,
    ):
        """
        Convert the unified page render plan into editable DOCX
        instructions.

        The wrapper keeps DocxExporter independent from resolver
        implementation details and makes integration easy to test.
        """

        return (
            EditablePageRenderResolver
            .build_page_plan(
                page=page,
                validation_report=(
                    validation_report
                ),
            )
        )

    @staticmethod
    def _render_page(
        word_document,
        page,
        numbering_manager: WordNumberingManager,
        list_sequence_resolver: (
            EditableListSequenceResolver
        ),
        validation_report=None,
    ) -> None:
        """
        Render one PDF page using the unified page render plan.

        Only RENDER_PARAGRAPH instructions are handled in this
        step. Tables, images, charts, vectors and fallback items
        remain deferred for their dedicated exporters.
        """

        editable_render_plan = (
            DocxExporter
            ._build_editable_render_plan(
                page=page,
                validation_report=(
                    validation_report
                ),
            )
        )

        paragraph_instructions = [
            instruction

            for instruction
            in editable_render_plan.instructions

            if (
                instruction.action
                == EditableRenderAction
                .RENDER_PARAGRAPH

                and instruction.layout_item
                is not None

                and str(
                    getattr(
                        instruction.source,
                        "text",
                        "",
                    )
                ).strip()
            )
        ]

        if not paragraph_instructions:
            return

        regions = [
            instruction.source

            for instruction
            in paragraph_instructions
        ]

        content_left = min(
            float(
                region.left
            )
            for region in regions
        )

        content_right = max(
            float(
                region.right
            )
            for region in regions
        )

        legacy_active_list_type = None
        legacy_active_number_id = None

        previous_region = None

        is_designed_cover = (
            getattr(
                getattr(
                    page,
                    "profile",
                    None,
                ),
                "page_type",
                None,
            )
            == PageType.DESIGNED_COVER
        )

        for instruction in (
            editable_render_plan.instructions
        ):
            if (
                instruction.action
                != EditableRenderAction
                .RENDER_PARAGRAPH
            ):
                # These instructions remain available for later
                # table, image, chart, vector and fallback
                # exporters.
                continue

            paragraph_plan = (
                instruction.layout_item
            )

            region = (
                instruction.source
            )

            if paragraph_plan is None:
                continue

            if not str(
                getattr(
                    region,
                    "text",
                    "",
                )
            ).strip():
                continue

            word_paragraph = (
                word_document.add_paragraph()
            )

            is_heading = (
                DocxExporter
                ._region_is_heading(
                    page=page,
                    region=region,
                )
            )

            DocxExporter._apply_region_layout(
                word_paragraph=word_paragraph,
                page=page,
                region=region,
                previous_region=previous_region,
                content_left=content_left,
                content_right=content_right,
                is_heading=is_heading,
            )

            is_list_item = (
                getattr(
                    region,
                    "list_type",
                    None,
                )
                in {
                    "bullet",
                    "number",
                }
            )

            list_sequence_binding = None

            if is_list_item:
                (
                    marker_font_name,
                    marker_font_size,
                ) = (
                    DocxExporter
                    ._resolve_list_marker_style(
                        region
                    )
                )

                list_sequence_binding = (
                    list_sequence_resolver
                    .resolve(
                        page=page,
                        paragraph=region,
                        marker_font_name=(
                            marker_font_name
                        ),
                        marker_font_size=(
                            marker_font_size
                        ),
                    )
                )

                if list_sequence_binding is not None:
                    DocxExporter._apply_list_numbering(
                        word_paragraph=word_paragraph,
                        number_id=(
                            list_sequence_binding
                            .number_id
                        ),
                        level=(
                            list_sequence_binding
                            .level
                        ),
                    )

                    # Numbering-level indentation is already
                    # defined in numbering.xml.
                    DocxExporter._clear_direct_list_indentation(
                        word_paragraph
                    )

                    legacy_active_list_type = None
                    legacy_active_number_id = None

                else:
                    # Backward-compatible list fallback.
                    continues_list = (
                        DocxExporter
                        ._continues_previous_list(
                            previous_region=(
                                previous_region
                            ),
                            current_region=region,
                        )
                    )

                    if (
                        not continues_list

                        or legacy_active_list_type
                        != region.list_type

                        or legacy_active_number_id
                        is None
                    ):
                        start_at = (
                            DocxExporter
                            ._extract_list_start_number(
                                getattr(
                                    region,
                                    "list_marker",
                                    None,
                                )
                            )
                        )

                        legacy_active_number_id = (
                            numbering_manager
                            .create_list(
                                list_type=(
                                    region.list_type
                                ),
                                start_at=start_at,
                                marker_font_name=(
                                    marker_font_name
                                ),
                                marker_font_size=(
                                    marker_font_size
                                ),
                            )
                        )

                        legacy_active_list_type = (
                            region.list_type
                        )

                    DocxExporter._apply_list_numbering(
                        word_paragraph=word_paragraph,
                        number_id=(
                            legacy_active_number_id
                        ),
                        level=getattr(
                            region,
                            "list_level",
                            0,
                        ),
                    )

                    DocxExporter._apply_list_indentation(
                        word_paragraph=word_paragraph,
                        region=region,
                        content_left=content_left,
                    )

            else:
                legacy_active_list_type = None
                legacy_active_number_id = None

            DocxExporter._render_paragraph_runs(
                word_paragraph=word_paragraph,
                paragraph=region,

                strip_textual_list_marker=(
                    (
                        list_sequence_binding
                        .strip_textual_marker
                    )
                    if list_sequence_binding
                    is not None

                    else (
                        getattr(
                            region,
                            "list_marker_source",
                            ListMarkerSource.UNKNOWN,
                        )
                        == ListMarkerSource.TEXT
                    )
                ),

                preserve_line_breaks=(
                    is_designed_cover
                    and is_heading
                ),
            )

            DocxExporter._normalize_alignment_indentation(
                word_paragraph=word_paragraph,
                paragraph_plan=paragraph_plan,
            )

            EditableLayoutResolver.apply_alignment(
                word_paragraph=word_paragraph,
                plan=paragraph_plan,
            )

            previous_region = region

    @staticmethod
    def _resolve_list_marker_style(
        region,
    ) -> tuple[str, float]:
        """
        Use the list content style for the generated Word marker.
    
        PDF bullets are frequently extracted using the legacy
        Symbol font, while the list text uses Arial or another
        Unicode font. Word's generated Unicode bullet should use
        the content font, not the PDF Symbol font.
        """
    
        marker = str(
            getattr(
                region,
                "list_marker",
                "",
            )
            or ""
        ).strip()
    
        fallback_font_size = 11.0
    
        for line in region.lines:
            visible_spans = sorted(
                [
                    span
    
                    for span in line.spans
    
                    if str(
                        span.text
                    ).strip()
                ],
                key=lambda span: (
                    span.left
                ),
            )
    
            for span in visible_spans:
                visible_text = str(
                    span.text
                ).strip()
    
                fallback_font_size = (
                    DocxExporter
                    ._round_word_font_size(
                        span.font_size
                    )
                )
    
                # Skip a span containing only the original marker.
                # The generated Word marker should inherit the
                # content font.
                if (
                    marker
                    and visible_text == marker
                ):
                    continue
                
                return (
                    FontNameResolver.resolve(
                        span.font
                    ),
    
                    DocxExporter
                    ._round_word_font_size(
                        span.font_size
                    ),
                )
    
        return (
            "Arial",
            fallback_font_size,
        )


    @staticmethod
    def _round_word_font_size(
        pdf_font_size: float,
    ) -> float:
        """
        Word supports half-point font-size increments.

        python-docx truncates values such as:

            18.43 → 18.0
            41.48 → 41.0

        Round explicitly before assigning the font size.
        """

        size = max(
            float(pdf_font_size),
            0.5,
        )

        return (
            math.floor(
                size * 2.0 + 0.5
            )
            / 2.0
        )

    @staticmethod
    def _continues_previous_list(
        previous_region,
        current_region,
    ) -> bool:
        """
        Determine whether two consecutive regions belong to the
        same Word list sequence.
        """

        if previous_region is None:
            return False

        if (
            previous_region.list_type
            != current_region.list_type
        ):
            return False

        previous_content_left = (
            previous_region.content_left
            if previous_region.content_left is not None
            else previous_region.left
        )

        current_content_left = (
            current_region.content_left
            if current_region.content_left is not None
            else current_region.left
        )

        if (
            abs(
                previous_content_left
                - current_content_left
            )
            > 18.0
        ):
            return False

        vertical_gap = (
            current_region.top
            - previous_region.bottom
        )

        if vertical_gap > 42.0:
            return False

        if current_region.list_type == "number":
            previous_number = (
                DocxExporter
                ._extract_list_start_number(
                    previous_region.list_marker
                )
            )

            current_number = (
                DocxExporter
                ._extract_list_start_number(
                    current_region.list_marker
                )
            )

            if (
                current_number
                != previous_number + 1
            ):
                return False

        return True

    @staticmethod
    def _extract_list_start_number(
        marker: str | None,
    ) -> int:
        if not marker:
            return 1

        match = re.match(
            r"^\s*(\d+)",
            marker,
        )

        if match is None:
            return 1

        try:
            return max(
                int(
                    match.group(1)
                ),
                1,
            )

        except ValueError:
            return 1

    @staticmethod
    def _apply_list_numbering(
        word_paragraph,
        number_id: int,
        level: int,
    ) -> None:
        """
        Attach a genuine Word numbering definition to a
        paragraph.
        """

        paragraph_properties = (
            word_paragraph
            ._p
            .get_or_add_pPr()
        )

        existing_number_properties = (
            paragraph_properties.find(
                qn("w:numPr")
            )
        )

        if existing_number_properties is not None:
            paragraph_properties.remove(
                existing_number_properties
            )

        number_properties = OxmlElement(
            "w:numPr"
        )

        indentation_level = OxmlElement(
            "w:ilvl"
        )

        indentation_level.set(
            qn("w:val"),
            str(
                max(
                    level,
                    0,
                )
            ),
        )

        number_identifier = OxmlElement(
            "w:numId"
        )

        number_identifier.set(
            qn("w:val"),
            str(number_id),
        )

        number_properties.append(
            indentation_level
        )

        number_properties.append(
            number_identifier
        )

        paragraph_properties.insert(
            0,
            number_properties,
        )

    @staticmethod
    def _apply_list_indentation(
        word_paragraph,
        region,
        content_left: float,
    ) -> None:
        """
        Reconstruct list marker and text positions from the
        original PDF geometry.
        """

        text_left = (
            region.content_left
            if region.content_left is not None
            else region.left
        )

        marker_left = (
            region.list_marker_left
            if region.list_marker_left is not None
            else text_left - 18.0
        )

        left_indent = max(
            text_left - content_left,
            12.0,
        )

        hanging_indent = max(
            text_left - marker_left,
            12.0,
        )

        hanging_indent = min(
            hanging_indent,
            left_indent,
        )

        paragraph_format = (
            word_paragraph.paragraph_format
        )

        paragraph_format.left_indent = Pt(
            left_indent
        )

        paragraph_format.first_line_indent = Pt(
            -hanging_indent
        )
        
        paragraph_format.tab_stops.add_tab_stop(
            Pt(left_indent),
            WD_TAB_ALIGNMENT.LEFT,
        )

    @staticmethod
    def _clear_direct_list_indentation(
        word_paragraph,
    ) -> None:
        """
        Remove direct paragraph indentation for sequence-based
        lists.

        The Word numbering level definition supplies left and
        hanging indentation. Keeping PDF-derived direct
        indentation here can override the numbering definition and
        produce doubled or incorrectly shifted indentation.
        """

        paragraph_format = (
            word_paragraph.paragraph_format
        )

        paragraph_format.left_indent = None
        paragraph_format.first_line_indent = None

    @staticmethod
    def _apply_region_layout(
        word_paragraph,
        page,
        region,
        previous_region,
        content_left: float,
        content_right: float,
        is_heading: bool,
    ) -> None:
        """
        Reconstruct paragraph indentation, vertical spacing,
        alignment and line spacing using the PDF coordinates.
        """

        paragraph_format = (
            word_paragraph.paragraph_format
        )

        left_indent = max(
            region.left - content_left,
            0.0,
        )

        first_line_indent = (
            DocxExporter
            ._calculate_first_line_indent(
                region
            )
        )

        paragraph_format.left_indent = Pt(
            left_indent
        )

        # The extracted right edge represents the original glyph
        # width, not the intended Word paragraph boundary.
        paragraph_format.right_indent = Pt(0)

        paragraph_format.first_line_indent = Pt(
            first_line_indent
        )

        if previous_region is None:
            spacing_before = 0.0

        else:
            raw_spacing_before = max(
                region.top
                - previous_region.bottom,
                0.0,
            )

            if DocxExporter._region_is_caption(
                region
            ):
                # Tables and charts will later be inserted before their
                # captions. Do not preserve their empty PDF area as
                # paragraph spacing.
                spacing_before = (
                    DocxExporter
                    .CAPTION_SPACING_BEFORE
                )

            else:
                is_designed_cover = (
                    getattr(
                        getattr(
                            page,
                            "profile",
                            None,
                        ),
                        "page_type",
                        None,
                    )
                    == PageType.DESIGNED_COVER
                )

                maximum_spacing = (
                    DocxExporter
                    .MAXIMUM_COVER_SPACING_BEFORE

                    if is_designed_cover

                    else DocxExporter
                    .MAXIMUM_BODY_SPACING_BEFORE
                )

                spacing_before = min(
                    raw_spacing_before,
                    maximum_spacing,
                )

        paragraph_format.space_before = Pt(
            spacing_before
        )

        paragraph_format.space_after = Pt(0)

        line_spacing = (
            DocxExporter
            ._estimate_region_line_spacing(
                region
            )
        )

        # Editable Word paragraphs should reflow naturally.
        # Exact PDF line spacing can clip text after editing or font
        # substitution.
        paragraph_format.line_spacing_rule = (
            WD_LINE_SPACING.SINGLE
        )

        paragraph_format.line_spacing = 1.0

        paragraph_format.keep_together = True
        paragraph_format.widow_control = False

        if is_heading:
            paragraph_format.keep_with_next = True

    @staticmethod
    def _region_is_caption(
        region,
    ) -> bool:
        """
        Detect table and figure caption paragraphs.
        """

        text = region.text.strip()

        return bool(
            re.match(
                r"^\(\s*(?:Table|Figure)\b",
                text,
                flags=re.IGNORECASE,
            )
        )

    @staticmethod
    def _calculate_first_line_indent(
        region,
    ) -> float:
        """
        Calculate first-line indentation from the first visible
        line of the region.
        """

        for line in region.lines:
            visible_spans = [
                span
                for span in line.spans
                if span.text.strip()
            ]

            if not visible_spans:
                continue

            first_line_left = min(
                span.left
                for span in visible_spans
            )

            return (
                first_line_left
                - region.left
            )

        return 0.0

    @staticmethod
    def _estimate_region_line_spacing(
        region,
    ) -> float:
        """
        Estimate line spacing using the original PDF line
        positions.

        For multiline regions, baseline movement between lines is
        preferred. For single-line regions, the extracted text
        height is used.
        """

        visible_lines = []

        for line in region.lines:
            visible_spans = [
                span
                for span in line.spans
                if span.text.strip()
            ]

            if not visible_spans:
                continue

            line_top = min(
                span.top
                for span in visible_spans
            )

            line_bottom = max(
                span.bottom
                for span in visible_spans
            )

            visible_lines.append({
                "top": line_top,
                "bottom": line_bottom,
                "height": max(
                    line_bottom - line_top,
                    1.0,
                ),
            })

        if not visible_lines:
            return 12.0

        line_advances = [
            current["top"] - previous["top"]

            for previous, current in zip(
                visible_lines,
                visible_lines[1:],
            )

            if (
                current["top"]
                - previous["top"]
            ) > 1.0
        ]

        if line_advances:
            return max(
                float(
                    median(
                        line_advances
                    )
                ),
                1.0,
            )

        line_heights = [
            line["height"]
            for line in visible_lines
        ]

        return max(
            float(
                median(
                    line_heights
                )
            ),
            1.0,
        )

    @staticmethod
    def _region_is_heading(
        page,
        region,
    ) -> bool:
        """
        Determine whether a paragraph region originated from a
        heading or subtitle block.
        """
    
        source_numbers = set(
            region.source_block_numbers
        )
    
        return any(
            block.block_number
            in source_numbers
    
            and block.block_type
            in {
                BlockType.HEADING,
                BlockType.SUBTITLE,
            }
    
            for block in page.blocks
        )

    @staticmethod
    def _create_word_paragraph(
        word_document,
        block_type: BlockType,
    ):
        """
        Create a Word heading or normal paragraph.
        """

        if block_type == BlockType.HEADING:
            return word_document.add_heading(
                level=1
            )

        return word_document.add_paragraph()

    @staticmethod
    def _apply_paragraph_indentation(
        word_paragraph,
        paragraph,
    ) -> None:
        """
        Apply horizontal paragraph geometry reconstructed
        from the PDF.
        """

        paragraph_format = (
            word_paragraph.paragraph_format
        )

        paragraph_format.left_indent = Pt(
            paragraph.style.left_indent
        )

        paragraph_format.right_indent = Pt(
            paragraph.style.right_indent
        )

        paragraph_format.first_line_indent = Pt(
            paragraph.style.first_line_indent
        )

    @staticmethod
    def _apply_paragraph_spacing(
        word_paragraph,
        paragraph,
    ) -> None:
        """
        Apply reconstructed paragraph and line spacing.
        """

        paragraph_format = (
            word_paragraph.paragraph_format
        )

        paragraph_format.space_before = Pt(
            paragraph.style.spacing_before
        )

        paragraph_format.space_after = Pt(
            paragraph.style.spacing_after
        )

        if paragraph.style.line_spacing > 0:
            paragraph_format.line_spacing_rule = (
                WD_LINE_SPACING.EXACTLY
            )

            paragraph_format.line_spacing = Pt(
                paragraph.style.line_spacing
            )

    @classmethod
    def _render_paragraph_runs(
        cls,
        word_paragraph,
        paragraph,
        strip_textual_list_marker: bool = False,
        preserve_line_breaks: bool = False,
    ) -> None:
        """
        Render an editable Word paragraph.

        Normal PDF line boundaries are converted to spaces so
        Word can reflow the paragraph. Intentional cover-page
        line breaks may be preserved.
        """

        run_plan = cls._build_reflow_run_plan(
            paragraph=paragraph,
            strip_textual_list_marker=(
                strip_textual_list_marker
            ),
            preserve_line_breaks=(
                preserve_line_breaks
            ),
        )

        for item in run_plan:
            if item is None:
                break_run = (
                    word_paragraph.add_run()
                )

                break_run.add_break(
                    WD_BREAK.LINE
                )

                continue

            if not item.text:
                continue

            run = word_paragraph.add_run(
                item.text
            )

            font = run.font

            font.size = Pt(
                cls._round_word_font_size(
                    item.font_size
                )
            )

            cls._apply_font_name(
                run=run,
                pdf_font_name=item.font_name,
            )

            font.color.rgb = WordRGBColor(
                item.color.red,
                item.color.green,
                item.color.blue,
            )

            run.bold = item.bold
            run.italic = item.italic
    
    @classmethod
    def _build_reflow_run_plan(
        cls,
        paragraph,
        strip_textual_list_marker: bool,
        preserve_line_breaks: bool,
    ) -> list[TextRun | None]:
        """
        Build formatted runs for one editable paragraph.

        Body-page line endings become normal spaces. Page 1 may
        retain intentional cover-layout line breaks.
        """

        visible_lines = [
            line
            for line in paragraph.lines
            if cls._line_has_text(line)
        ]

        plan: list[
            TextRun | None
        ] = []

        marker_remaining = (
            paragraph.list_marker
            if strip_textual_list_marker
            else None
        )

        previous_line = None

        for line_index, line in enumerate(
            visible_lines
        ):
            line_runs = (
                cls._build_editable_line_runs(
                    line
                )
            )

            if not line_runs:
                continue

            if (
                line_index == 0
                and marker_remaining
            ):
                marker_remaining = (
                    cls._strip_marker_from_text_runs(
                        runs=line_runs,
                        marker_remaining=(
                            marker_remaining
                        ),
                    )
                )

                line_runs = [
                    run
                    for run in line_runs
                    if run.text
                ]

            if not line_runs:
                previous_line = line
                continue

            if plan and previous_line is not None:
                if preserve_line_breaks:
                    plan.append(None)
                else:
                    cls._append_reflow_boundary(
                        plan=plan,
                        next_runs=line_runs,
                        previous_line=previous_line,
                        current_line=line,
                        paragraph=paragraph,
                    )

            for text_run in line_runs:
                cls._append_text_run(
                    plan=plan,
                    text_run=text_run,
                )

            previous_line = line

        return plan
    
    @classmethod
    def _build_editable_line_runs(
        cls,
        line,
    ) -> list[TextRun]:
        """
        Convert PDF spans into Word runs while restoring spaces
        between separately formatted spans.
        """

        spans = sorted(
            [
                span
                for span in line.spans
                if span.text
            ],
            key=lambda span: span.left,
        )

        runs: list[TextRun] = []

        previous_span = None
        previous_raw_text = None

        for span in spans:
            raw_text = (
                cls._normalize_inline_text(
                    span.text
                )
            )

            visible_text = (
                raw_text.strip()
            )

            if not visible_text:
                continue

            if (
                previous_span is not None
                and cls._needs_space_between_spans(
                    previous_span=previous_span,
                    current_span=span,
                    previous_raw_text=(
                        previous_raw_text or ""
                    ),
                    current_raw_text=raw_text,
                )
            ):
                cls._append_text_run(
                    plan=runs,
                    text_run=cls._text_run_from_span(
                        span=previous_span,
                        text=" ",
                    ),
                )

            cls._append_text_run(
                plan=runs,
                text_run=cls._text_run_from_span(
                    span=span,
                    text=visible_text,
                ),
            )

            previous_span = span
            previous_raw_text = raw_text

        return runs
    
    @staticmethod
    def _normalize_inline_text(
        text: str,
    ) -> str:
        """
        Normalize PDF whitespace without removing meaningful
        spaces between words.
        """

        normalized = (
            text
            .replace(
                "\u00a0",
                " ",
            )
            .replace(
                "\u2007",
                " ",
            )
            .replace(
                "\u202f",
                " ",
            )
        )

        normalized = re.sub(
            r"[\t\r\n\f\v]+",
            " ",
            normalized,
        )

        normalized = re.sub(
            r" {2,}",
            " ",
            normalized,
        )

        return normalized
    
    @classmethod
    def _needs_space_between_spans(
        cls,
        previous_span,
        current_span,
        previous_raw_text: str,
        current_raw_text: str,
    ) -> bool:
        """
        Restore a missing space between adjacent PDF spans.

        This commonly occurs where bold text changes to regular
        text.
        """

        if not previous_raw_text or not current_raw_text:
            return False

        if (
            previous_raw_text[-1:].isspace()
            or current_raw_text[:1].isspace()
        ):
            return True

        previous_visible = (
            previous_raw_text.rstrip()
        )

        current_visible = (
            current_raw_text.lstrip()
        )

        if not previous_visible or not current_visible:
            return False

        if current_visible[0] in {
            ".",
            ",",
            ";",
            ":",
            "!",
            "?",
            "%",
            ")",
            "]",
            "}",
        }:
            return False

        if previous_visible[-1] in {
            "(",
            "[",
            "{",
            "/",
            "\\",
            "–",
            "—",
        }:
            return False

        horizontal_gap = (
            current_span.left
            - previous_span.right
        )

        if horizontal_gap <= 0:
            return False

        reference_font_size = max(
            min(
                previous_span.font_size,
                current_span.font_size,
            ),
            1.0,
        )

        threshold = max(
            cls.INLINE_SPACE_MINIMUM_GAP,
            (
                reference_font_size
                * cls.INLINE_SPACE_GAP_FACTOR
            ),
        )

        return horizontal_gap >= threshold
    
    @classmethod
    def _append_reflow_boundary(
        cls,
        plan: list[TextRun | None],
        next_runs: list[TextRun],
        previous_line,
        current_line,
        paragraph,
    ) -> None:
        """
        Join two visual PDF lines into one editable Word
        paragraph.

        Normally a space is inserted. A conservative broken-word
        rule joins clear one-letter fragments without a space.
        """

        previous_run = next(
            (
                item
                for item in reversed(plan)
                if isinstance(
                    item,
                    TextRun,
                )
            ),
            None,
        )

        if previous_run is None or not next_runs:
            return

        first_next_run = next_runs[0]

        previous_run.text = (
            previous_run.text.rstrip()
        )

        first_next_run.text = (
            first_next_run.text.lstrip()
        )

        if not previous_run.text:
            return

        if not first_next_run.text:
            return

        broken_prefix = (
            cls._extract_broken_word_prefix(
                previous_run=previous_run,
                next_run=first_next_run,
                previous_line=previous_line,
                paragraph=paragraph,
            )
        )

        if broken_prefix is not None:
            # Append the fragment without a space:
            #
            # dapibu + s → dapibus
            # ant + e    → ante
            prefix_run = cls._copy_text_run(
                source=first_next_run,
                text=broken_prefix,
            )

            cls._append_text_run(
                plan=plan,
                text_run=prefix_run,
            )

            remainder = (
                first_next_run.text[
                    len(broken_prefix):
                ]
                .lstrip()
            )

            if remainder:
                first_next_run.text = (
                    " " + remainder
                )
            else:
                next_runs.pop(0)

                if next_runs:
                    next_runs[0].text = (
                        " "
                        + next_runs[0]
                        .text
                        .lstrip()
                    )

            return

        # Preserve an explicit line-ending hyphen.
        if previous_run.text.endswith("-"):
            return

        if first_next_run.text[0] in {
            ".",
            ",",
            ";",
            ":",
            "!",
            "?",
            "%",
            ")",
            "]",
            "}",
        }:
            return

        cls._append_text_run(
            plan=plan,
            text_run=cls._copy_text_run(
                source=previous_run,
                text=" ",
            ),
        )
    
    @classmethod
    def _extract_broken_word_prefix(
        cls,
        previous_run: TextRun,
        next_run: TextRun,
        previous_line,
        paragraph,
    ) -> str | None:
        """
        Detect obvious one-character fragments moved onto the
        next PDF line.

        The rule is intentionally conservative because a normal
        line may legitimately begin with words such as "a".
        """

        previous_match = re.search(
            r"([A-Za-z]+)$",
            previous_run.text,
        )

        next_match = re.match(
            r"^([a-z])(?=\s|$)",
            next_run.text,
        )

        if (
            previous_match is None
            or next_match is None
        ):
            return None

        previous_word = (
            previous_match
            .group(1)
        )

        if (
            len(previous_word)
            < cls.BROKEN_WORD_MINIMUM_PREVIOUS_LENGTH
        ):
            return None

        if (
            previous_word.lower()
            in cls.COMMON_COMPLETE_WORDS
        ):
            return None

        visible_spans = [
            span
            for span in previous_line.spans
            if span.text.strip()
        ]

        if not visible_spans:
            return None

        previous_line_right = max(
            span.right
            for span in visible_spans
        )

        available_width = max(
            paragraph.right
            - paragraph.left,
            1.0,
        )

        used_width_ratio = (
            previous_line_right
            - paragraph.left
        ) / available_width

        if (
            used_width_ratio
            < cls.BROKEN_WORD_LINE_FILL_RATIO
        ):
            return None

        return next_match.group(1)
    
    @classmethod
    def _strip_marker_from_text_runs(
        cls,
        runs: list[TextRun],
        marker_remaining: str,
    ) -> str | None:
        """
        Remove textual numbering because Word now generates the
        list number itself.
        """

        remaining: str | None = (
            marker_remaining
        )

        for text_run in runs:
            if not remaining:
                break

            (
                cleaned_text,
                remaining,
            ) = cls._remove_textual_list_marker(
                text=text_run.text,
                marker_remaining=remaining,
            )

            text_run.text = cleaned_text

        return remaining
    
    @staticmethod
    def _text_run_from_span(
        span,
        text: str,
    ) -> TextRun:
        font_name_lower = (
            span.font or ""
        ).lower()

        bold = (
            bool(
                span.flags
                & (1 << 4)
            )
            or "bold" in font_name_lower
            or "semibold" in font_name_lower
            or "extrabold" in font_name_lower
        )

        italic = (
            bool(
                span.flags
                & (1 << 1)
            )
            or "italic" in font_name_lower
            or "oblique" in font_name_lower
        )

        return TextRun(
            text=text,
            font_name=span.font,
            font_size=span.font_size,
            color=span.color,
            bold=bold,
            italic=italic,
        )


    @staticmethod
    def _copy_text_run(
        source: TextRun,
        text: str,
    ) -> TextRun:
        return TextRun(
            text=text,
            font_name=source.font_name,
            font_size=source.font_size,
            color=source.color,
            bold=source.bold,
            italic=source.italic,
        )
    
    @classmethod
    def _append_text_run(
        cls,
        plan,
        text_run: TextRun,
    ) -> None:
        """
        Add a run, merging it with the previous run when both
        have identical formatting.
        """

        if not text_run.text:
            return

        previous = (
            plan[-1]
            if plan
            else None
        )

        if (
            isinstance(
                previous,
                TextRun,
            )
            and cls._text_runs_have_same_style(
                previous,
                text_run,
            )
        ):
            previous.text += (
                text_run.text
            )
            return

        plan.append(
            text_run
        )


    @staticmethod
    def _text_runs_have_same_style(
        first: TextRun,
        second: TextRun,
    ) -> bool:
        return (
            first.font_name
            == second.font_name

            and abs(
                first.font_size
                - second.font_size
            ) <= 0.01

            and first.color
            == second.color

            and first.bold
            == second.bold

            and first.italic
            == second.italic
        )
    
    @staticmethod
    def _remove_textual_list_marker(
        text: str,
        marker_remaining: str,
    ) -> tuple[str, str | None]:
        """
        Remove a textual marker such as "1." because Word now
        generates the number itself.
        """

        if not text:
            return text, marker_remaining

        leading_length = (
            len(text)
            - len(text.lstrip())
        )

        visible_text = text.lstrip()

        if visible_text.startswith(
            marker_remaining
        ):
            visible_text = visible_text[
                len(marker_remaining):
            ]

            return (
                visible_text.lstrip(),
                None,
            )

        if marker_remaining.startswith(
            visible_text
        ):
            remaining = marker_remaining[
                len(visible_text):
            ]

            return (
                "",
                remaining or None,
            )

        # Marker and content may have been combined unexpectedly.
        marker_pattern = re.compile(
            r"^\s*"
            + re.escape(
                marker_remaining
            )
            + r"\s*"
        )

        cleaned_text = marker_pattern.sub(
            "",
            text,
            count=1,
        )

        if cleaned_text != text:
            return cleaned_text, None

        return text, marker_remaining
                
    @staticmethod
    def _apply_font_name(
        run,
        pdf_font_name: str,
    ) -> None:
        """
        Resolve and apply a PDF font name to every Word
        font slot.

        Setting only run.font.name is sometimes insufficient
        because Word stores separate font names for different
        character categories.
        """

        word_font_name = FontNameResolver.resolve(
            pdf_font_name
        )

        run.font.name = word_font_name

        run_properties = run._element.get_or_add_rPr()

        font_properties = run_properties.get_or_add_rFonts()

        font_properties.set(
            qn("w:ascii"),
            word_font_name,
        )

        font_properties.set(
            qn("w:hAnsi"),
            word_font_name,
        )

        font_properties.set(
            qn("w:eastAsia"),
            word_font_name,
        )

        font_properties.set(
            qn("w:cs"),
            word_font_name,
        )

    @staticmethod
    def _line_has_text(line) -> bool:
        """
        Return True when a PDF line contains visible text.
        """

        return any(
            span.text.strip()
            for span in line.spans
        )

    @staticmethod
    def _block_has_text(block) -> bool:
        """
        Return True when a block contains visible text.
        """

        return any(
            span.text.strip()
            for line in block.lines
            for span in line.spans
        )